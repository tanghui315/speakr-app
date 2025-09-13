# Speakr - Audio Transcription and Summarization App
import os
import sys
from flask import Flask, render_template, request, jsonify, send_file, redirect, url_for, flash, Response, make_response
from urllib.parse import urlparse, urljoin, quote
from email.utils import encode_rfc2231
try:
    from flask import Markup
except ImportError:
    from markupsafe import Markup
from flask_sqlalchemy import SQLAlchemy
from datetime import datetime, timedelta
from openai import OpenAI # Keep using the OpenAI library
import json
from werkzeug.utils import secure_filename
from werkzeug.exceptions import RequestEntityTooLarge
from werkzeug.middleware.proxy_fix import ProxyFix
from sqlalchemy import select
from sqlalchemy.orm import joinedload
import threading
from dotenv import load_dotenv # Import load_dotenv
import httpx 
import re
import subprocess
import mimetypes
import markdown
import bleach

# Add common audio MIME type mappings that might be missing
mimetypes.add_type('audio/mp4', '.m4a')
mimetypes.add_type('audio/aac', '.aac')
mimetypes.add_type('audio/x-m4a', '.m4a')
mimetypes.add_type('audio/webm', '.webm')
mimetypes.add_type('audio/flac', '.flac')
mimetypes.add_type('audio/ogg', '.ogg')
from flask_login import LoginManager, UserMixin, login_user, logout_user, login_required, current_user
from flask_bcrypt import Bcrypt
from flask_wtf import FlaskForm
from flask_wtf.csrf import CSRFProtect
from wtforms import StringField, PasswordField, SubmitField, BooleanField
from wtforms.validators import DataRequired, Length, Email, EqualTo, ValidationError
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
import pytz
from babel.dates import format_datetime
import ast
import logging
import secrets
import time
from src.audio_chunking import AudioChunkingService, ChunkProcessingError, ChunkingNotSupportedError

# Optional imports for embedding functionality
try:
    import numpy as np
    from sentence_transformers import SentenceTransformer
    from sklearn.metrics.pairwise import cosine_similarity
    EMBEDDINGS_AVAILABLE = True
except ImportError as e:
    EMBEDDINGS_AVAILABLE = False
    # Create dummy classes to prevent import errors
    class SentenceTransformer:
        def __init__(self, *args, **kwargs):
            pass
        def encode(self, *args, **kwargs):
            return []
    
    np = None
    cosine_similarity = None

# Load environment variables from .env file
load_dotenv()

# Early check for Inquire Mode configuration (needed for startup message)
ENABLE_INQUIRE_MODE = os.environ.get('ENABLE_INQUIRE_MODE', 'false').lower() == 'true'

# Log embedding status on startup
if ENABLE_INQUIRE_MODE and EMBEDDINGS_AVAILABLE:
    print("✅ Inquire Mode: Full semantic search enabled (embeddings available)")
elif ENABLE_INQUIRE_MODE and not EMBEDDINGS_AVAILABLE:
    print("⚠️  Inquire Mode: Basic text search only (embedding dependencies not available)")
    print("   To enable semantic search, install: pip install sentence-transformers==2.7.0 huggingface-hub>=0.19.0")
elif not ENABLE_INQUIRE_MODE:
    print("ℹ️  Inquire Mode: Disabled (set ENABLE_INQUIRE_MODE=true to enable)")

# Configure logging
log_level = os.environ.get('LOG_LEVEL', 'INFO').upper()
handler = logging.StreamHandler(sys.stdout)
handler.setLevel(log_level)
formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
handler.setFormatter(formatter)

# Get the root logger
root_logger = logging.getLogger()
root_logger.setLevel(log_level)
root_logger.addHandler(handler)

# Also configure Flask's logger
app_logger = logging.getLogger('werkzeug')
app_logger.setLevel(log_level)
app_logger.addHandler(handler)

# --- Rate Limiting Setup (will be configured after app creation) ---
# TEMPORARILY INCREASED FOR TESTING - REVERT FOR PRODUCTION!
limiter = Limiter(
    get_remote_address,
    app=None,  # Defer initialization
    default_limits=["5000 per day", "1000 per hour"]  # Increased from 200/day, 50/hour for testing
)

def auto_close_json(json_string):
    """
    Attempts to close an incomplete JSON string by appending necessary brackets and braces.
    This is a simplified parser and may not handle all edge cases, but is
    designed to fix unterminated strings from API responses.
    """
    if not isinstance(json_string, str):
        return json_string

    stack = []
    in_string = False
    escape_next = False

    for char in json_string:
        if escape_next:
            escape_next = False
            continue

        if char == '\\':
            escape_next = True
            continue

        if char == '"':
            # We don't handle escaped quotes inside strings perfectly,
            # but this is a simple heuristic.
            if not escape_next:
                in_string = not in_string

        if not in_string:
            if char == '{':
                stack.append('}')
            elif char == '[':
                stack.append(']')
            elif char == '}':
                if stack and stack[-1] == '}':
                    stack.pop()
            elif char == ']':
                if stack and stack[-1] == ']':
                    stack.pop()

    # If we are inside a string at the end, close it.
    if in_string:
        json_string += '"'
    
    # Close any remaining open structures
    while stack:
        json_string += stack.pop()

    return json_string

def safe_json_loads(json_string, fallback_value=None):
    """
    Safely parse JSON with preprocessing to handle common LLM JSON formatting issues.
    
    Args:
        json_string (str): The JSON string to parse
        fallback_value: Value to return if parsing fails (default: None)
    
    Returns:
        Parsed JSON object or fallback_value if parsing fails
    """
    if not json_string or not isinstance(json_string, str):
        app.logger.warning(f"Invalid JSON input: {type(json_string)} - {json_string}")
        return fallback_value
    
    # Step 1: Clean the input string
    cleaned_json = json_string.strip()
    
    # Step 2: Extract JSON from markdown code blocks if present
    json_match = re.search(r'```(?:json)?\s*(.*?)\s*```', cleaned_json, re.DOTALL)
    if json_match:
        cleaned_json = json_match.group(1).strip()
    
    # Step 3: Try multiple parsing strategies
    parsing_strategies = [
        # Strategy 1: Direct parsing (for well-formed JSON)
        lambda x: json.loads(x),
        
        # Strategy 2: Fix common escape issues
        lambda x: json.loads(preprocess_json_escapes(x)),
        
        # Strategy 3: Use ast.literal_eval as fallback for simple cases
        lambda x: ast.literal_eval(x) if x.startswith(('{', '[')) else None,
        
        # Strategy 4: Extract JSON object/array using regex
        lambda x: json.loads(extract_json_object(x)),
        
        # Strategy 5: Auto-close incomplete JSON and parse
        lambda x: json.loads(auto_close_json(x)),
    ]
    
    for i, strategy in enumerate(parsing_strategies):
        try:
            result = strategy(cleaned_json)
            if result is not None:
                if i > 0:  # Log if we had to use a fallback strategy
                    app.logger.info(f"JSON parsed successfully using strategy {i+1}")
                return result
        except (json.JSONDecodeError, ValueError, SyntaxError) as e:
            if i == 0:  # Only log the first failure to avoid spam
                app.logger.debug(f"JSON parsing strategy {i+1} failed: {e}")
            continue
    
    # All strategies failed
    app.logger.error(f"All JSON parsing strategies failed for: {cleaned_json[:200]}...")
    return fallback_value

def preprocess_json_escapes(json_string):
    """
    Preprocess JSON string to fix common escape issues from LLM responses.
    Uses a more sophisticated approach to handle nested quotes properly.
    """
    if not json_string:
        return json_string
    
    result = []
    i = 0
    in_string = False
    escape_next = False
    expecting_value = False  # Track if we're expecting a value (after :)
    
    while i < len(json_string):
        char = json_string[i]
        
        if escape_next:
            # This character is escaped, add it as-is
            result.append(char)
            escape_next = False
        elif char == '\\':
            # This is an escape character
            result.append(char)
            escape_next = True
        elif char == ':' and not in_string:
            # We found a colon, next string will be a value
            result.append(char)
            expecting_value = True
        elif char == ',' and not in_string:
            # We found a comma, reset expecting_value
            result.append(char)
            expecting_value = False
        elif char == '"':
            if not in_string:
                # Starting a string
                in_string = True
                result.append(char)
            else:
                # We're in a string, check if this quote should be escaped
                # Look ahead to see if this is the end of the string value
                j = i + 1
                while j < len(json_string) and json_string[j].isspace():
                    j += 1
                
                # For keys (not expecting_value), only end on colon
                # For values (expecting_value), end on comma, closing brace, or closing bracket
                if expecting_value:
                    end_chars = ',}]'
                else:
                    end_chars = ':'
                
                if j < len(json_string) and json_string[j] in end_chars:
                    # This is the end of the string
                    in_string = False
                    result.append(char)
                    if not expecting_value:
                        # We just finished a key, next will be expecting value
                        expecting_value = True
                else:
                    # This is an inner quote that should be escaped
                    result.append('\\"')
        else:
            result.append(char)
        
        i += 1
    
    return ''.join(result)

def extract_json_object(text):
    """
    Extract the first complete JSON object or array from text using regex.
    """
    # Look for JSON object
    obj_match = re.search(r'\{.*\}', text, re.DOTALL)
    if obj_match:
        return obj_match.group(0)
    
    # Look for JSON array
    arr_match = re.search(r'\[.*\]', text, re.DOTALL)
    if arr_match:
        return arr_match.group(0)
    
    # Return original if no JSON structure found
    return text

# Initialize Flask-Bcrypt
bcrypt = Bcrypt()

# Helper function to convert markdown to HTML
def sanitize_html(text):
    """
    Sanitize HTML content to prevent XSS attacks and code execution.
    This function removes dangerous content while preserving safe formatting.
    """
    if not text:
        return ""
    
    # First, remove any template syntax that could be dangerous
    # Remove Jinja2/Flask template syntax
    text = re.sub(r'\{\{.*?\}\}', '', text, flags=re.DOTALL)
    text = re.sub(r'\{%.*?%\}', '', text, flags=re.DOTALL)
    
    # Remove other template-like syntax
    text = re.sub(r'<%.*?%>', '', text, flags=re.DOTALL)
    text = re.sub(r'<\?.*?\?>', '', text, flags=re.DOTALL)
    
    # Define allowed tags and attributes for safe HTML
    allowed_tags = [
        'p', 'br', 'strong', 'b', 'em', 'i', 'u', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6',
        'ul', 'ol', 'li', 'blockquote', 'code', 'pre', 'a', 'img', 'table', 'thead', 
        'tbody', 'tr', 'th', 'td', 'dl', 'dt', 'dd', 'div', 'span', 'hr', 'sup', 'sub'
    ]
    
    allowed_attributes = {
        'a': ['href', 'title'],
        'img': ['src', 'alt', 'title', 'width', 'height'],
        'code': ['class'],  # For syntax highlighting
        'pre': ['class'],   # For syntax highlighting
        'div': ['class'],   # For code blocks
        'span': ['class'],  # For syntax highlighting
        'th': ['align'],
        'td': ['align'],
        'table': ['class']
    }
    
    # Sanitize the HTML to remove dangerous content
    sanitized_html = bleach.clean(
        text,
        tags=allowed_tags,
        attributes=allowed_attributes,
        protocols=['http', 'https', 'mailto'],
        strip=True  # Strip disallowed tags instead of escaping them
    )
    
    return sanitized_html

def md_to_html(text):
    if not text:
        return ""
    
    # Fix list spacing
    def fix_list_spacing(text):
        lines = text.split('\n')
        result = []
        in_list = False
        
        for line in lines:
            stripped = line.strip()
            
            # Check if this line is a list item (starts with -, *, +, or number.)
            is_list_item = (
                stripped.startswith(('- ', '* ', '+ ')) or
                (stripped and stripped[0].isdigit() and '. ' in stripped[:10])
            )
            
            # If we're starting a new list or continuing a list, ensure proper spacing
            if is_list_item:
                if not in_list and result and result[-1].strip():
                    # Starting a new list - add blank line before
                    result.append('')
                in_list = True
            elif in_list and stripped and not is_list_item:
                # Ending a list - add blank line after the list
                if result and result[-1].strip():
                    result.append('')
                in_list = False
            
            result.append(line)
        
        return '\n'.join(result)
    
    # Fix list spacing
    processed_text = fix_list_spacing(text)
    
    # Convert markdown to HTML with extensions for tables, code highlighting, etc.
    html = markdown.markdown(processed_text, extensions=[
        'fenced_code',      # Fenced code blocks
        'tables',           # Table support
        'attr_list',        # Attribute lists
        'def_list',         # Definition lists
        'footnotes',        # Footnotes
        'abbr',             # Abbreviations
        'codehilite',       # Syntax highlighting for code blocks
        'smarty'            # Smart quotes, dashes, etc.
    ])
    
    # Apply sanitization to the generated HTML
    return sanitize_html(html)

def format_transcription_for_llm(transcription_text):
    """
    Formats transcription for LLM. If it's our simplified JSON, convert it to plain text.
    Otherwise, return as is.
    """
    try:
        transcription_data = json.loads(transcription_text)
        if isinstance(transcription_data, list):
            # It's our simplified JSON format
            formatted_lines = []
            for segment in transcription_data:
                speaker = segment.get('speaker', 'Unknown Speaker')
                sentence = segment.get('sentence', '')
                formatted_lines.append(f"[{speaker}]: {sentence}")
            return "\n".join(formatted_lines)
    except (json.JSONDecodeError, TypeError):
        # Not a JSON, or not the format we expect, so return as is.
        pass
    return transcription_text

def clean_llm_response(text):
    """
    Clean LLM responses by removing thinking tags and excessive whitespace.
    This handles responses from reasoning models that include <think> tags.
    """
    if not text:
        return ""
    
    # Remove thinking tags and their content
    # Handle both <think> and <thinking> tags with various closing formats
    cleaned = re.sub(r'<think(?:ing)?>.*?</think(?:ing)?>', '', text, flags=re.DOTALL | re.IGNORECASE)
    
    # Also handle unclosed thinking tags (in case the model doesn't close them)
    cleaned = re.sub(r'<think(?:ing)?>.*$', '', cleaned, flags=re.DOTALL | re.IGNORECASE)
    
    # Remove any remaining XML-like tags that might be related to thinking
    # but preserve markdown formatting
    cleaned = re.sub(r'<(?!/?(?:code|pre|blockquote|p|br|hr|ul|ol|li|h[1-6]|em|strong|b|i|a|img)(?:\s|>|/))[^>]+>', '', cleaned)
    
    # Clean up excessive whitespace while preserving intentional formatting
    # Remove leading/trailing whitespace from each line
    lines = cleaned.split('\n')
    cleaned_lines = []
    for line in lines:
        # Preserve lines that are part of code blocks or lists
        if line.strip() or (len(cleaned_lines) > 0 and cleaned_lines[-1].strip().startswith(('```', '-', '*', '1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.'))):
            cleaned_lines.append(line.rstrip())
    
    # Join lines and remove multiple consecutive blank lines
    cleaned = '\n'.join(cleaned_lines)
    cleaned = re.sub(r'\n{3,}', '\n\n', cleaned)
    
    # Final strip to remove leading/trailing whitespace
    return cleaned.strip()

def extract_thinking_content(text):
    """
    Extract thinking content from LLM responses.
    Returns a tuple of (thinking_content, main_content).
    """
    if not text:
        return ("", "")
    
    # Find all thinking tags and their content
    thinking_pattern = r'<think(?:ing)?>.*?</think(?:ing)?>'
    thinking_matches = re.findall(thinking_pattern, text, flags=re.DOTALL | re.IGNORECASE)
    
    # Extract the content from within the tags
    thinking_content = ""
    for match in thinking_matches:
        # Remove the opening and closing tags
        content = re.sub(r'^<think(?:ing)?>', '', match, flags=re.IGNORECASE)
        content = re.sub(r'</think(?:ing)?>$', '', content, flags=re.IGNORECASE)
        if thinking_content:
            thinking_content += "\n\n"
        thinking_content += content.strip()
    
    # Get the main content by removing thinking tags
    main_content = clean_llm_response(text)
    
    return (thinking_content, main_content)

def process_streaming_with_thinking(stream):
    """
    Generator that processes a streaming response and separates thinking content.
    Yields SSE-formatted data with 'delta' for regular content and 'thinking' for thinking content.
    """
    content_buffer = ""
    in_thinking = False
    thinking_buffer = ""
    
    for chunk in stream:
        content = chunk.choices[0].delta.content
        if content:
            content_buffer += content
            
            # Process the buffer to detect and handle thinking tags
            while True:
                if not in_thinking:
                    # Look for opening thinking tag
                    think_start = re.search(r'<think(?:ing)?>', content_buffer, re.IGNORECASE)
                    if think_start:
                        # Send any content before the thinking tag
                        before_thinking = content_buffer[:think_start.start()]
                        if before_thinking:
                            yield f"data: {json.dumps({'delta': before_thinking})}\n\n"
                        
                        # Start capturing thinking content
                        in_thinking = True
                        content_buffer = content_buffer[think_start.end():]
                        thinking_buffer = ""
                    else:
                        # No thinking tag found, send accumulated content
                        if content_buffer:
                            yield f"data: {json.dumps({'delta': content_buffer})}\n\n"
                        content_buffer = ""
                        break
                else:
                    # We're inside a thinking tag, look for closing tag
                    think_end = re.search(r'</think(?:ing)?>', content_buffer, re.IGNORECASE)
                    if think_end:
                        # Capture thinking content up to the closing tag
                        thinking_buffer += content_buffer[:think_end.start()]
                        
                        # Send the thinking content as a special type
                        if thinking_buffer.strip():
                            yield f"data: {json.dumps({'thinking': thinking_buffer.strip()})}\n\n"
                        
                        # Continue processing after the closing tag
                        in_thinking = False
                        content_buffer = content_buffer[think_end.end():]
                        thinking_buffer = ""
                    else:
                        # Still inside thinking tag, accumulate content
                        thinking_buffer += content_buffer
                        content_buffer = ""
                        break
    
    # Handle any remaining content
    if in_thinking and thinking_buffer:
        # Unclosed thinking tag - send as thinking content
        yield f"data: {json.dumps({'thinking': thinking_buffer.strip()})}\n\n"
    elif content_buffer:
        # Regular content
        yield f"data: {json.dumps({'delta': content_buffer})}\n\n"
    
    # Signal the end of the stream
    yield f"data: {json.dumps({'end_of_stream': True})}\n\n"

app = Flask(__name__, 
            template_folder='../templates',
            static_folder='../static')
# Use environment variables or default paths for Docker compatibility
app.config['SQLALCHEMY_DATABASE_URI'] = os.environ.get('SQLALCHEMY_DATABASE_URI', 'sqlite:////data/instance/transcriptions.db')
app.config['UPLOAD_FOLDER'] = os.environ.get('UPLOAD_FOLDER', '/data/uploads')
# MAX_CONTENT_LENGTH will be set dynamically after database initialization
# Set a secret key for session management and CSRF protection
app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', 'default-dev-key-change-in-production')

# Apply ProxyFix to handle headers from a reverse proxy (like Nginx or Caddy)
# This is crucial for request.is_secure to work correctly behind an SSL-terminating proxy.
app.wsgi_app = ProxyFix(app.wsgi_app, x_for=1, x_proto=1, x_host=1, x_prefix=1)

# --- Secure Session Cookie Configuration ---
# For local network usage, disable secure cookies to allow HTTP connections
# Only enable secure cookies in production when HTTPS is actually being used
app.config['SESSION_COOKIE_SECURE'] = False  # Allow HTTP for local network usage
app.config['SESSION_COOKIE_HTTPONLY'] = True  # Still protect against XSS
app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'  # CSRF protection

db = SQLAlchemy()
db.init_app(app)

# Initialize Flask-Login and other extensions
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'
login_manager.login_message_category = 'info'
bcrypt.init_app(app)
limiter.init_app(app)  # Initialize the limiter (uses in-memory storage by default)

csrf = CSRFProtect(app)

# Add context processor to make 'now' available to all templates
@app.context_processor
def inject_now():
    return {'now': datetime.now()}

# --- Timezone Formatting Filter ---
@app.template_filter('localdatetime')
def local_datetime_filter(dt):
    """Format a UTC datetime object to the user's local timezone."""
    if dt is None:
        return ""
    
    # Get timezone from .env, default to UTC
    user_tz_name = os.environ.get('TIMEZONE', 'UTC')
    try:
        user_tz = pytz.timezone(user_tz_name)
    except pytz.UnknownTimeZoneError:
        user_tz = pytz.utc
        app.logger.warning(f"Invalid TIMEZONE '{user_tz_name}' in .env. Defaulting to UTC.")

    # If the datetime object is naive, assume it's UTC
    if dt.tzinfo is None:
        dt = pytz.utc.localize(dt)

    # Convert to the user's timezone
    local_dt = dt.astimezone(user_tz)
    
    # Format it nicely
    return format_datetime(local_dt, format='medium', locale='en_US')

# Ensure upload and instance directories exist
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Ensure upload and instance directories exist
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
# Assuming the instance folder is handled correctly by Flask or created by setup.sh
# os.makedirs(os.path.dirname(app.config['SQLALCHEMY_DATABASE_URI'].replace('sqlite:///', '/')), exist_ok=True)


# --- User loader for Flask-Login ---
@login_manager.user_loader
def load_user(user_id):
    return db.session.get(User, int(user_id))

# --- Embedding and Chunking Utilities ---

# Initialize embedding model (lazy loading)
_embedding_model = None

def get_embedding_model():
    """Get or initialize the sentence transformer model."""
    global _embedding_model
    
    if not EMBEDDINGS_AVAILABLE:
        return None
        
    if _embedding_model is None:
        try:
            _embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
            app.logger.info("Embedding model loaded successfully")
        except Exception as e:
            app.logger.error(f"Failed to load embedding model: {e}")
            return None
    return _embedding_model

def chunk_transcription(transcription, max_chunk_length=500, overlap=50):
    """
    Split transcription into overlapping chunks for better context retrieval.
    
    Args:
        transcription (str): The full transcription text
        max_chunk_length (int): Maximum characters per chunk
        overlap (int): Character overlap between chunks
    
    Returns:
        list: List of text chunks
    """
    if not transcription or len(transcription) <= max_chunk_length:
        return [transcription] if transcription else []
    
    chunks = []
    start = 0
    
    while start < len(transcription):
        end = start + max_chunk_length
        
        # Try to break at sentence boundaries
        if end < len(transcription):
            # Look for sentence endings within the last 100 characters
            sentence_end = -1
            for i in range(max(0, end - 100), end):
                if transcription[i] in '.!?':
                    # Check if it's not an abbreviation
                    if i + 1 < len(transcription) and transcription[i + 1].isspace():
                        sentence_end = i + 1
            
            if sentence_end > start:
                end = sentence_end
        
        chunk = transcription[start:end].strip()
        if chunk:
            chunks.append(chunk)
        
        # Move start position with overlap
        start = max(start + 1, end - overlap)
        
        # Prevent infinite loop
        if start >= len(transcription):
            break
    
    return chunks

def generate_embeddings(texts):
    """
    Generate embeddings for a list of texts.
    
    Args:
        texts (list): List of text strings
    
    Returns:
        list: List of embedding vectors as numpy arrays, or empty list if embeddings unavailable
    """
    if not EMBEDDINGS_AVAILABLE:
        app.logger.warning("Embeddings not available - skipping embedding generation")
        return []
        
    model = get_embedding_model()
    if not model or not texts:
        return []
    
    try:
        embeddings = model.encode(texts)
        return [embedding.astype(np.float32) for embedding in embeddings]
    except Exception as e:
        app.logger.error(f"Error generating embeddings: {e}")
        return []

def serialize_embedding(embedding):
    """Convert numpy array to binary for database storage."""
    if embedding is None or not EMBEDDINGS_AVAILABLE:
        return None
    return embedding.tobytes()

def deserialize_embedding(binary_data):
    """Convert binary data back to numpy array."""
    if binary_data is None or not EMBEDDINGS_AVAILABLE:
        return None
    return np.frombuffer(binary_data, dtype=np.float32)

def process_recording_chunks(recording_id):
    """
    Process a recording by creating chunks and generating embeddings.
    This should be called after a recording is transcribed.
    """
    try:
        recording = db.session.get(Recording, recording_id)
        if not recording or not recording.transcription:
            return False
        
        # Delete existing chunks for this recording
        TranscriptChunk.query.filter_by(recording_id=recording_id).delete()
        
        # Create chunks
        chunks = chunk_transcription(recording.transcription)
        
        if not chunks:
            return True
        
        # Generate embeddings
        embeddings = generate_embeddings(chunks)
        
        # Store chunks in database
        for i, (chunk_text, embedding) in enumerate(zip(chunks, embeddings)):
            chunk = TranscriptChunk(
                recording_id=recording_id,
                user_id=recording.user_id,
                chunk_index=i,
                content=chunk_text,
                embedding=serialize_embedding(embedding) if embedding is not None else None
            )
            db.session.add(chunk)
        
        db.session.commit()
        app.logger.info(f"Created {len(chunks)} chunks for recording {recording_id}")
        return True
        
    except Exception as e:
        app.logger.error(f"Error processing chunks for recording {recording_id}: {e}")
        db.session.rollback()
        return False

def basic_text_search_chunks(user_id, query, filters=None, top_k=5):
    """
    Basic text search fallback when embeddings are not available.
    Uses simple text matching instead of semantic search.
    """
    try:
        # Build base query for chunks
        chunks_query = TranscriptChunk.query.filter_by(user_id=user_id)
        
        # Apply filters if provided
        if filters:
            if filters.get('tag_ids'):
                chunks_query = chunks_query.join(Recording).join(
                    RecordingTag, Recording.id == RecordingTag.recording_id
                ).filter(RecordingTag.tag_id.in_(filters['tag_ids']))
            
            if filters.get('speaker_names'):
                # Filter by participants field in recordings instead of chunk speaker_name
                if not any(hasattr(desc, 'name') and desc.name == 'recording' for desc in chunks_query.column_descriptions):
                    chunks_query = chunks_query.join(Recording)
                
                # Build OR conditions for each speaker name in participants
                speaker_conditions = []
                for speaker_name in filters['speaker_names']:
                    speaker_conditions.append(
                        Recording.participants.ilike(f'%{speaker_name}%')
                    )
                
                chunks_query = chunks_query.filter(db.or_(*speaker_conditions))
                app.logger.info(f"Applied speaker filter for: {filters['speaker_names']}")
            
            if filters.get('recording_ids'):
                chunks_query = chunks_query.filter(
                    TranscriptChunk.recording_id.in_(filters['recording_ids'])
                )
            
            if filters.get('date_from') or filters.get('date_to'):
                chunks_query = chunks_query.join(Recording)
                if filters.get('date_from'):
                    chunks_query = chunks_query.filter(Recording.meeting_date >= filters['date_from'])
                if filters.get('date_to'):
                    chunks_query = chunks_query.filter(Recording.meeting_date <= filters['date_to'])
        
        # Simple text search - split query into words and search for them
        query_words = query.lower().split()
        if query_words:
            # Create a filter that matches any of the query words in the content
            text_conditions = []
            for word in query_words:
                text_conditions.append(TranscriptChunk.content.ilike(f'%{word}%'))
            
            # Combine conditions with OR
            from sqlalchemy import or_
            chunks_query = chunks_query.filter(or_(*text_conditions))
        
        # Get chunks and return with dummy similarity scores
        chunks = chunks_query.limit(top_k).all()
        
        # Return chunks with dummy similarity scores (1.0 for found chunks)
        return [(chunk, 1.0) for chunk in chunks]
        
    except Exception as e:
        app.logger.error(f"Error in basic text search: {e}")
        return []

def semantic_search_chunks(user_id, query, filters=None, top_k=5):
    """
    Perform semantic search on transcript chunks with filtering.
    
    Args:
        user_id (int): User ID for permission filtering
        query (str): Search query
        filters (dict): Optional filters for tags, speakers, dates, recording_ids
        top_k (int): Number of top chunks to return
    
    Returns:
        list: List of relevant chunks with similarity scores
    """
    try:
        # If embeddings are not available, fall back to basic text search
        if not EMBEDDINGS_AVAILABLE:
            app.logger.info("Embeddings not available - using basic text search as fallback")
            return basic_text_search_chunks(user_id, query, filters, top_k)
        
        # Generate embedding for the query
        model = get_embedding_model()
        if not model:
            return basic_text_search_chunks(user_id, query, filters, top_k)
        
        query_embedding = model.encode([query])[0]
        
        # Build base query for chunks with eager loading of recording relationship
        chunks_query = TranscriptChunk.query.options(joinedload(TranscriptChunk.recording)).filter_by(user_id=user_id)
        
        # Apply filters if provided
        if filters:
            if filters.get('tag_ids'):
                # Join with recordings that have specified tags
                chunks_query = chunks_query.join(Recording).join(
                    RecordingTag, Recording.id == RecordingTag.recording_id
                ).filter(RecordingTag.tag_id.in_(filters['tag_ids']))
            
            if filters.get('speaker_names'):
                # Filter by participants field in recordings instead of chunk speaker_name
                if not any(hasattr(desc, 'name') and desc.name == 'recording' for desc in chunks_query.column_descriptions):
                    chunks_query = chunks_query.join(Recording)
                
                # Build OR conditions for each speaker name in participants
                speaker_conditions = []
                for speaker_name in filters['speaker_names']:
                    speaker_conditions.append(
                        Recording.participants.ilike(f'%{speaker_name}%')
                    )
                
                chunks_query = chunks_query.filter(db.or_(*speaker_conditions))
                app.logger.info(f"Applied speaker filter for: {filters['speaker_names']}")
            
            if filters.get('recording_ids'):
                chunks_query = chunks_query.filter(
                    TranscriptChunk.recording_id.in_(filters['recording_ids'])
                )
            
            if filters.get('date_from') or filters.get('date_to'):
                chunks_query = chunks_query.join(Recording)
                if filters.get('date_from'):
                    chunks_query = chunks_query.filter(Recording.meeting_date >= filters['date_from'])
                if filters.get('date_to'):
                    chunks_query = chunks_query.filter(Recording.meeting_date <= filters['date_to'])
        
        # Get chunks that have embeddings
        chunks = chunks_query.filter(TranscriptChunk.embedding.isnot(None)).all()
        
        if not chunks:
            return []
        
        # Calculate similarities
        chunk_similarities = []
        for chunk in chunks:
            try:
                chunk_embedding = deserialize_embedding(chunk.embedding)
                if chunk_embedding is not None:
                    similarity = cosine_similarity(
                        query_embedding.reshape(1, -1),
                        chunk_embedding.reshape(1, -1)
                    )[0][0]
                    chunk_similarities.append((chunk, float(similarity)))
            except Exception as e:
                app.logger.warning(f"Error calculating similarity for chunk {chunk.id}: {e}")
                continue
        
        # Sort by similarity and return top k
        chunk_similarities.sort(key=lambda x: x[1], reverse=True)
        return chunk_similarities[:top_k]
        
    except Exception as e:
        app.logger.error(f"Error in semantic search: {e}")
        return []

# --- Database Models ---
class User(db.Model, UserMixin):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(20), unique=True, nullable=False)
    email = db.Column(db.String(120), unique=True, nullable=False)
    password = db.Column(db.String(60), nullable=False)
    is_admin = db.Column(db.Boolean, default=False)
    recordings = db.relationship('Recording', backref='owner', lazy=True)
    transcription_language = db.Column(db.String(10), nullable=True) # For ISO 639-1 codes
    output_language = db.Column(db.String(50), nullable=True) # For full language names like "Spanish"
    ui_language = db.Column(db.String(10), nullable=True, default='en') # For UI language preference (en, es, fr, zh)
    summary_prompt = db.Column(db.Text, nullable=True)
    extract_events = db.Column(db.Boolean, default=False)  # Enable event extraction from transcripts
    name = db.Column(db.String(100), nullable=True)
    job_title = db.Column(db.String(100), nullable=True)
    company = db.Column(db.String(100), nullable=True)
    diarize = db.Column(db.Boolean, default=False)
    
    def __repr__(self):
        return f"User('{self.username}', '{self.email}')"
class Speaker(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100), nullable=False)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    last_used = db.Column(db.DateTime, default=datetime.utcnow)
    use_count = db.Column(db.Integer, default=1)
    
    # Relationship to user
    user = db.relationship('User', backref=db.backref('speakers', lazy=True, cascade='all, delete-orphan'))
    
    def to_dict(self):
        return {
            'id': self.id,
            'name': self.name,
            'created_at': self.created_at,
            'last_used': self.last_used,
            'use_count': self.use_count
        }

class SystemSetting(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    key = db.Column(db.String(100), unique=True, nullable=False)
    value = db.Column(db.Text, nullable=True)
    description = db.Column(db.Text, nullable=True)
    setting_type = db.Column(db.String(50), nullable=False, default='string')  # string, integer, boolean, float
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)
    
    def to_dict(self):
        return {
            'id': self.id,
            'key': self.key,
            'value': self.value,
            'description': self.description,
            'setting_type': self.setting_type,
            'created_at': self.created_at,
            'updated_at': self.updated_at
        }
    
    @staticmethod
    def get_setting(key, default_value=None):
        """Get a system setting value by key, with optional default."""
        setting = SystemSetting.query.filter_by(key=key).first()
        if setting:
            # Convert value based on type
            if setting.setting_type == 'integer':
                try:
                    return int(setting.value) if setting.value is not None else default_value
                except (ValueError, TypeError):
                    return default_value
            elif setting.setting_type == 'boolean':
                return setting.value.lower() in ('true', '1', 'yes') if setting.value else default_value
            elif setting.setting_type == 'float':
                try:
                    return float(setting.value) if setting.value is not None else default_value
                except (ValueError, TypeError):
                    return default_value
            else:  # string
                return setting.value if setting.value is not None else default_value
        return default_value
    
    @staticmethod
    def set_setting(key, value, description=None, setting_type='string'):
        """Set a system setting value."""
        setting = SystemSetting.query.filter_by(key=key).first()
        if setting:
            setting.value = str(value) if value is not None else None
            setting.updated_at = datetime.utcnow()
            if description:
                setting.description = description
            if setting_type:
                setting.setting_type = setting_type
        else:
            setting = SystemSetting(
                key=key,
                value=str(value) if value is not None else None,
                description=description,
                setting_type=setting_type
            )
            db.session.add(setting)
        db.session.commit()
        return setting

# Many-to-many relationship table for recordings and tags
class RecordingTag(db.Model):
    __tablename__ = 'recording_tags'
    recording_id = db.Column(db.Integer, db.ForeignKey('recording.id'), primary_key=True)
    tag_id = db.Column(db.Integer, db.ForeignKey('tag.id'), primary_key=True)
    added_at = db.Column(db.DateTime, default=datetime.utcnow, nullable=True)
    order = db.Column(db.Integer, nullable=False, default=0)
    
    # Relationships
    recording = db.relationship('Recording', back_populates='tag_associations')
    tag = db.relationship('Tag', back_populates='recording_associations')

class Tag(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(50), nullable=False)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    color = db.Column(db.String(7), default='#3B82F6')  # Hex color for UI
    
    # Custom settings for this tag
    custom_prompt = db.Column(db.Text, nullable=True)  # Custom summarization prompt
    default_language = db.Column(db.String(10), nullable=True)  # Default transcription language
    default_min_speakers = db.Column(db.Integer, nullable=True)  # Default min speakers for ASR
    default_max_speakers = db.Column(db.Integer, nullable=True)  # Default max speakers for ASR
    
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)
    
    # Relationships
    user = db.relationship('User', backref=db.backref('tags', lazy=True, cascade='all, delete-orphan'))
    # Use association object for many-to-many with order tracking
    recording_associations = db.relationship('RecordingTag', back_populates='tag', cascade='all, delete-orphan')
    
    # Unique constraint: tag name must be unique per user
    __table_args__ = (db.UniqueConstraint('name', 'user_id', name='_user_tag_uc'),)
    
    def to_dict(self):
        return {
            'id': self.id,
            'name': self.name,
            'color': self.color,
            'custom_prompt': self.custom_prompt,
            'default_language': self.default_language,
            'default_min_speakers': self.default_min_speakers,
            'default_max_speakers': self.default_max_speakers,
            'created_at': self.created_at.isoformat() if self.created_at else None,
            'recording_count': len(self.recording_associations)
        }

class Event(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    recording_id = db.Column(db.Integer, db.ForeignKey('recording.id'), nullable=False)
    title = db.Column(db.String(200), nullable=False)
    description = db.Column(db.Text, nullable=True)
    start_datetime = db.Column(db.DateTime, nullable=False)
    end_datetime = db.Column(db.DateTime, nullable=True)
    location = db.Column(db.String(500), nullable=True)
    attendees = db.Column(db.Text, nullable=True)  # JSON list of attendees
    reminder_minutes = db.Column(db.Integer, nullable=True, default=15)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

    # Relationship
    recording = db.relationship('Recording', backref=db.backref('events', lazy=True, cascade='all, delete-orphan'))

    def to_dict(self):
        return {
            'id': self.id,
            'recording_id': self.recording_id,
            'title': self.title,
            'description': self.description,
            'start_datetime': self.start_datetime.isoformat() if self.start_datetime else None,
            'end_datetime': self.end_datetime.isoformat() if self.end_datetime else None,
            'location': self.location,
            'attendees': json.loads(self.attendees) if self.attendees else [],
            'reminder_minutes': self.reminder_minutes,
            'created_at': self.created_at.isoformat() if self.created_at else None
        }

class Share(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    public_id = db.Column(db.String(32), unique=True, nullable=False, default=lambda: secrets.token_urlsafe(16))
    recording_id = db.Column(db.Integer, db.ForeignKey('recording.id'), nullable=False)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    share_summary = db.Column(db.Boolean, default=True)
    share_notes = db.Column(db.Boolean, default=True)

    user = db.relationship('User', backref=db.backref('shares', lazy=True, cascade='all, delete-orphan'))
    recording = db.relationship('Recording', backref=db.backref('shares', lazy=True, cascade='all, delete-orphan'))

    def to_dict(self):
        return {
            'id': self.id,
            'public_id': self.public_id,
            'recording_id': self.recording_id,
            'created_at': local_datetime_filter(self.created_at),
            'share_summary': self.share_summary,
            'share_notes': self.share_notes,
            'recording_title': self.recording.title if self.recording else "N/A"
        }

class Recording(db.Model):
    # Add user_id foreign key to associate recordings with users
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=True)
    id = db.Column(db.Integer, primary_key=True)
    # Title will now often be AI-generated, maybe start with filename?
    title = db.Column(db.String(200), nullable=True) # Allow Null initially
    participants = db.Column(db.String(500))
    notes = db.Column(db.Text)
    transcription = db.Column(db.Text, nullable=True)
    summary = db.Column(db.Text, nullable=True) # <-- ADDED: Summary field
    status = db.Column(db.String(50), default='PENDING') # PENDING, PROCESSING, SUMMARIZING, COMPLETED, FAILED
    audio_path = db.Column(db.String(500))
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    meeting_date = db.Column(db.Date, nullable=True) # <-- ADDED: Meeting Date field
    file_size = db.Column(db.Integer)  # Store file size in bytes
    original_filename = db.Column(db.String(500), nullable=True) # Store the original uploaded filename
    is_inbox = db.Column(db.Boolean, default=True)  # New recordings are marked as inbox by default
    is_highlighted = db.Column(db.Boolean, default=False)  # Recordings can be highlighted by the user
    mime_type = db.Column(db.String(100), nullable=True)
    completed_at = db.Column(db.DateTime, nullable=True)
    processing_time_seconds = db.Column(db.Integer, nullable=True)
    processing_source = db.Column(db.String(50), default='upload')  # upload, auto_process, recording
    error_message = db.Column(db.Text, nullable=True)  # Store detailed error messages
    
    # Relationships
    tag_associations = db.relationship('RecordingTag', back_populates='recording', cascade='all, delete-orphan', order_by='RecordingTag.order')
    
    @property
    def tags(self):
        """Get tags ordered by the order they were added to this recording."""
        return [assoc.tag for assoc in sorted(self.tag_associations, key=lambda x: x.order)]

    def to_dict(self):
        return {
            'id': self.id,
            'title': self.title,
            'participants': self.participants,
            'notes': self.notes,
            'notes_html': md_to_html(self.notes) if self.notes else "",
            'transcription': self.transcription,
            'summary': self.summary,
            'summary_html': md_to_html(self.summary) if self.summary else "",
            'status': self.status,
            'created_at': local_datetime_filter(self.created_at),
            'completed_at': local_datetime_filter(self.completed_at),
            'processing_time_seconds': self.processing_time_seconds,
            'meeting_date': f"{self.meeting_date.isoformat()}T00:00:00" if self.meeting_date else None, # <-- ADDED: Include meeting_date with time component
            'file_size': self.file_size,
            'original_filename': self.original_filename, # <-- ADDED: Include original filename
            'user_id': self.user_id,
            'is_inbox': self.is_inbox,
            'is_highlighted': self.is_highlighted,
            'mime_type': self.mime_type,
            'tags': [tag.to_dict() for tag in self.tags] if self.tags else [],
            'events': [event.to_dict() for event in self.events] if self.events else []
        }

class TranscriptChunk(db.Model):
    """Stores chunked transcription segments for efficient retrieval and embedding."""
    id = db.Column(db.Integer, primary_key=True)
    recording_id = db.Column(db.Integer, db.ForeignKey('recording.id'), nullable=False)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    chunk_index = db.Column(db.Integer, nullable=False)  # Order within the recording
    content = db.Column(db.Text, nullable=False)  # The actual text chunk
    start_time = db.Column(db.Float, nullable=True)  # Start time in seconds (if available)
    end_time = db.Column(db.Float, nullable=True)  # End time in seconds (if available)
    speaker_name = db.Column(db.String(100), nullable=True)  # Speaker for this chunk
    embedding = db.Column(db.LargeBinary, nullable=True)  # Stored as binary vector
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    
    # Relationships
    recording = db.relationship('Recording', backref=db.backref('chunks', lazy=True, cascade='all, delete-orphan'))
    user = db.relationship('User', backref=db.backref('transcript_chunks', lazy=True, cascade='all, delete-orphan'))
    
    def to_dict(self):
        return {
            'id': self.id,
            'recording_id': self.recording_id,
            'chunk_index': self.chunk_index,
            'content': self.content,
            'start_time': self.start_time,
            'end_time': self.end_time,
            'speaker_name': self.speaker_name,
            'created_at': self.created_at.isoformat() if self.created_at else None
        }

class TranscriptTemplate(db.Model):
    """Stores user-defined templates for transcript formatting."""
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    name = db.Column(db.String(100), nullable=False)
    template = db.Column(db.Text, nullable=False)
    description = db.Column(db.String(500), nullable=True)
    is_default = db.Column(db.Boolean, default=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)

    # Relationships
    user = db.relationship('User', backref=db.backref('transcript_templates', lazy=True, cascade='all, delete-orphan'))

    def to_dict(self):
        return {
            'id': self.id,
            'name': self.name,
            'template': self.template,
            'description': self.description,
            'is_default': self.is_default,
            'created_at': self.created_at.isoformat() if self.created_at else None,
            'updated_at': self.updated_at.isoformat() if self.updated_at else None
        }

class InquireSession(db.Model):
    """Tracks inquire mode sessions and their filtering criteria."""
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    session_name = db.Column(db.String(200), nullable=True)  # Optional user-defined name
    
    # Filter criteria (JSON stored as text)
    filter_tags = db.Column(db.Text, nullable=True)  # JSON array of tag IDs
    filter_speakers = db.Column(db.Text, nullable=True)  # JSON array of speaker names
    filter_date_from = db.Column(db.Date, nullable=True)
    filter_date_to = db.Column(db.Date, nullable=True)
    filter_recording_ids = db.Column(db.Text, nullable=True)  # JSON array of specific recording IDs
    
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    last_used = db.Column(db.DateTime, default=datetime.utcnow)
    
    # Relationships
    user = db.relationship('User', backref=db.backref('inquire_sessions', lazy=True, cascade='all, delete-orphan'))
    
    def to_dict(self):
        return {
            'id': self.id,
            'session_name': self.session_name,
            'filter_tags': json.loads(self.filter_tags) if self.filter_tags else [],
            'filter_speakers': json.loads(self.filter_speakers) if self.filter_speakers else [],
            'filter_date_from': self.filter_date_from.isoformat() if self.filter_date_from else None,
            'filter_date_to': self.filter_date_to.isoformat() if self.filter_date_to else None,
            'filter_recording_ids': json.loads(self.filter_recording_ids) if self.filter_recording_ids else [],
            'created_at': self.created_at.isoformat() if self.created_at else None,
            'last_used': self.last_used.isoformat() if self.last_used else None
        }

# --- Forms for Authentication ---
# --- Custom Password Validator ---
def password_check(form, field):
    password = field.data
    if len(password) < 8:
        raise ValidationError('Password must be at least 8 characters long.')
    if not re.search(r'[A-Z]', password):
        raise ValidationError('Password must contain at least one uppercase letter.')
    if not re.search(r'[a-z]', password):
        raise ValidationError('Password must contain at least one lowercase letter.')
    if not re.search(r'[0-9]', password):
        raise ValidationError('Password must contain at least one number.')
    if not re.search(r'[!@#$%^&*(),.?":{}|<>]', password):
        raise ValidationError('Password must contain at least one special character.')

# --- Share Routes ---
@app.route('/share/<string:public_id>', methods=['GET'])
def view_shared_recording(public_id):
    share = Share.query.filter_by(public_id=public_id).first_or_404()
    recording = share.recording
    
    # Create a limited dictionary for the public view
    recording_data = {
        'id': recording.id,
        'public_id': share.public_id,
        'title': recording.title,
        'participants': recording.participants,
        'transcription': recording.transcription,
        'summary': md_to_html(recording.summary) if share.share_summary else None,
        'notes': md_to_html(recording.notes) if share.share_notes else None,
        'meeting_date': f"{recording.meeting_date.isoformat()}T00:00:00" if recording.meeting_date else None,
        'mime_type': recording.mime_type
    }
    
    return render_template('share.html', recording=recording_data)

@app.route('/api/recording/<int:recording_id>/share', methods=['GET'])
@login_required
def get_existing_share(recording_id):
    """Check if a share already exists for this recording."""
    recording = db.session.get(Recording, recording_id)
    if not recording or recording.user_id != current_user.id:
        return jsonify({'error': 'Recording not found or you do not have permission to view it.'}), 404
    
    existing_share = Share.query.filter_by(
        recording_id=recording.id,
        user_id=current_user.id
    ).order_by(Share.created_at.desc()).first()
    
    if existing_share:
        share_url = url_for('view_shared_recording', public_id=existing_share.public_id, _external=True)
        return jsonify({
            'success': True,
            'exists': True,
            'share_url': share_url,
            'share': existing_share.to_dict()
        }), 200
    else:
        return jsonify({
            'success': True,
            'exists': False
        }), 200

@app.route('/api/recording/<int:recording_id>/share', methods=['POST'])
@login_required
def create_share(recording_id):
    if not request.is_secure:
        return jsonify({'error': 'Sharing is only available over a secure (HTTPS) connection.'}), 403
        
    recording = db.session.get(Recording, recording_id)
    if not recording or recording.user_id != current_user.id:
        return jsonify({'error': 'Recording not found or you do not have permission to share it.'}), 404
        
    data = request.json
    share_summary = data.get('share_summary', True)
    share_notes = data.get('share_notes', True)
    force_new = data.get('force_new', False)
    
    # Check if ANY share already exists for this recording by this user
    # Get the most recent one if there are multiple
    existing_share = Share.query.filter_by(
        recording_id=recording.id,
        user_id=current_user.id
    ).order_by(Share.created_at.desc()).first()
    
    if existing_share and not force_new:
        # Update the share permissions if they've changed
        if existing_share.share_summary != share_summary or existing_share.share_notes != share_notes:
            existing_share.share_summary = share_summary
            existing_share.share_notes = share_notes
            db.session.commit()
        
        # Return existing share info
        share_url = url_for('view_shared_recording', public_id=existing_share.public_id, _external=True)
        return jsonify({
            'success': True, 
            'share_url': share_url, 
            'share': existing_share.to_dict(),
            'existing': True,
            'message': 'Using existing share link for this recording'
        }), 200
    
    # Create new share (either no existing share or force_new is True)
    share = Share(
        recording_id=recording.id,
        user_id=current_user.id,
        share_summary=share_summary,
        share_notes=share_notes
    )
    db.session.add(share)
    db.session.commit()
    
    share_url = url_for('view_shared_recording', public_id=share.public_id, _external=True)
    
    return jsonify({
        'success': True, 
        'share_url': share_url, 
        'share': share.to_dict(),
        'existing': False
    }), 201

@app.route('/api/shares', methods=['GET'])
@login_required
def get_shares():
    shares = Share.query.filter_by(user_id=current_user.id).order_by(Share.created_at.desc()).all()
    return jsonify([share.to_dict() for share in shares])

@app.route('/api/share/<int:share_id>', methods=['PUT'])
@login_required
def update_share(share_id):
    share = Share.query.filter_by(id=share_id, user_id=current_user.id).first_or_404()
    data = request.json
    
    if 'share_summary' in data:
        share.share_summary = data['share_summary']
    if 'share_notes' in data:
        share.share_notes = data['share_notes']
        
    db.session.commit()
    return jsonify({'success': True, 'share': share.to_dict()})

@app.route('/api/share/<int:share_id>', methods=['DELETE'])
@login_required
def delete_share(share_id):
    share = Share.query.filter_by(id=share_id, user_id=current_user.id).first_or_404()
    db.session.delete(share)
    db.session.commit()
    return jsonify({'success': True})

# --- User Preferences API Endpoint ---
@app.route('/api/user/preferences', methods=['POST'])
@login_required
def save_user_preferences():
    """Save user preferences including UI language"""
    data = request.json
    
    if 'language' in data:
        current_user.ui_language = data['language']
    
    db.session.commit()
    
    return jsonify({
        'success': True,
        'message': 'Preferences saved successfully',
        'ui_language': current_user.ui_language
    })

# --- System Info API Endpoint ---
@app.route('/api/system/info', methods=['GET'])
def get_system_info():
    """Get system information including version and model details."""
    try:
        # Use the same version detection logic as startup
        version = get_version()
        
        return jsonify({
            'version': version,
            'llm_endpoint': TEXT_MODEL_BASE_URL,
            'llm_model': TEXT_MODEL_NAME,
            'whisper_endpoint': os.environ.get('TRANSCRIPTION_BASE_URL', 'https://api.openai.com/v1'),
            'asr_enabled': USE_ASR_ENDPOINT,
            'asr_endpoint': ASR_BASE_URL if USE_ASR_ENDPOINT else None
        })
    except Exception as e:
        app.logger.error(f"Error getting system info: {e}")
        return jsonify({'error': 'Unable to retrieve system information'}), 500

# --- Tag API Endpoints ---
@app.route('/api/tags', methods=['GET'])
@login_required
def get_tags():
    """Get all tags for the current user."""
    tags = Tag.query.filter_by(user_id=current_user.id).order_by(Tag.name).all()
    return jsonify([tag.to_dict() for tag in tags])

@app.route('/api/tags', methods=['POST'])
@login_required
def create_tag():
    """Create a new tag."""
    data = request.get_json()
    
    if not data or not data.get('name'):
        return jsonify({'error': 'Tag name is required'}), 400
    
    # Check if tag with same name already exists for this user
    existing_tag = Tag.query.filter_by(name=data['name'], user_id=current_user.id).first()
    if existing_tag:
        return jsonify({'error': 'Tag with this name already exists'}), 400
    
    tag = Tag(
        name=data['name'],
        user_id=current_user.id,
        color=data.get('color', '#3B82F6'),
        custom_prompt=data.get('custom_prompt'),
        default_language=data.get('default_language'),
        default_min_speakers=data.get('default_min_speakers'),
        default_max_speakers=data.get('default_max_speakers')
    )
    
    db.session.add(tag)
    db.session.commit()
    
    return jsonify(tag.to_dict()), 201

@app.route('/api/tags/<int:tag_id>', methods=['PUT'])
@login_required
def update_tag(tag_id):
    """Update a tag."""
    tag = Tag.query.filter_by(id=tag_id, user_id=current_user.id).first_or_404()
    data = request.get_json()
    
    if 'name' in data:
        # Check if new name conflicts with another tag
        existing_tag = Tag.query.filter_by(name=data['name'], user_id=current_user.id).filter(Tag.id != tag_id).first()
        if existing_tag:
            return jsonify({'error': 'Another tag with this name already exists'}), 400
        tag.name = data['name']
    
    if 'color' in data:
        tag.color = data['color']
    if 'custom_prompt' in data:
        tag.custom_prompt = data['custom_prompt']
    if 'default_language' in data:
        tag.default_language = data['default_language']
    if 'default_min_speakers' in data:
        tag.default_min_speakers = data['default_min_speakers']
    if 'default_max_speakers' in data:
        tag.default_max_speakers = data['default_max_speakers']
    
    tag.updated_at = datetime.utcnow()
    db.session.commit()
    
    return jsonify(tag.to_dict())

@app.route('/api/tags/<int:tag_id>', methods=['DELETE'])
@login_required
def delete_tag(tag_id):
    """Delete a tag."""
    tag = Tag.query.filter_by(id=tag_id, user_id=current_user.id).first_or_404()
    db.session.delete(tag)
    db.session.commit()
    return jsonify({'success': True})

@app.route('/api/recordings/<int:recording_id>/tags', methods=['POST'])
@login_required
def add_tag_to_recording(recording_id):
    """Add a tag to a recording."""
    recording = Recording.query.filter_by(id=recording_id, user_id=current_user.id).first_or_404()
    data = request.get_json()
    
    tag_id = data.get('tag_id')
    if not tag_id:
        return jsonify({'error': 'tag_id is required'}), 400
    
    tag = Tag.query.filter_by(id=tag_id, user_id=current_user.id).first_or_404()
    
    # Check if tag is already associated with recording
    existing_association = RecordingTag.query.filter_by(recording_id=recording_id, tag_id=tag_id).first()
    if not existing_association:
        # Get the next order number for this recording
        max_order = db.session.query(db.func.max(RecordingTag.order)).filter_by(recording_id=recording_id).scalar() or 0
        
        # Create new association with proper order
        new_association = RecordingTag(
            recording_id=recording_id,
            tag_id=tag_id,
            order=max_order + 1,
            added_at=datetime.utcnow()
        )
        db.session.add(new_association)
        db.session.commit()
    
    return jsonify({'success': True, 'tags': [t.to_dict() for t in recording.tags]})

@app.route('/api/recordings/<int:recording_id>/tags/<int:tag_id>', methods=['DELETE'])
@login_required
def remove_tag_from_recording(recording_id, tag_id):
    """Remove a tag from a recording."""
    recording = Recording.query.filter_by(id=recording_id, user_id=current_user.id).first_or_404()
    tag = Tag.query.filter_by(id=tag_id, user_id=current_user.id).first_or_404()
    
    # Find and remove the association
    association = RecordingTag.query.filter_by(recording_id=recording_id, tag_id=tag_id).first()
    if association:
        db.session.delete(association)
        db.session.commit()
    
    return jsonify({'success': True, 'tags': [t.to_dict() for t in recording.tags]})

class RegistrationForm(FlaskForm):
    username = StringField('Username', validators=[DataRequired(), Length(min=2, max=20)])
    email = StringField('Email', validators=[DataRequired(), Email()])
    password = PasswordField('Password', validators=[DataRequired(), password_check])
    confirm_password = PasswordField('Confirm Password', validators=[DataRequired(), EqualTo('password')])
    submit = SubmitField('Sign Up')
    
    def validate_username(self, username):
        user = User.query.filter_by(username=username.data).first()
        if user:
            raise ValidationError('That username is already taken. Please choose a different one.')
    
    def validate_email(self, email):
        user = User.query.filter_by(email=email.data).first()
        if user:
            raise ValidationError('That email is already registered. Please use a different one.')

class LoginForm(FlaskForm):
    email = StringField('Email', validators=[DataRequired(), Email()])
    password = PasswordField('Password', validators=[DataRequired()])
    remember = BooleanField('Remember Me')
    submit = SubmitField('Login')

# Function to check and add columns if they don't exist
def add_column_if_not_exists(engine, table_name, column_name, column_type):
    """Add a column to a table if it doesn't exist."""
    from sqlalchemy import text
    
    try:
        # Check if column exists
        with engine.connect() as conn:
            # For SQLite, we can query the pragma table_info using text()
            result = conn.execute(text(f"PRAGMA table_info({table_name})"))
            columns = [row[1] for row in result]
            
            if column_name not in columns:
                app.logger.info(f"Adding column {column_name} to {table_name}")
                conn.execute(text(f"ALTER TABLE {table_name} ADD COLUMN {column_name} {column_type}"))
                return True
            return False
    except Exception as e:
        app.logger.error(f"Error checking/adding column {column_name} to {table_name}: {e}")
        return False

with app.app_context():
    db.create_all()
    
    # Check and add new columns if they don't exist
    engine = db.engine
    try:
        # Add is_inbox column with default value of 1 (True)
        if add_column_if_not_exists(engine, 'recording', 'is_inbox', 'BOOLEAN DEFAULT 1'):
            app.logger.info("Added is_inbox column to recording table")
        
        # Add is_highlighted column with default value of 0 (False)
        if add_column_if_not_exists(engine, 'recording', 'is_highlighted', 'BOOLEAN DEFAULT 0'):
            app.logger.info("Added is_highlighted column to recording table")

        # Add language preference columns to User table
        if add_column_if_not_exists(engine, 'user', 'transcription_language', 'VARCHAR(10)'):
            app.logger.info("Added transcription_language column to user table")

        # Add extract_events column to User table
        if add_column_if_not_exists(engine, 'user', 'extract_events', 'BOOLEAN DEFAULT 0'):
            app.logger.info("Added extract_events column to user table")
        if add_column_if_not_exists(engine, 'user', 'output_language', 'VARCHAR(50)'):
            app.logger.info("Added output_language column to user table")
        if add_column_if_not_exists(engine, 'user', 'summary_prompt', 'TEXT'):
            app.logger.info("Added summary_prompt column to user table")
        if add_column_if_not_exists(engine, 'user', 'name', 'VARCHAR(100)'):
            app.logger.info("Added name column to user table")
        if add_column_if_not_exists(engine, 'user', 'job_title', 'VARCHAR(100)'):
            app.logger.info("Added job_title column to user table")
        if add_column_if_not_exists(engine, 'user', 'company', 'VARCHAR(100)'):
            app.logger.info("Added company column to user table")
        if add_column_if_not_exists(engine, 'user', 'diarize', 'BOOLEAN'):
            app.logger.info("Added diarize column to user table")
        if add_column_if_not_exists(engine, 'user', 'ui_language', 'VARCHAR(10) DEFAULT "en"'):
            app.logger.info("Added ui_language column to user table")
        if add_column_if_not_exists(engine, 'recording', 'mime_type', 'VARCHAR(100)'):
            app.logger.info("Added mime_type column to recording table")
        if add_column_if_not_exists(engine, 'recording', 'completed_at', 'DATETIME'):
            app.logger.info("Added completed_at column to recording table")
        if add_column_if_not_exists(engine, 'recording', 'processing_time_seconds', 'INTEGER'):
            app.logger.info("Added processing_time_seconds column to recording table")
        if add_column_if_not_exists(engine, 'recording', 'processing_source', 'VARCHAR(50) DEFAULT "upload"'):
            app.logger.info("Added processing_source column to recording table")
        if add_column_if_not_exists(engine, 'recording', 'error_message', 'TEXT'):
            app.logger.info("Added error_message column to recording table")
            
        # Add columns to recording_tags for order tracking
        if add_column_if_not_exists(engine, 'recording_tags', 'added_at', 'DATETIME'):
            app.logger.info("Added added_at column to recording_tags table")
        if add_column_if_not_exists(engine, 'recording_tags', 'order', '"order" INTEGER DEFAULT 0'):
            app.logger.info("Added order column to recording_tags table")
            
            # Update existing records to have proper order values (approximate by tag_id)
            try:
                from sqlalchemy import text
                with engine.connect() as conn:
                    # Get existing associations without order values and assign them
                    existing_associations = conn.execute(text('''
                        SELECT recording_id, tag_id, 
                               ROW_NUMBER() OVER (PARTITION BY recording_id ORDER BY tag_id) as row_num
                        FROM recording_tags 
                        WHERE "order" = 0
                    ''')).fetchall()
                    
                    for assoc in existing_associations:
                        conn.execute(text('''
                            UPDATE recording_tags 
                            SET "order" = :order_num 
                            WHERE recording_id = :rec_id AND tag_id = :tag_id
                        '''), {"order_num": assoc.row_num, "rec_id": assoc.recording_id, "tag_id": assoc.tag_id})
                    
                    conn.commit()
                    app.logger.info(f"Updated order values for {len(existing_associations)} existing tag associations")
            except Exception as e:
                app.logger.warning(f"Could not update existing tag order values: {e}")
        
        # Initialize default system settings
        if not SystemSetting.query.filter_by(key='transcript_length_limit').first():
            SystemSetting.set_setting(
                key='transcript_length_limit',
                value='30000',
                description='Maximum number of characters to send from transcript to LLM for summarization and chat. Use -1 for no limit.',
                setting_type='integer'
            )
            app.logger.info("Initialized default transcript_length_limit setting")
            
        if not SystemSetting.query.filter_by(key='max_file_size_mb').first():
            SystemSetting.set_setting(
                key='max_file_size_mb',
                value='250',
                description='Maximum file size allowed for audio uploads in megabytes (MB).',
                setting_type='integer'
            )
            app.logger.info("Initialized default max_file_size_mb setting")
        
        if not SystemSetting.query.filter_by(key='asr_timeout_seconds').first():
            SystemSetting.set_setting(
                key='asr_timeout_seconds',
                value='1800',
                description='Maximum time in seconds to wait for ASR transcription to complete. Default is 1800 seconds (30 minutes).',
                setting_type='integer'
            )
            app.logger.info("Initialized default asr_timeout_seconds setting")
        
        if not SystemSetting.query.filter_by(key='admin_default_summary_prompt').first():
            default_prompt = """Generate a comprehensive summary that includes the following sections:
- **Key Issues Discussed**: A bulleted list of the main topics
- **Key Decisions Made**: A bulleted list of any decisions reached
- **Action Items**: A bulleted list of tasks assigned, including who is responsible if mentioned"""
            SystemSetting.set_setting(
                key='admin_default_summary_prompt',
                value=default_prompt,
                description='Default summarization prompt used when users have not set their own prompt. This serves as the base prompt for all users.',
                setting_type='string'
            )
            app.logger.info("Initialized admin_default_summary_prompt setting")
        
        if not SystemSetting.query.filter_by(key='recording_disclaimer').first():
            SystemSetting.set_setting(
                key='recording_disclaimer',
                value='',
                description='Legal disclaimer shown to users before recording starts. Supports Markdown formatting. Leave empty to disable.',
                setting_type='string'
            )
            app.logger.info("Initialized recording_disclaimer setting")
        
        # Process existing recordings for inquire mode (chunk and embed them)
        # Only run if inquire mode is enabled
        if ENABLE_INQUIRE_MODE:
            # Use a file lock to prevent multiple workers from running this simultaneously
            import fcntl
            import tempfile
            lock_file_path = os.path.join(tempfile.gettempdir(), 'inquire_migration.lock')
            
            try:
                with open(lock_file_path, 'w') as lock_file:
                    # Try to acquire exclusive lock (non-blocking)
                    try:
                        fcntl.flock(lock_file.fileno(), fcntl.LOCK_EX | fcntl.LOCK_NB)
                        app.logger.info("Acquired migration lock, checking for existing recordings that need chunking for inquire mode...")
                        
                        completed_recordings = Recording.query.filter_by(status='COMPLETED').all()
                        recordings_needing_processing = []
                        
                        for recording in completed_recordings:
                            if recording.transcription:  # Has transcription
                                chunk_count = TranscriptChunk.query.filter_by(recording_id=recording.id).count()
                                if chunk_count == 0:  # No chunks yet
                                    recordings_needing_processing.append(recording)
                        
                        if recordings_needing_processing:
                            app.logger.info(f"Found {len(recordings_needing_processing)} recordings that need chunking for inquire mode")
                            app.logger.info("Processing first 10 recordings automatically. Use admin API or migration script for remaining recordings.")
                            
                            # Process first 10 recordings automatically to avoid long startup times
                            batch_size = min(10, len(recordings_needing_processing))
                            processed = 0
                            
                            for i in range(batch_size):
                                recording = recordings_needing_processing[i]
                                try:
                                    success = process_recording_chunks(recording.id)
                                    if success:
                                        processed += 1
                                        app.logger.info(f"Processed chunks for recording: {recording.title} ({recording.id})")
                                except Exception as e:
                                    app.logger.warning(f"Failed to process chunks for recording {recording.id}: {e}")
                            
                            remaining = len(recordings_needing_processing) - processed
                            if remaining > 0:
                                app.logger.info(f"Successfully processed {processed} recordings. {remaining} recordings remaining.")
                                app.logger.info("Use the admin migration API or run 'python migrate_existing_recordings.py' to process remaining recordings.")
                            else:
                                app.logger.info(f"Successfully processed all {processed} recordings for inquire mode.")
                        else:
                            app.logger.info("All existing recordings are already processed for inquire mode.")
                        
                    except BlockingIOError:
                        app.logger.info("Migration already running in another worker, skipping...")
                    
            except Exception as e:
                app.logger.warning(f"Error during existing recordings migration: {e}")
                app.logger.info("Existing recordings can be migrated later using the admin API or migration script.")
            
    except Exception as e:
        app.logger.error(f"Error during database migration: {e}")

# --- API client setup for OpenRouter ---
# Use environment variables from .env
TEXT_MODEL_API_KEY = os.environ.get("TEXT_MODEL_API_KEY")
# Strip any inline comments from URLs (users might add "# comment" in .env files)
TEXT_MODEL_BASE_URL = os.environ.get("TEXT_MODEL_BASE_URL", "https://openrouter.ai/api/v1")
if TEXT_MODEL_BASE_URL:
    TEXT_MODEL_BASE_URL = TEXT_MODEL_BASE_URL.split('#')[0].strip()
TEXT_MODEL_NAME = os.environ.get("TEXT_MODEL_NAME", "openai/gpt-3.5-turbo") # Default if not set

# Set up HTTP client with custom headers for OpenRouter app identification
app_headers = {
    "HTTP-Referer": "https://github.com/murtaza-nasir/speakr",  # Your app's repo URL for OpenRouter visibility
    "X-Title": "Speakr - AI Audio Transcription",  # Your app name for OpenRouter visibility
    "User-Agent": "Speakr/1.0 (https://github.com/murtaza-nasir/speakr)"  # Custom user agent for better tracking
}

http_client_no_proxy = httpx.Client(
    verify=True,
    headers=app_headers
)

try:
    # Always attempt to create client - use API key if provided, otherwise use placeholder
    api_key = TEXT_MODEL_API_KEY or "not-needed"
    client = OpenAI(
        api_key=api_key,
        base_url=TEXT_MODEL_BASE_URL,
        http_client=http_client_no_proxy
    )
    app.logger.info(f"LLM client initialized for endpoint: {TEXT_MODEL_BASE_URL}. Using model: {TEXT_MODEL_NAME}")
    if "openrouter" in TEXT_MODEL_BASE_URL.lower():
        app.logger.info("OpenRouter integration: App identification headers added for visibility in logs")
except Exception as client_init_e:
    app.logger.error(f"Failed to initialize LLM client: {client_init_e}", exc_info=True)
    client = None

def call_llm_completion(messages, temperature=0.7, response_format=None, stream=False, max_tokens=None):
    """
    Centralized function for LLM API calls with proper error handling and logging.
    
    Args:
        messages: List of message dicts with 'role' and 'content'
        temperature: Sampling temperature (0-1)
        response_format: Optional response format dict (e.g., {"type": "json_object"})
        stream: Whether to stream the response
        max_tokens: Optional maximum tokens to generate
        
    Returns:
        OpenAI completion object or generator (if streaming)
    """
    if not client:
        raise ValueError("LLM client not initialized")
    
    if not TEXT_MODEL_API_KEY:
        raise ValueError("TEXT_MODEL_API_KEY not configured")
    
    try:
        completion_args = {
            "model": TEXT_MODEL_NAME,
            "messages": messages,
            "temperature": temperature,
            "stream": stream
        }
        
        if response_format:
            completion_args["response_format"] = response_format
            
        if max_tokens:
            completion_args["max_tokens"] = max_tokens
            
        return client.chat.completions.create(**completion_args)
        
    except Exception as e:
        app.logger.error(f"LLM API call failed: {e}")
        raise

# Store details for the transcription client (potentially different)
transcription_api_key = os.environ.get("TRANSCRIPTION_API_KEY", "")
# Strip any inline comments from URLs (users might add "# comment" in .env files)
transcription_base_url = os.environ.get("TRANSCRIPTION_BASE_URL", "")
if transcription_base_url:
    transcription_base_url = transcription_base_url.split('#')[0].strip()


# ASR endpoint configuration
USE_ASR_ENDPOINT = os.environ.get('USE_ASR_ENDPOINT', 'false').lower() == 'true'
# Strip any inline comments from ASR_BASE_URL (users might add "# comment" in .env files)
ASR_BASE_URL = os.environ.get('ASR_BASE_URL')
if ASR_BASE_URL:
    # Remove everything after '#' if it exists (inline comment)
    ASR_BASE_URL = ASR_BASE_URL.split('#')[0].strip()

# When using ASR endpoint, automatically enable diarization and set sensible defaults
# Users can still override these if needed, but they default to the expected ASR behavior
if USE_ASR_ENDPOINT:
    # Default to diarization enabled for ASR (can be overridden by setting ASR_DIARIZE=false)
    ASR_DIARIZE = os.environ.get('ASR_DIARIZE', 'true').lower() == 'true'
    # Default speaker range for most conversations (None means auto-detect)
    ASR_MIN_SPEAKERS = os.environ.get('ASR_MIN_SPEAKERS')
    ASR_MAX_SPEAKERS = os.environ.get('ASR_MAX_SPEAKERS')
else:
    # When not using ASR, these settings are irrelevant
    ASR_DIARIZE = False
    ASR_MIN_SPEAKERS = None
    ASR_MAX_SPEAKERS = None

# Audio chunking configuration for large files with OpenAI Whisper API
ENABLE_CHUNKING = os.environ.get('ENABLE_CHUNKING', 'true').lower() == 'true'
CHUNK_SIZE_MB = int(os.environ.get('CHUNK_SIZE_MB', '20'))  # 20MB default for safety margin
CHUNK_OVERLAP_SECONDS = int(os.environ.get('CHUNK_OVERLAP_SECONDS', '3'))  # 3 seconds overlap

# Initialize chunking service
chunking_service = AudioChunkingService(
    max_chunk_size_mb=CHUNK_SIZE_MB,
    overlap_seconds=CHUNK_OVERLAP_SECONDS
) if ENABLE_CHUNKING else None

# Get and log version information at startup
def get_version():
    # Try reading VERSION file first (works in Docker)
    try:
        with open('VERSION', 'r') as f:
            return f.read().strip()
    except FileNotFoundError:
        pass
    
    # Fall back to git tags (works in development)
    try:
        import subprocess
        return subprocess.check_output(['git', 'describe', '--tags', '--abbrev=0'], 
                                     stderr=subprocess.DEVNULL).decode().strip()
    except:
        pass
    
    # Final fallback
    return "unknown"

version = get_version()

app.logger.info(f"=== Speakr {version} Starting Up ===")
app.logger.info(f"Using LLM endpoint: {TEXT_MODEL_BASE_URL} with model: {TEXT_MODEL_NAME}")

# Validate transcription service configuration
if USE_ASR_ENDPOINT:
    if not ASR_BASE_URL:
        app.logger.error("ERROR: ASR endpoint is enabled but ASR_BASE_URL is not configured!")
        app.logger.error("Please set ASR_BASE_URL in your .env file or disable USE_ASR_ENDPOINT")
        sys.exit(1)
    app.logger.info(f"Using ASR endpoint for transcription at: {ASR_BASE_URL}")
else:
    # Check if Whisper API is configured
    if not transcription_base_url or not transcription_api_key:
        app.logger.error("ERROR: No transcription service configured!")
        app.logger.error("You must configure either:")
        app.logger.error("  1. ASR endpoint: Set USE_ASR_ENDPOINT=true and ASR_BASE_URL=<your-asr-url>")
        app.logger.error("  2. Whisper API: Set TRANSCRIPTION_BASE_URL and TRANSCRIPTION_API_KEY")
        app.logger.error("Please check your .env file configuration.")
        sys.exit(1)
    app.logger.info(f"Using Whisper API at: {transcription_base_url}")

# --- Background Transcription & Summarization Task ---
def format_api_error_message(error_str):
    """
    Formats API error messages to be more user-friendly.
    Specifically handles token limit errors with helpful suggestions.
    """
    error_lower = error_str.lower()
    
    # Check for token limit errors
    if 'maximum context length' in error_lower and 'tokens' in error_lower:
        return "[Summary generation failed: The transcription is too long for AI processing. Request your admin to try using a different LLM with a larger context size, or set a limit for the transcript_length_limit in the system settings.]"
    
    # Check for other common API errors
    if 'rate limit' in error_lower:
        return "[Summary generation failed: API rate limit exceeded. Please try again in a few minutes.]"
    
    if 'insufficient funds' in error_lower or 'quota exceeded' in error_lower:
        return "[Summary generation failed: API quota exceeded. Please contact support.]"
    
    if 'timeout' in error_lower:
        return "[Summary generation failed: Request timed out. Please try again.]"
    
    # For other errors, show a generic message
    return f"[Summary generation failed: {error_str}]"

def generate_title_task(app_context, recording_id):
    """Generates only a title for a recording based on transcription.
    
    Args:
        app_context: Flask app context
        recording_id: ID of the recording
    """
    with app_context:
        recording = db.session.get(Recording, recording_id)
        if not recording:
            app.logger.error(f"Error: Recording {recording_id} not found for title generation.")
            return
            
        if client is None:
            app.logger.warning(f"Skipping title generation for {recording_id}: OpenRouter client not configured.")
            # Still mark as completed even if we can't generate a title
            recording.status = 'COMPLETED'
            recording.completed_at = datetime.utcnow()
            db.session.commit()
            return
            
        if not recording.transcription or len(recording.transcription.strip()) < 10:
            app.logger.warning(f"Transcription for recording {recording_id} is too short or empty. Skipping title generation.")
            # Still mark as completed even if we can't generate a title
            recording.status = 'COMPLETED'
            recording.completed_at = datetime.utcnow()
            db.session.commit()
            return
        
        # Get configurable transcript length limit and format transcription for LLM
        transcript_limit = SystemSetting.get_setting('transcript_length_limit', 30000)
        if transcript_limit == -1:
            raw_transcription = recording.transcription
        else:
            raw_transcription = recording.transcription[:transcript_limit]
            
        # Convert ASR JSON to clean text format
        transcript_text = format_transcription_for_llm(raw_transcription)
        
        
        # Get user language preference
        user_output_language = None
        if recording.owner:
            user_output_language = recording.owner.output_language
            
        language_directive = f"Please provide the title in {user_output_language}." if user_output_language else ""
        
        prompt_text = f"""Create a short title for this conversation:

{transcript_text}

Requirements:
- Maximum 8 words
- No phrases like "Discussion about" or "Meeting on"  
- Just the main topic

{language_directive}

Title:"""

        system_message_content = "You are an AI assistant that generates concise titles for audio transcriptions. Respond only with the title."
        if user_output_language:
            system_message_content += f" Ensure your response is in {user_output_language}."
        
            
        try:
            completion = call_llm_completion(
                messages=[
                    {"role": "system", "content": system_message_content},
                    {"role": "user", "content": prompt_text}
                ],
                temperature=0.7,
                max_tokens=5000
            )
            
            
            raw_response = completion.choices[0].message.content
            reasoning = getattr(completion.choices[0].message, 'reasoning', None)
            
            # Use reasoning content if main content is empty (fallback for reasoning models)
            if not raw_response and reasoning:
                app.logger.info(f"Title generation for recording {recording_id}: Using reasoning field as fallback")
                # Try to extract a title from the reasoning field
                lines = reasoning.strip().split('\n')
                # Look for the last line that might be the title
                for line in reversed(lines):
                    line = line.strip()
                    if line and not line.startswith('I') and len(line.split()) <= 8:
                        raw_response = line
                        break
            
            title = clean_llm_response(raw_response) if raw_response else ""
            
            if title:
                recording.title = title
                app.logger.info(f"Title generated for recording {recording_id}: {title}")
            else:
                app.logger.warning(f"Empty title generated for recording {recording_id}")
                
        except Exception as e:
            app.logger.error(f"Error generating title for recording {recording_id}: {str(e)}")
            app.logger.error(f"Exception details:", exc_info=True)
        
        # Always set status to COMPLETED after title generation (successful or not)
        # This ensures transcription processing is marked as complete
        recording.status = 'COMPLETED'
        recording.completed_at = datetime.utcnow()
        db.session.commit()
        
        # Process chunks for semantic search after completion (if inquire mode is enabled)
        if ENABLE_INQUIRE_MODE:
            try:
                process_recording_chunks(recording_id)
            except Exception as e:
                app.logger.error(f"Error processing chunks for completed recording {recording_id}: {e}")

def generate_summary_only_task(app_context, recording_id):
    """Generates only a summary for a recording (no title, no JSON response).
    
    Args:
        app_context: Flask app context
        recording_id: ID of the recording
    """
    with app_context:
        recording = db.session.get(Recording, recording_id)
        if not recording:
            app.logger.error(f"Error: Recording {recording_id} not found for summary generation.")
            return
            
        if client is None:
            app.logger.warning(f"Skipping summary generation for {recording_id}: OpenRouter client not configured.")
            recording.summary = "[Summary skipped: OpenRouter client not configured]"
            db.session.commit()
            return
            
        recording.status = 'SUMMARIZING'
        db.session.commit()
        
        app.logger.info(f"Requesting summary from OpenRouter for recording {recording_id} using model {TEXT_MODEL_NAME}...")
        
        if not recording.transcription or len(recording.transcription.strip()) < 10:
            app.logger.warning(f"Transcription for recording {recording_id} is too short or empty. Skipping summarization.")
            recording.summary = "[Summary skipped due to short transcription]"
            recording.status = 'COMPLETED'
            db.session.commit()
            return
        
        # Get user preferences and tag custom prompts
        user_summary_prompt = None
        user_output_language = None
        tag_custom_prompt = None
        
        # Collect all custom prompts from tags in the order they were added to this recording
        tag_custom_prompts = []
        if recording.tags:
            # Tags are now automatically ordered by the order they were added to this recording
            for tag in recording.tags:
                if tag.custom_prompt and tag.custom_prompt.strip():
                    tag_custom_prompts.append({
                        'name': tag.name,
                        'prompt': tag.custom_prompt.strip()
                    })
                    app.logger.info(f"Found custom prompt from tag '{tag.name}' for recording {recording_id}")
        
        # Create merged prompt if we have multiple tag prompts
        if tag_custom_prompts:
            if len(tag_custom_prompts) == 1:
                tag_custom_prompt = tag_custom_prompts[0]['prompt']
                app.logger.info(f"Using single custom prompt from tag '{tag_custom_prompts[0]['name']}' for recording {recording_id}")
            else:
                # Merge multiple prompts seamlessly as unified instructions
                merged_parts = []
                for tag_prompt in tag_custom_prompts:
                    merged_parts.append(tag_prompt['prompt'])
                tag_custom_prompt = "\n\n".join(merged_parts)
                tag_names = [tp['name'] for tp in tag_custom_prompts]
                app.logger.info(f"Combined custom prompts from {len(tag_custom_prompts)} tags in order added ({', '.join(tag_names)}) for recording {recording_id}")
        else:
            tag_custom_prompt = None
        
        if recording.owner:
            user_summary_prompt = recording.owner.summary_prompt
            user_output_language = recording.owner.output_language
        
        # Format transcription for LLM (convert JSON to clean text format like clipboard copy)
        formatted_transcription = format_transcription_for_llm(recording.transcription)
        
        # Get configurable transcript length limit
        transcript_limit = SystemSetting.get_setting('transcript_length_limit', 30000)
        if transcript_limit == -1:
            transcript_text = formatted_transcription
        else:
            transcript_text = formatted_transcription[:transcript_limit]
        
        language_directive = f"IMPORTANT: You MUST provide the summary in {user_output_language}. The entire response must be in {user_output_language}." if user_output_language else ""
        
        # Determine which summarization instructions to use
        # Priority order: tag custom prompt > user summary prompt > admin default prompt > hardcoded fallback
        summarization_instructions = ""
        if tag_custom_prompt:
            app.logger.info(f"Using tag custom prompt for recording {recording_id}")
            summarization_instructions = tag_custom_prompt
        elif user_summary_prompt:
            app.logger.info(f"Using user custom prompt for recording {recording_id}")
            summarization_instructions = user_summary_prompt
        else:
            # Get admin default prompt from system settings
            admin_default_prompt = SystemSetting.get_setting('admin_default_summary_prompt', None)
            if admin_default_prompt:
                app.logger.info(f"Using admin default prompt for recording {recording_id}")
                summarization_instructions = admin_default_prompt
            else:
                # Fallback to hardcoded default if admin hasn't set one
                summarization_instructions = """Generate a comprehensive summary that includes the following sections:
- **Key Issues Discussed**: A bulleted list of the main topics
- **Key Decisions Made**: A bulleted list of any decisions reached
- **Action Items**: A bulleted list of tasks assigned, including who is responsible if mentioned"""
                app.logger.info(f"Using hardcoded default prompt for recording {recording_id}")
        
        # Build context information
        current_date = datetime.now().strftime("%B %d, %Y")
        context_parts = []
        context_parts.append(f"Current date: {current_date}")
        
        # Add selected tags information
        if recording.tags:
            tag_names = [tag.name for tag in recording.tags]
            context_parts.append(f"Tags applied to this transcript by the user: {', '.join(tag_names)}")
        
        # Add user profile information if available
        if recording.owner:
            user_context_parts = []
            if recording.owner.name:
                user_context_parts.append(f"Name: {recording.owner.name}")
            if recording.owner.job_title:
                user_context_parts.append(f"Job title: {recording.owner.job_title}")
            if recording.owner.company:
                user_context_parts.append(f"Company: {recording.owner.company}")
            
            if user_context_parts:
                context_parts.append(f"Information about the user: {', '.join(user_context_parts)}")
        
        context_section = "Context:\n" + "\n".join(f"- {part}" for part in context_parts)
        
        # Build SYSTEM message: Initial instructions + Context + Language
        system_message_content = "You are an AI assistant that generates comprehensive summaries for meeting transcripts. Respond only with the summary in Markdown format. Do NOT use markdown code blocks (```markdown). Provide raw markdown content directly."
        system_message_content += f"\n\n{context_section}"
        if user_output_language:
            system_message_content += f"\n\nLanguage Requirement: You MUST generate the entire summary in {user_output_language}. This is mandatory."
        
        # Build USER message: Transcription + Summarization Instructions + Language Directive
        prompt_text = f"""Transcription:
\"\"\"
{transcript_text}
\"\"\"

Summarization Instructions:
{summarization_instructions}

{language_directive}"""
            
        # Debug logging: Log the complete prompt being sent to the LLM
        app.logger.info(f"Sending summarization prompt to LLM (length: {len(prompt_text)} chars). Set LOG_LEVEL=DEBUG to see full prompt details.")
        app.logger.debug(f"=== SUMMARIZATION DEBUG for recording {recording_id} ===")
        app.logger.debug(f"System message: {system_message_content}")
        app.logger.debug(f"User prompt (length: {len(prompt_text)} chars):\n{prompt_text}")
        app.logger.debug(f"=== END SUMMARIZATION DEBUG for recording {recording_id} ===")
            
        try:
            completion = call_llm_completion(
                messages=[
                    {"role": "system", "content": system_message_content},
                    {"role": "user", "content": prompt_text}
                ],
                temperature=0.5,
                max_tokens=int(os.environ.get("SUMMARY_MAX_TOKENS", "3000"))
            )
            
            raw_response = completion.choices[0].message.content
            app.logger.info(f"Raw LLM response for recording {recording_id}: '{raw_response}'")
            
            summary = clean_llm_response(raw_response) if raw_response else ""
            app.logger.info(f"Processed summary length for recording {recording_id}: {len(summary)} characters")
            
            if summary:
                recording.summary = summary
                recording.status = 'COMPLETED'
                recording.completed_at = datetime.utcnow()
                db.session.commit()
                app.logger.info(f"Summary generated successfully for recording {recording_id}")

                # Extract events if enabled for this user
                if recording.owner and recording.owner.extract_events:
                    extract_events_from_transcript(recording_id, formatted_transcription, summary)
            else:
                app.logger.warning(f"Empty summary generated for recording {recording_id}")
                recording.summary = "[Summary not generated]"
                recording.status = 'COMPLETED'
                db.session.commit()
                
        except Exception as e:
            error_msg = handle_openai_api_error(e, "summary")
            app.logger.error(f"Error generating summary for recording {recording_id}: {str(e)}")
            recording.summary = error_msg
            recording.status = 'FAILED'
            db.session.commit()

def extract_events_from_transcript(recording_id, transcript_text, summary_text):
    """Extract calendar events from transcript using LLM.

    Args:
        recording_id: ID of the recording
        transcript_text: The formatted transcript text
        summary_text: The generated summary text
    """
    try:
        recording = db.session.get(Recording, recording_id)
        if not recording or not recording.owner or not recording.owner.extract_events:
            return  # Event extraction not enabled for this user

        app.logger.info(f"Extracting events for recording {recording_id}")

        # Build comprehensive context information
        current_date = datetime.now()
        context_parts = []

        # CRITICAL: Determine the reference date for relative date calculations
        reference_date = None
        reference_date_source = ""

        if recording.meeting_date:
            # Prefer meeting date if available
            reference_date = recording.meeting_date
            reference_date_source = "Meeting Date"
            context_parts.append(f"**MEETING DATE (use this for relative date calculations): {recording.meeting_date.strftime('%A, %B %d, %Y')}**")
        elif recording.created_at:
            # Fall back to upload date
            reference_date = recording.created_at.date()
            reference_date_source = "Upload Date (no meeting date available)"
            context_parts.append(f"**REFERENCE DATE (use this for relative date calculations): {recording.created_at.strftime('%A, %B %d, %Y')}**")

        context_parts.append(f"Today's actual date: {current_date.strftime('%A, %B %d, %Y')}")
        context_parts.append(f"Current time: {current_date.strftime('%I:%M %p')}")

        # Add additional recording context
        if recording.created_at:
            context_parts.append(f"Recording uploaded on: {recording.created_at.strftime('%B %d, %Y at %I:%M %p')}")
        if recording.meeting_date and reference_date_source == "Meeting Date":
            # Calculate days between meeting and today for context
            days_since = (current_date.date() - recording.meeting_date).days
            if days_since == 0:
                context_parts.append("This meeting happened today")
            elif days_since == 1:
                context_parts.append("This meeting happened yesterday")
            else:
                context_parts.append(f"This meeting happened {days_since} days ago")

        # Add user context for better understanding
        if recording.owner:
            user_context = []
            if recording.owner.name:
                user_context.append(f"User's name: {recording.owner.name}")
            if recording.owner.job_title:
                user_context.append(f"Job title: {recording.owner.job_title}")
            if recording.owner.company:
                user_context.append(f"Company: {recording.owner.company}")
            if user_context:
                context_parts.append("User information: " + ", ".join(user_context))

        # Add participants if available
        if recording.participants:
            context_parts.append(f"Participants in the meeting: {recording.participants}")

        context_section = "\n".join(context_parts)

        # Prepare the prompt for event extraction
        event_prompt = f"""You are analyzing a meeting transcript to extract calendar events. Use the context below to correctly interpret relative dates and times.

IMPORTANT CONTEXT:
{context_section}

INSTRUCTIONS:
1. **CRITICAL**: Use the MEETING DATE shown above as your reference point for ALL relative date calculations
2. When people say "next Wednesday" or "tomorrow" or "next week", calculate from the MEETING DATE, not today's date
3. Example: If the meeting date is September 13, 2025 and someone says "next Wednesday", that means September 17, 2025
4. If no specific time is mentioned for an event, use 09:00:00 (9 AM) as the default start time
5. Pay attention to time zones if mentioned
6. Extract ONLY events that are explicitly discussed as future appointments, meetings, or deadlines
7. Do NOT create events for past occurrences or general discussions

For each event found, extract:
- Title: A clear, concise title for the event
- Description: Brief description including context from the meeting
- Start date/time: The calculated actual date/time (in ISO format YYYY-MM-DDTHH:MM:SS, use 09:00:00 if no time specified)
- End date/time: When the event ends (if mentioned, in ISO format, default to 1 hour after start if not specified)
- Location: Where the event will take place (if mentioned)
- Attendees: List of people who should attend (if mentioned)
- Reminder minutes: How many minutes before to remind (default 15)

Transcript Summary:
{summary_text}

Transcript excerpt (for additional context):
{transcript_text[:8000]}

RESPONSE FORMAT:
Respond with a JSON object containing an "events" array. If no events are found, return a JSON object with an empty events array.

Example response:
{{
  "events": [
    {{
      "title": "Project Review Meeting",
      "description": "Quarterly review to discuss project progress and next steps as discussed in the meeting",
      "start_datetime": "2025-07-22T14:00:00",
      "end_datetime": "2025-07-22T15:30:00",
      "location": "Conference Room A",
      "attendees": ["John Smith", "Jane Doe", "Bob Johnson"],
      "reminder_minutes": 15
    }}
  ]
}}

CRITICAL RULES:
1. **BASE ALL DATE CALCULATIONS ON THE MEETING DATE PROVIDED IN THE CONTEXT ABOVE**
2. Only extract events that are FUTURE relative to the MEETING DATE (not today's date)
3. Convert all relative dates using the MEETING DATE as the reference point
4. Example: If the meeting date is September 13, 2025 (Friday) and someone says:
   - "next Wednesday" = September 17, 2025
   - "tomorrow" = September 14, 2025
   - "next week" = week of September 15-19, 2025
5. IMPORTANT: If no time is mentioned, always use 09:00:00 (9 AM) as the start time, NOT midnight
6. Include context from the discussion in the description
7. Do NOT invent or assume events not explicitly discussed
8. If unsure about a date/time, do not include that event"""

        completion = call_llm_completion(
            messages=[
                {"role": "system", "content": """You are an expert at extracting calendar events from meeting transcripts. You excel at:
1. Understanding relative date references ("next Tuesday", "tomorrow", "in two weeks") and converting them to absolute dates
2. Identifying genuine future appointments, meetings, and deadlines from conversations
3. Distinguishing between actual planned events vs. general discussions
4. Extracting participant names and meeting details accurately

You must respond with valid JSON format only."""},
                {"role": "user", "content": event_prompt}
            ],
            temperature=0.2,
            response_format={"type": "json_object"},
            max_tokens=3000
        )

        response_content = completion.choices[0].message.content
        events_data = safe_json_loads(response_content, {})

        # Handle both {"events": [...]} and direct array format
        if isinstance(events_data, dict) and 'events' in events_data:
            events_list = events_data['events']
        elif isinstance(events_data, list):
            events_list = events_data
        else:
            events_list = []

        app.logger.info(f"Found {len(events_list)} events for recording {recording_id}")

        # Save events to database
        for event_data in events_list:
            try:
                # Parse dates
                start_dt = None
                end_dt = None

                if 'start_datetime' in event_data:
                    try:
                        # Try ISO format first
                        start_dt = datetime.fromisoformat(event_data['start_datetime'].replace('Z', '+00:00'))
                    except:
                        # Try other common formats
                        from dateutil import parser
                        try:
                            start_dt = parser.parse(event_data['start_datetime'])
                        except:
                            app.logger.warning(f"Could not parse start_datetime: {event_data['start_datetime']}")
                            continue  # Skip this event if we can't parse the date

                if 'end_datetime' in event_data and event_data['end_datetime']:
                    try:
                        end_dt = datetime.fromisoformat(event_data['end_datetime'].replace('Z', '+00:00'))
                    except:
                        from dateutil import parser
                        try:
                            end_dt = parser.parse(event_data['end_datetime'])
                        except:
                            pass  # End time is optional

                # Create event record
                event = Event(
                    recording_id=recording_id,
                    title=event_data.get('title', 'Untitled Event')[:200],
                    description=event_data.get('description', ''),
                    start_datetime=start_dt,
                    end_datetime=end_dt,
                    location=event_data.get('location', '')[:500] if event_data.get('location') else None,
                    attendees=json.dumps(event_data.get('attendees', [])) if event_data.get('attendees') else None,
                    reminder_minutes=event_data.get('reminder_minutes', 15)
                )

                db.session.add(event)
                app.logger.info(f"Added event '{event.title}' for recording {recording_id}")

            except Exception as e:
                app.logger.error(f"Error saving event for recording {recording_id}: {str(e)}")
                continue

        db.session.commit()

    except Exception as e:
        app.logger.error(f"Error extracting events for recording {recording_id}: {str(e)}")
        db.session.rollback()

def extract_audio_from_video(video_filepath, output_format='mp3', cleanup_original=True):
    """Extract audio from video containers using FFmpeg.
    
    Uses MP3 codec for optimal compatibility and predictable file sizes.
    64kbps MP3 provides good speech quality at ~480KB per minute.
    """
    try:
        # Generate output filename with audio extension
        base_filepath, file_ext = os.path.splitext(video_filepath)
        temp_audio_filepath = f"{base_filepath}_audio_temp.{output_format}"
        final_audio_filepath = f"{base_filepath}_audio.{output_format}"
        
        app.logger.info(f"Extracting audio from video: {video_filepath} -> {temp_audio_filepath}")
        
        # Extract audio using FFmpeg - using high-quality MP3 for better transcription
        subprocess.run([
            'ffmpeg', '-i', video_filepath, '-y',
            '-vn',  # No video
            '-codec:a', 'libmp3lame',  # Use LAME MP3 encoder explicitly
            '-b:a', '128k',  # 128kbps bitrate for high quality
            '-ar', '44100',  # 44.1kHz sample rate for better quality
            '-ac', '1',  # Mono (sufficient for speech, reduces file size)
            '-compression_level', '2',  # Better compression
            temp_audio_filepath
        ], check=True, capture_output=True, text=True)
        
        app.logger.info(f"Successfully extracted audio to {temp_audio_filepath}")
        
        # Optionally preserve temp file for debugging (set PRESERVE_TEMP_AUDIO=true in env)
        if os.getenv('PRESERVE_TEMP_AUDIO', 'false').lower() == 'true':
            import shutil
            shutil.copy2(temp_audio_filepath, temp_audio_filepath.replace('_temp', '_debug'))
            app.logger.info(f"Debug: Preserved temp audio file as {temp_audio_filepath.replace('_temp', '_debug')}")
        
        # Rename temp file to final filename
        os.rename(temp_audio_filepath, final_audio_filepath)
        
        # Clean up original video file if requested
        if cleanup_original:
            try:
                os.remove(video_filepath)
                app.logger.info(f"Cleaned up original video file: {video_filepath}")
            except Exception as e:
                app.logger.warning(f"Failed to clean up original video file {video_filepath}: {str(e)}")
        
        return final_audio_filepath, f'audio/{output_format}'
        
    except subprocess.CalledProcessError as e:
        app.logger.error(f"FFmpeg audio extraction failed for {video_filepath}: {e.stderr}")
        raise Exception(f"Audio extraction failed: {e.stderr}")
    except FileNotFoundError:
        app.logger.error("FFmpeg command not found. Please ensure FFmpeg is installed and in the system's PATH.")
        raise Exception("Audio conversion tool (FFmpeg) not found on server.")
    except Exception as e:
        app.logger.error(f"Error extracting audio from {video_filepath}: {str(e)}")
        raise

def transcribe_audio_asr(app_context, recording_id, filepath, original_filename, start_time, mime_type=None, language=None, diarize=False, min_speakers=None, max_speakers=None, tag_id=None):
    """Transcribes audio using the ASR webservice."""
    with app_context:
        recording = db.session.get(Recording, recording_id)
        if not recording:
            app.logger.error(f"Error: Recording {recording_id} not found for ASR transcription.")
            return

        try:
            app.logger.info(f"Starting ASR transcription for recording {recording_id}...")
            recording.status = 'PROCESSING'
            db.session.commit()

            # Check if we need to extract audio from video container
            actual_filepath = filepath
            actual_content_type = mime_type or mimetypes.guess_type(original_filename)[0] or 'application/octet-stream'
            actual_filename = original_filename

            # List of video MIME types that need audio extraction
            video_mime_types = [
                'video/mp4', 'video/quicktime', 'video/x-msvideo', 'video/webm',
                'video/avi', 'video/x-ms-wmv', 'video/3gpp'
            ]
            
            # Check if file is a video container by MIME type or extension
            is_video = (
                actual_content_type.startswith('video/') or 
                actual_content_type in video_mime_types or
                original_filename.lower().endswith(('.mp4', '.mov', '.avi', '.mkv', '.webm', '.wmv', '.3gp'))
            )
            
            if is_video:
                app.logger.info(f"Video container detected ({actual_content_type}), extracting audio...")
                try:
                    # Extract audio from video
                    audio_filepath, audio_mime_type = extract_audio_from_video(filepath, 'mp3')
                    
                    # Update paths and MIME type for ASR processing
                    actual_filepath = audio_filepath
                    actual_content_type = audio_mime_type
                    actual_filename = os.path.basename(audio_filepath)
                    
                    # Update recording with extracted audio path and new MIME type
                    recording.audio_path = audio_filepath
                    recording.mime_type = audio_mime_type
                    db.session.commit()
                    
                    app.logger.info(f"Audio extracted successfully: {audio_filepath}")
                except Exception as e:
                    app.logger.error(f"Failed to extract audio from video: {str(e)}")
                    recording.status = 'FAILED'
                    recording.error_msg = f"Audio extraction failed: {str(e)}"
                    db.session.commit()
                    return

            # Keep track of whether we've already tried WAV conversion
            wav_conversion_attempted = False
            wav_converted_filepath = None
            
            # Retry loop for handling 500 errors with WAV conversion
            max_attempts = 2
            for attempt in range(max_attempts):
                try:
                    # Use converted MP3 if available from previous attempt
                    current_filepath = wav_converted_filepath if wav_converted_filepath else actual_filepath
                    current_content_type = 'audio/mpeg' if wav_converted_filepath else actual_content_type
                    current_filename = os.path.basename(current_filepath)
                    
                    with open(current_filepath, 'rb') as audio_file:
                        url = f"{ASR_BASE_URL}/asr"
                        params = {
                            'encode': True,
                            'task': 'transcribe',
                            'output': 'json'
                        }
                        if language:
                            params['language'] = language
                        if diarize:
                            params['diarize'] = diarize
                        if min_speakers:
                            params['min_speakers'] = min_speakers
                        if max_speakers:
                            params['max_speakers'] = max_speakers

                        content_type = current_content_type
                        app.logger.info(f"Using MIME type {content_type} for ASR upload.")
                        files = {'audio_file': (current_filename, audio_file, content_type)}
                        
                        with httpx.Client() as client:
                            # Get configurable ASR timeout from database (default 30 minutes)
                            asr_timeout_seconds = SystemSetting.get_setting('asr_timeout_seconds', 1800)
                            timeout = httpx.Timeout(None, connect=30.0, read=float(asr_timeout_seconds), write=30.0, pool=30.0)
                            app.logger.info(f"Sending ASR request to {url} with params: {params} (timeout: {asr_timeout_seconds}s)")
                            response = client.post(url, params=params, files=files, timeout=timeout)
                            app.logger.info(f"ASR request completed with status: {response.status_code}")
                            response.raise_for_status()
                            
                            # Parse the JSON response from ASR (moved here so it's accessible)
                            asr_response_data = response.json()
                    
                    # If we reach here, the request was successful
                    break
                    
                except httpx.HTTPStatusError as e:
                    # Check if it's a 500 error and we haven't tried WAV conversion yet
                    if e.response.status_code == 500 and attempt == 0 and not wav_conversion_attempted:
                        app.logger.warning(f"ASR returned 500 error for recording {recording_id}, attempting high-quality MP3 conversion and retry...")
                        
                        # Convert to high-quality MP3 for better compatibility
                        filename_lower = actual_filename.lower()
                        if not filename_lower.endswith('.mp3'):
                            try:
                                base_filepath, file_ext = os.path.splitext(actual_filepath)
                                temp_mp3_filepath = f"{base_filepath}_temp.mp3"
                                
                                app.logger.info(f"Converting {actual_filename} to high-quality MP3 format for retry...")
                                subprocess.run(
                                    ['ffmpeg', '-i', actual_filepath, '-y', '-acodec', 'libmp3lame', '-b:a', '128k', '-ar', '44100', temp_mp3_filepath],
                                    check=True, capture_output=True, text=True
                                )
                                app.logger.info(f"Successfully converted {actual_filepath} to {temp_mp3_filepath}")
                                
                                wav_converted_filepath = temp_mp3_filepath  # Keep variable name for compatibility
                                wav_conversion_attempted = True
                                # Continue to next iteration to retry with WAV
                                continue
                            except subprocess.CalledProcessError as conv_error:
                                app.logger.error(f"Failed to convert to WAV: {conv_error}")
                                # Re-raise the original HTTP error if conversion fails
                                raise e
                        else:
                            # Already a WAV file, can't convert further
                            app.logger.error(f"File is already WAV but still getting 500 error")
                            raise e
                    else:
                        # Not a 500 error or already tried conversion, propagate the error
                        raise e
            
            # DEBUG: Preserve converted file for quality checking
            if wav_converted_filepath and os.path.exists(wav_converted_filepath):
                try:
                    # Get file size and basic info for debugging
                    converted_size = os.path.getsize(wav_converted_filepath)
                    converted_size_mb = converted_size / (1024 * 1024)
                    
                    # Create a debug copy in a known location
                    debug_dir = os.path.join(app.config['UPLOAD_FOLDER'], 'debug_converted')
                    os.makedirs(debug_dir, exist_ok=True)
                    
                    # Copy the converted file with a timestamp (MP3 now)
                    from shutil import copy2
                    file_ext = os.path.splitext(wav_converted_filepath)[1] or '.mp3'
                    debug_filename = f"debug_{recording_id}_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}{file_ext}"
                    debug_filepath = os.path.join(debug_dir, debug_filename)
                    copy2(wav_converted_filepath, debug_filepath)
                    
                    app.logger.info(f"DEBUG: Converted file preserved at: {debug_filepath}")
                    app.logger.info(f"DEBUG: Converted file size: {converted_size_mb:.2f} MB ({converted_size} bytes)")
                    app.logger.info(f"DEBUG: Original file: {actual_filename}")
                    app.logger.info(f"DEBUG: Recording ID: {recording_id}")
                    app.logger.info(f"DEBUG: You can download this file from the container at: {debug_filepath}")
                    
                except Exception as debug_error:
                    app.logger.warning(f"DEBUG: Failed to preserve converted file: {debug_error}")
            
            # Clean up the original temporary converted file (but keep debug copy)
            try:
                if wav_converted_filepath and os.path.exists(wav_converted_filepath):
                    os.remove(wav_converted_filepath)
                    app.logger.info(f"Cleaned up original temporary converted file: {wav_converted_filepath}")
            except Exception as cleanup_error:
                app.logger.warning(f"Failed to clean up temporary converted file: {cleanup_error}")
            
            # Debug logging for ASR response
            app.logger.info(f"ASR response keys: {list(asr_response_data.keys())}")
            
            # Log the complete raw JSON response (truncated for readability)
            import json as json_module
            raw_json_str = json_module.dumps(asr_response_data, indent=2)
            if len(raw_json_str) > 5000:
                app.logger.info(f"Raw ASR response (first 5000 chars): {raw_json_str[:5000]}...")
            else:
                app.logger.info(f"Raw ASR response: {raw_json_str}")
            
            if 'segments' in asr_response_data:
                app.logger.info(f"Number of segments: {len(asr_response_data['segments'])}")
                
                # Collect all unique speakers from the response
                all_speakers = set()
                segments_with_speakers = 0
                segments_without_speakers = 0
                
                for segment in asr_response_data['segments']:
                    if 'speaker' in segment and segment['speaker'] is not None:
                        all_speakers.add(segment['speaker'])
                        segments_with_speakers += 1
                    else:
                        segments_without_speakers += 1
                
                app.logger.info(f"Unique speakers found in raw response: {sorted(list(all_speakers))}")
                app.logger.info(f"Segments with speakers: {segments_with_speakers}, without speakers: {segments_without_speakers}")
                
                # Log first few segments for debugging
                for i, segment in enumerate(asr_response_data['segments'][:5]):
                    segment_keys = list(segment.keys())
                    app.logger.info(f"Segment {i} keys: {segment_keys}")
                    app.logger.info(f"Segment {i}: speaker='{segment.get('speaker')}', text='{segment.get('text', '')[:50]}...'")
            
            # Simplify the JSON data
            simplified_segments = []
            if 'segments' in asr_response_data and isinstance(asr_response_data['segments'], list):
                last_known_speaker = None
                
                for i, segment in enumerate(asr_response_data['segments']):
                    speaker = segment.get('speaker')
                    text = segment.get('text', '').strip()
                    
                    # If segment doesn't have a speaker, use the previous segment's speaker
                    if speaker is None:
                        if last_known_speaker is not None:
                            speaker = last_known_speaker
                            app.logger.info(f"Assigned speaker '{speaker}' to segment {i} from previous segment")
                        else:
                            speaker = 'UNKNOWN_SPEAKER'
                            app.logger.warning(f"No previous speaker available for segment {i}, using UNKNOWN_SPEAKER")
                    else:
                        # Update the last known speaker when we have a valid one
                        last_known_speaker = speaker
                    
                    simplified_segments.append({
                        'speaker': speaker,
                        'sentence': text,
                        'start_time': segment.get('start'),
                        'end_time': segment.get('end')
                    })
            
            # Log final simplified segments count
            app.logger.info(f"Created {len(simplified_segments)} simplified segments")
            null_speaker_count = sum(1 for seg in simplified_segments if seg['speaker'] is None)
            if null_speaker_count > 0:
                app.logger.warning(f"Found {null_speaker_count} segments with null speakers in final output")
            
            # Store the simplified JSON as a string
            recording.transcription = json.dumps(simplified_segments)
            
            # Commit the transcription data
            db.session.commit()
            app.logger.info(f"ASR transcription completed for recording {recording_id}.")
            
            # Generate title immediately
            generate_title_task(app_context, recording_id)
            
            # Always auto-generate summary for all recordings
            app.logger.info(f"Auto-generating summary for recording {recording_id}")
            generate_summary_only_task(app_context, recording_id)

        except Exception as e:
            db.session.rollback()
            
            # Handle timeout errors specifically
            error_msg = str(e)
            if "timed out" in error_msg.lower() or "timeout" in error_msg.lower():
                asr_timeout = SystemSetting.get_setting('asr_timeout_seconds', 1800)
                app.logger.error(f"ASR processing TIMED OUT for recording {recording_id} after {asr_timeout} seconds. Consider increasing 'asr_timeout_seconds' in Admin Dashboard > System Settings.")
                user_error_msg = f"ASR processing timed out after {asr_timeout} seconds. The file may be too long for the current timeout setting."
            else:
                # For non-timeout errors, include more detail
                app.logger.error(f"ASR processing FAILED for recording {recording_id}: {error_msg}")
                user_error_msg = f"ASR processing failed: {error_msg}"
            
            recording = db.session.get(Recording, recording_id)
            if recording:
                recording.status = 'FAILED'
                recording.transcription = user_error_msg
                db.session.commit()

def transcribe_audio_task(app_context, recording_id, filepath, filename_for_asr, start_time, language=None, min_speakers=None, max_speakers=None, tag_id=None):
    """Runs the transcription and summarization in a background thread.
    
    Args:
        app_context: Flask app context
        recording_id: ID of the recording to process
        filepath: Path to the audio file
        filename_for_asr: Filename to use for ASR
        start_time: Processing start time
        language: Optional language code override (from upload form)
        min_speakers: Optional minimum speakers override (from upload form)
        max_speakers: Optional maximum speakers override (from upload form)
        tag_id: Optional tag ID to apply custom prompt from
    """
    if USE_ASR_ENDPOINT:
        with app_context:
            recording = db.session.get(Recording, recording_id)
            # Environment variable ASR_DIARIZE overrides user setting
            if 'ASR_DIARIZE' in os.environ:
                diarize_setting = ASR_DIARIZE
            elif USE_ASR_ENDPOINT:
                # When using ASR endpoint, use the configured ASR_DIARIZE value
                diarize_setting = ASR_DIARIZE
            else:
                diarize_setting = recording.owner.diarize if recording.owner else False
            
            # Use language from upload form if provided, otherwise use user's default
            if language:
                user_transcription_language = language
            else:
                user_transcription_language = recording.owner.transcription_language if recording.owner else None
        # Use min/max speakers from upload form (already processed with precedence hierarchy)
        # If None, ASR will auto-detect the number of speakers
        final_min_speakers = min_speakers
        final_max_speakers = max_speakers
        
        transcribe_audio_asr(app_context, recording_id, filepath, filename_for_asr, start_time, 
                           mime_type=recording.mime_type, 
                           language=user_transcription_language, 
                           diarize=diarize_setting,
                           min_speakers=final_min_speakers,
                           max_speakers=final_max_speakers,
                           tag_id=tag_id)
        
        # After ASR task completes, calculate processing time
        with app_context:
            recording = db.session.get(Recording, recording_id)
            if recording.status in ['COMPLETED', 'FAILED']:
                end_time = datetime.utcnow()
                recording.processing_time_seconds = (end_time - start_time).total_seconds()
                db.session.commit()
        return

    with app_context: # Need app context for db operations in thread
        recording = db.session.get(Recording, recording_id)
        if not recording:
            app.logger.error(f"Error: Recording {recording_id} not found for transcription.")
            return

        try:
            app.logger.info(f"Starting transcription for recording {recording_id} ({filename_for_asr})...")
            recording.status = 'PROCESSING'
            db.session.commit()

            # Check if chunking is needed for large files
            needs_chunking = (chunking_service and 
                            ENABLE_CHUNKING and 
                            chunking_service.needs_chunking(filepath, USE_ASR_ENDPOINT))
            
            if needs_chunking:
                app.logger.info(f"File {filepath} is large ({os.path.getsize(filepath)/1024/1024:.1f}MB), using chunking for transcription")
                transcription_text = transcribe_with_chunking(app_context, recording_id, filepath, filename_for_asr)
            else:
                # --- Standard transcription for smaller files ---
                transcription_text = transcribe_single_file(filepath, recording)
            
            recording.transcription = transcription_text
            app.logger.info(f"Transcription completed for recording {recording_id}. Text length: {len(recording.transcription)}")
            
            # Generate title immediately
            generate_title_task(app_context, recording_id)
            
            # Always auto-generate summary for all recordings
            app.logger.info(f"Auto-generating summary for recording {recording_id}")
            generate_summary_only_task(app_context, recording_id)

        except Exception as e:
            db.session.rollback() # Rollback if any step failed critically
            app.logger.error(f"Processing FAILED for recording {recording_id}: {str(e)}", exc_info=True)
            # Retrieve recording again in case session was rolled back
            recording = db.session.get(Recording, recording_id)
            if recording:
                 # Ensure status reflects failure even after rollback/retrieve attempt
                if recording.status not in ['COMPLETED', 'FAILED']: # Avoid overwriting final state
                    recording.status = 'FAILED'
                if not recording.transcription: # If transcription itself failed
                     recording.transcription = f"Processing failed: {str(e)}"
                # Add error note to summary if appropriate stage was reached
                if recording.status == 'SUMMARIZING' and not recording.summary:
                     recording.summary = f"[Processing failed during summarization: {str(e)}]"
                
                end_time = datetime.utcnow()
                recording.processing_time_seconds = (end_time - start_time).total_seconds()
                db.session.commit()

def transcribe_single_file(filepath, recording):
    """Transcribe a single audio file using OpenAI Whisper API."""
    
    # Check if we need to extract audio from video container
    actual_filepath = filepath
    mime_type = recording.mime_type if recording else None
    
    # Detect video containers
    is_video = False
    if mime_type:
        is_video = mime_type.startswith('video/')
    else:
        # Fallback to extension-based detection
        is_video = filepath.lower().endswith(('.mp4', '.mov', '.avi', '.mkv', '.webm', '.wmv', '.3gp'))
    
    if is_video:
        app.logger.info(f"Video container detected for Whisper transcription, extracting audio...")
        try:
            # Extract audio from video
            audio_filepath, audio_mime_type = extract_audio_from_video(filepath, 'wav')
            actual_filepath = audio_filepath
            
            # Update recording with extracted audio path and new MIME type if recording exists
            if recording:
                recording.audio_path = audio_filepath
                recording.mime_type = audio_mime_type
                db.session.commit()
            
            app.logger.info(f"Audio extracted successfully for Whisper: {audio_filepath}")
        except Exception as e:
            app.logger.error(f"Failed to extract audio from video for Whisper: {str(e)}")
            if recording:
                recording.status = 'FAILED'
                recording.error_msg = f"Audio extraction failed: {str(e)}"
                db.session.commit()
            raise Exception(f"Audio extraction failed: {str(e)}")
    
    # List of formats supported by Whisper API
    WHISPER_SUPPORTED_FORMATS = ['flac', 'm4a', 'mp3', 'mp4', 'mpeg', 'mpga', 'oga', 'ogg', 'wav', 'webm']
    
    # Check if the file format needs conversion
    file_ext = os.path.splitext(actual_filepath)[1].lower().lstrip('.')
    converted_filepath = None
    
    try:
        with open(actual_filepath, 'rb') as audio_file:
            transcription_client = OpenAI(
                api_key=transcription_api_key,
                base_url=transcription_base_url,
                http_client=http_client_no_proxy
            )
            whisper_model = os.environ.get("WHISPER_MODEL", "Systran/faster-distil-whisper-large-v3")
            
            user_transcription_language = None
            if recording and recording.owner:
                user_transcription_language = recording.owner.transcription_language
            
            transcription_language = user_transcription_language

            transcription_params = {
                "model": whisper_model,
                "file": audio_file
            }

            if transcription_language:
                transcription_params["language"] = transcription_language
                app.logger.info(f"Using transcription language: {transcription_language}")
            else:
                app.logger.info("Transcription language not set, using auto-detection or service default.")

            transcript = transcription_client.audio.transcriptions.create(**transcription_params)
            return transcript.text
            
    except Exception as e:
        # Check if it's a format error
        error_message = str(e)
        if "Invalid file format" in error_message or "Supported formats" in error_message:
            app.logger.warning(f"Unsupported audio format '{file_ext}' detected, converting to MP3...")
            
            # Convert to MP3
            import tempfile
            temp_mp3_filepath = None
            try:
                with tempfile.NamedTemporaryFile(suffix='.mp3', delete=False) as temp_mp3:
                    temp_mp3_filepath = temp_mp3.name
                
                # Use ffmpeg to convert to MP3 with consistent settings
                subprocess.run(
                    ['ffmpeg', '-i', actual_filepath, '-y', '-acodec', 'libmp3lame', '-b:a', '128k', '-ar', '44100', temp_mp3_filepath],
                    check=True,
                    capture_output=True
                )
                app.logger.info(f"Successfully converted {actual_filepath} to MP3 format")
                converted_filepath = temp_mp3_filepath
                
                # Retry transcription with converted file
                with open(converted_filepath, 'rb') as audio_file:
                    transcription_client = OpenAI(
                        api_key=transcription_api_key,
                        base_url=transcription_base_url,
                        http_client=http_client_no_proxy
                    )
                    
                    transcription_params = {
                        "model": whisper_model,
                        "file": audio_file
                    }

                    if transcription_language:
                        transcription_params["language"] = transcription_language

                    transcript = transcription_client.audio.transcriptions.create(**transcription_params)
                    return transcript.text
                    
            finally:
                # Clean up temporary converted file
                if converted_filepath and os.path.exists(converted_filepath):
                    try:
                        os.unlink(converted_filepath)
                        app.logger.info(f"Cleaned up temporary converted file: {converted_filepath}")
                    except Exception as cleanup_error:
                        app.logger.warning(f"Failed to clean up temporary file {converted_filepath}: {cleanup_error}")
        else:
            # Re-raise if it's not a format error
            raise

def transcribe_with_chunking(app_context, recording_id, filepath, filename_for_asr):
    """Transcribe a large audio file using chunking."""
    import tempfile
    
    with app_context:
        recording = db.session.get(Recording, recording_id)
        if not recording:
            raise ValueError(f"Recording {recording_id} not found")
    
    # Create temporary directory for chunks
    with tempfile.TemporaryDirectory() as temp_dir:
        try:
            # Create chunks
            app.logger.info(f"Creating chunks for large file: {filepath}")
            chunks = chunking_service.create_chunks(filepath, temp_dir)
            
            if not chunks:
                raise ChunkProcessingError("No chunks were created from the audio file")
            
            app.logger.info(f"Created {len(chunks)} chunks, processing each with Whisper API...")
            
            # Process each chunk with proper timeout and retry handling
            chunk_results = []
            
            # Create HTTP client with proper timeouts
            timeout_config = httpx.Timeout(
                connect=30.0,    # 30 seconds to establish connection
                read=300.0,      # 5 minutes to read response (for large audio files)
                write=60.0,      # 1 minute to write request
                pool=10.0        # 10 seconds to get connection from pool
            )
            
            http_client_with_timeout = httpx.Client(
                verify=True,
                timeout=timeout_config,
                limits=httpx.Limits(max_connections=5, max_keepalive_connections=2)
            )
            
            transcription_client = OpenAI(
                api_key=transcription_api_key,
                base_url=transcription_base_url,
                http_client=http_client_with_timeout,
                max_retries=3,  # Increased retries for better reliability
                timeout=300.0   # 5 minute timeout for API calls
            )
            whisper_model = os.environ.get("WHISPER_MODEL", "Systran/faster-distil-whisper-large-v3")
            
            # Get user language preference
            user_transcription_language = None
            with app_context:
                recording = db.session.get(Recording, recording_id)
                if recording and recording.owner:
                    user_transcription_language = recording.owner.transcription_language
            
            for i, chunk in enumerate(chunks):
                max_chunk_retries = 3
                chunk_retry_count = 0
                chunk_success = False
                
                while chunk_retry_count < max_chunk_retries and not chunk_success:
                    try:
                        retry_suffix = f" (retry {chunk_retry_count + 1}/{max_chunk_retries})" if chunk_retry_count > 0 else ""
                        app.logger.info(f"Processing chunk {i+1}/{len(chunks)}: {chunk['filename']} ({chunk['size_mb']:.1f}MB){retry_suffix}")
                        
                        # Log detailed timing for each step
                        step_start_time = time.time()
                        
                        # Step 1: File opening
                        file_open_start = time.time()
                        with open(chunk['path'], 'rb') as chunk_file:
                            file_open_time = time.time() - file_open_start
                            app.logger.info(f"Chunk {i+1}: File opened in {file_open_time:.2f}s")
                            
                            # Step 2: Prepare transcription parameters
                            param_start = time.time()
                            transcription_params = {
                                "model": whisper_model,
                                "file": chunk_file
                            }
                            
                            if user_transcription_language:
                                transcription_params["language"] = user_transcription_language
                            
                            param_time = time.time() - param_start
                            app.logger.info(f"Chunk {i+1}: Parameters prepared in {param_time:.2f}s")
                            
                            # Step 3: API call with detailed timing
                            api_start = time.time()
                            app.logger.info(f"Chunk {i+1}: Starting API call to {transcription_base_url}")
                            
                            # Log connection details
                            app.logger.info(f"Chunk {i+1}: Using timeout config - connect: 30s, read: 300s, write: 60s")
                            app.logger.info(f"Chunk {i+1}: Max retries: 2, API timeout: 300s")
                            
                            try:
                                transcript = transcription_client.audio.transcriptions.create(**transcription_params)
                            except Exception as chunk_error:
                                # Check if it's a format error (unlikely for chunks since they're MP3, but handle it)
                                error_msg = str(chunk_error)
                                if "Invalid file format" in error_msg or "Supported formats" in error_msg:
                                    app.logger.warning(f"Chunk {i+1} format issue, attempting conversion...")
                                    # Convert chunk to MP3 if needed
                                    import tempfile
                                    with tempfile.NamedTemporaryFile(suffix='.mp3', delete=False) as temp_mp3:
                                        temp_mp3_path = temp_mp3.name
                                    try:
                                        subprocess.run(
                                            ['ffmpeg', '-i', chunk['path'], '-y', '-acodec', 'libmp3lame', '-b:a', '128k', '-ar', '44100', temp_mp3_path],
                                            check=True,
                                            capture_output=True
                                        )
                                        with open(temp_mp3_path, 'rb') as converted_chunk:
                                            transcription_params['file'] = converted_chunk
                                            transcript = transcription_client.audio.transcriptions.create(**transcription_params)
                                    finally:
                                        if os.path.exists(temp_mp3_path):
                                            os.unlink(temp_mp3_path)
                                else:
                                    raise
                            
                            api_time = time.time() - api_start
                            app.logger.info(f"Chunk {i+1}: API call completed in {api_time:.2f}s")
                            
                            # Step 4: Process response
                            response_start = time.time()
                            chunk_result = {
                                'index': chunk['index'],
                                'start_time': chunk['start_time'],
                                'end_time': chunk['end_time'],
                                'duration': chunk['duration'],
                                'size_mb': chunk['size_mb'],
                                'transcription': transcript.text,
                                'filename': chunk['filename'],
                                'processing_time': api_time  # Store the actual API processing time
                            }
                            chunk_results.append(chunk_result)
                            response_time = time.time() - response_start
                            
                            total_time = time.time() - step_start_time
                            app.logger.info(f"Chunk {i+1}: Response processed in {response_time:.2f}s")
                            app.logger.info(f"Chunk {i+1}: Total processing time: {total_time:.2f}s")
                            app.logger.info(f"Chunk {i+1} transcribed successfully: {len(transcript.text)} characters")
                            chunk_success = True
                            
                    except Exception as chunk_error:
                        chunk_retry_count += 1
                        error_msg = str(chunk_error)
                        
                        if chunk_retry_count < max_chunk_retries:
                            # Determine wait time based on error type
                            if "timeout" in error_msg.lower() or "timed out" in error_msg.lower():
                                wait_time = 30  # 30 seconds for timeout errors
                            elif "rate limit" in error_msg.lower():
                                wait_time = 60  # 1 minute for rate limit errors
                            else:
                                wait_time = 15  # 15 seconds for other errors
                            
                            app.logger.warning(f"Chunk {i+1} failed (attempt {chunk_retry_count}/{max_chunk_retries}): {chunk_error}. Retrying in {wait_time} seconds...")
                            time.sleep(wait_time)
                        else:
                            app.logger.error(f"Chunk {i+1} failed after {max_chunk_retries} attempts: {chunk_error}")
                            # Add failed chunk to results
                            chunk_result = {
                                'index': chunk['index'],
                                'start_time': chunk['start_time'],
                                'end_time': chunk['end_time'],
                                'transcription': f"[Chunk {i+1} transcription failed after {max_chunk_retries} attempts: {str(chunk_error)}]",
                                'filename': chunk['filename']
                            }
                            chunk_results.append(chunk_result)
                
                # Add small delay between chunks to avoid overwhelming the API
                if i < len(chunks) - 1:  # Don't delay after the last chunk
                    time.sleep(2)
            
            # Merge transcriptions
            app.logger.info(f"Merging {len(chunk_results)} chunk transcriptions...")
            merged_transcription = chunking_service.merge_transcriptions(chunk_results)
            
            if not merged_transcription.strip():
                raise ChunkProcessingError("Merged transcription is empty")
            
            # Log detailed performance statistics and analysis
            chunking_service.log_processing_statistics(chunk_results)
            
            # Get performance recommendations
            recommendations = chunking_service.get_performance_recommendations(chunk_results)
            if recommendations:
                app.logger.info("=== PERFORMANCE RECOMMENDATIONS ===")
                for i, rec in enumerate(recommendations, 1):
                    app.logger.info(f"{i}. {rec}")
                app.logger.info("=== END RECOMMENDATIONS ===")
            
            app.logger.info(f"Chunked transcription completed. Final length: {len(merged_transcription)} characters")
            return merged_transcription
            
        except Exception as e:
            app.logger.error(f"Chunking transcription failed for {filepath}: {e}")
            # Clean up chunks if they exist
            if 'chunks' in locals():
                chunking_service.cleanup_chunks(chunks)
            raise ChunkProcessingError(f"Chunked transcription failed: {str(e)}")
        finally:
            # Cleanup is handled by tempfile.TemporaryDirectory context manager
            pass

@app.route('/speakers', methods=['GET'])
@login_required
def get_speakers():
    """Get all speakers for the current user, ordered by usage frequency and recency."""
    try:
        speakers = Speaker.query.filter_by(user_id=current_user.id)\
                               .order_by(Speaker.use_count.desc(), Speaker.last_used.desc())\
                               .all()
        return jsonify([speaker.to_dict() for speaker in speakers])
    except Exception as e:
        app.logger.error(f"Error fetching speakers: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/speakers/search', methods=['GET'])
@login_required
def search_speakers():
    """Search speakers by name for autocomplete functionality."""
    try:
        query = request.args.get('q', '').strip()
        if not query:
            return jsonify([])
        
        speakers = Speaker.query.filter_by(user_id=current_user.id)\
                               .filter(Speaker.name.ilike(f'%{query}%'))\
                               .order_by(Speaker.use_count.desc(), Speaker.last_used.desc())\
                               .limit(10)\
                               .all()
        
        return jsonify([speaker.to_dict() for speaker in speakers])
    except Exception as e:
        app.logger.error(f"Error searching speakers: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/speakers', methods=['POST'])
@login_required
def create_speaker():
    """Create a new speaker or update existing one."""
    try:
        data = request.json
        name = data.get('name', '').strip()
        
        if not name:
            return jsonify({'error': 'Speaker name is required'}), 400
        
        # Check if speaker already exists for this user
        existing_speaker = Speaker.query.filter_by(user_id=current_user.id, name=name).first()
        
        if existing_speaker:
            # Update usage statistics
            existing_speaker.use_count += 1
            existing_speaker.last_used = datetime.utcnow()
            db.session.commit()
            return jsonify(existing_speaker.to_dict())
        else:
            # Create new speaker
            speaker = Speaker(
                name=name,
                user_id=current_user.id,
                use_count=1,
                created_at=datetime.utcnow(),
                last_used=datetime.utcnow()
            )
            db.session.add(speaker)
            db.session.commit()
            return jsonify(speaker.to_dict()), 201
            
    except Exception as e:
        db.session.rollback()
        app.logger.error(f"Error creating speaker: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/speakers/<int:speaker_id>', methods=['DELETE'])
@login_required
def delete_speaker(speaker_id):
    """Delete a speaker."""
    try:
        speaker = Speaker.query.filter_by(id=speaker_id, user_id=current_user.id).first()
        if not speaker:
            return jsonify({'error': 'Speaker not found'}), 404
        
        db.session.delete(speaker)
        db.session.commit()
        return jsonify({'success': True})
        
    except Exception as e:
        db.session.rollback()
        app.logger.error(f"Error deleting speaker: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/speakers/delete_all', methods=['DELETE'])
@login_required
def delete_all_speakers():
    """Delete all speakers for the current user."""
    try:
        deleted_count = Speaker.query.filter_by(user_id=current_user.id).delete()
        db.session.commit()
        return jsonify({'success': True, 'deleted_count': deleted_count})
        
    except Exception as e:
        db.session.rollback()
        app.logger.error(f"Error deleting all speakers: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/recording/<int:recording_id>/download/transcript')
@login_required
def download_transcript_with_template(recording_id):
    """Download transcript with custom template formatting."""
    try:
        import re
        from datetime import timedelta

        recording = db.session.get(Recording, recording_id)
        if not recording:
            return jsonify({'error': 'Recording not found'}), 404

        # Check if the recording belongs to the current user
        if recording.user_id and recording.user_id != current_user.id:
            return jsonify({'error': 'You do not have permission to access this recording'}), 403

        if not recording.transcription:
            return jsonify({'error': 'No transcription available for this recording'}), 400

        # Get template ID from query params
        template_id = request.args.get('template_id', type=int)

        # Get the template
        if template_id:
            template = TranscriptTemplate.query.filter_by(
                id=template_id,
                user_id=current_user.id
            ).first()
        else:
            # Use default template
            template = TranscriptTemplate.query.filter_by(
                user_id=current_user.id,
                is_default=True
            ).first()

        # If no template found, use a basic format
        if not template:
            template_format = "[{{speaker}}]: {{text}}"
        else:
            template_format = template.template

        # Helper functions for formatting
        def format_time(seconds):
            """Format seconds to HH:MM:SS"""
            if seconds is None:
                return "00:00:00"
            td = timedelta(seconds=seconds)
            hours = int(td.total_seconds() // 3600)
            minutes = int((td.total_seconds() % 3600) // 60)
            secs = int(td.total_seconds() % 60)
            return f"{hours:02d}:{minutes:02d}:{secs:02d}"

        def format_srt_time(seconds):
            """Format seconds to SRT format HH:MM:SS,mmm"""
            if seconds is None:
                return "00:00:00,000"
            td = timedelta(seconds=seconds)
            hours = int(td.total_seconds() // 3600)
            minutes = int((td.total_seconds() % 3600) // 60)
            secs = int(td.total_seconds() % 60)
            millis = int((td.total_seconds() % 1) * 1000)
            return f"{hours:02d}:{minutes:02d}:{secs:02d},{millis:03d}"

        # Parse transcription JSON
        try:
            transcription_data = json.loads(recording.transcription)
        except:
            return jsonify({'error': 'Invalid transcription format'}), 400

        # Generate formatted transcript
        output_lines = []
        for index, segment in enumerate(transcription_data, 1):
            line = template_format

            # Replace variables
            replacements = {
                '{{index}}': str(index),
                '{{speaker}}': segment.get('speaker', 'Unknown'),
                '{{text}}': segment.get('sentence', ''),
                '{{start_time}}': format_time(segment.get('start_time')),
                '{{end_time}}': format_time(segment.get('end_time')),
            }

            for key, value in replacements.items():
                line = line.replace(key, value)

            # Handle filters
            # Upper case filter
            line = re.sub(r'{{(.*?)\|upper}}', lambda m: replacements.get('{{' + m.group(1) + '}}', '').upper(), line)
            # SRT time filter
            line = re.sub(r'{{start_time\|srt}}', format_srt_time(segment.get('start_time')), line)
            line = re.sub(r'{{end_time\|srt}}', format_srt_time(segment.get('end_time')), line)

            output_lines.append(line)

        # Join lines
        formatted_transcript = '\n'.join(output_lines)

        # Create response
        response = make_response(formatted_transcript)
        filename = f"{recording.title or 'transcript'}_{template.name if template else 'formatted'}.txt"
        filename = re.sub(r'[^a-zA-Z0-9_\-\.]', '_', filename)
        response.headers['Content-Type'] = 'text/plain; charset=utf-8'
        response.headers['Content-Disposition'] = f'attachment; filename="{filename}"'

        return response

    except Exception as e:
        app.logger.error(f"Error downloading transcript: {e}")
        return jsonify({'error': 'Failed to generate transcript download'}), 500

@app.route('/recording/<int:recording_id>/download/summary')
@login_required
def download_summary_word(recording_id):
    """Download recording summary as a Word document."""
    try:
        from docx import Document
        from docx.shared import Inches
        import re
        from io import BytesIO
        
        recording = db.session.get(Recording, recording_id)
        if not recording:
            return jsonify({'error': 'Recording not found'}), 404
            
        # Check if the recording belongs to the current user
        if recording.user_id and recording.user_id != current_user.id:
            return jsonify({'error': 'You do not have permission to access this recording'}), 403
            
        if not recording.summary:
            return jsonify({'error': 'No summary available for this recording'}), 400
        
        # Create Word document
        doc = Document()
        
        # Add title
        title = doc.add_heading(f'Summary: {recording.title or "Untitled Recording"}', 0)
        
        # Add metadata
        doc.add_paragraph(f'Uploaded: {recording.created_at.strftime("%Y-%m-%d %H:%M")}')
        if recording.meeting_date:
            doc.add_paragraph(f'Recording Date: {recording.meeting_date.strftime("%Y-%m-%d")}')
        if recording.participants:
            doc.add_paragraph(f'Participants: {recording.participants}')
        if recording.tags:
            tags_str = ', '.join([tag.name for tag in recording.tags])
            doc.add_paragraph(f'Tags: {tags_str}')
        doc.add_paragraph('')  # Empty line
        
        # Convert markdown-like formatting to Word formatting
        def add_formatted_content(paragraph, text):
            # Handle bold text (**text**)
            bold_pattern = r'\*\*(.*?)\*\*'
            parts = re.split(bold_pattern, text)
            
            for i, part in enumerate(parts):
                if i % 2 == 0:  # Regular text
                    if part:
                        paragraph.add_run(part)
                else:  # Bold text
                    if part:
                        paragraph.add_run(part).bold = True
        
        # Split summary into paragraphs and process
        summary_lines = recording.summary.split('\n')
        current_paragraph = None
        
        for line in summary_lines:
            line = line.strip()
            if not line:
                current_paragraph = None
                continue
                
            # Check if line starts with markdown heading
            if line.startswith('# '):
                doc.add_heading(line[2:], 1)
                current_paragraph = None
            elif line.startswith('## '):
                doc.add_heading(line[3:], 2)
                current_paragraph = None
            elif line.startswith('### '):
                doc.add_heading(line[4:], 3)
                current_paragraph = None
            elif line.startswith('- ') or line.startswith('* '):
                # Bullet point
                p = doc.add_paragraph(style='List Bullet')
                add_formatted_content(p, line[2:])
                current_paragraph = None
            elif re.match(r'^\d+\.', line):
                # Numbered list
                p = doc.add_paragraph(style='List Number')
                add_formatted_content(p, re.sub(r'^\d+\.\s*', '', line))
                current_paragraph = None
            else:
                # Regular paragraph
                if current_paragraph is None:
                    current_paragraph = doc.add_paragraph()
                else:
                    current_paragraph.add_run('\n')
                add_formatted_content(current_paragraph, line)
        
        # Save to BytesIO
        doc_stream = BytesIO()
        doc.save(doc_stream)
        doc_stream.seek(0)
        
        # Create safe filename
        safe_title = re.sub(r'[<>:"/\\|?*]', '', recording.title or 'Untitled')
        safe_title = re.sub(r'[-\s]+', '-', safe_title).strip('-')
        filename = f'summary-{safe_title}.docx' if safe_title else f'summary-recording-{recording_id}.docx'

        # Create ASCII fallback for send_file - if title has non-ASCII chars, use generic name with ID
        ascii_filename = filename.encode('ascii', 'ignore').decode('ascii')
        if not ascii_filename.strip() or ascii_filename.strip() in ['summary-.docx', 'summary-recording-.docx']:
            ascii_filename = f'summary-recording-{recording_id}.docx'

        response = send_file(
            doc_stream,
            as_attachment=False,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        # Properly encode filename for international characters
        # Check if filename contains non-ASCII characters
        try:
            # Try to encode as ASCII - if this works, use simple format
            filename.encode('ascii')
            # ASCII-only filename, use simple format
            response.headers['Content-Disposition'] = f'attachment; filename="{filename}"'
        except UnicodeEncodeError:
            # Contains non-ASCII characters, use proper RFC 2231 encoding
            try:
                # Use Python's built-in RFC 2231 encoder
                encoded_value = encode_rfc2231(filename, charset='utf-8')
                header_value = f'attachment; filename*={encoded_value}'
                app.logger.error(f"DEBUG CHINESE FILENAME (RFC2231): Original='{filename}', Header='{header_value}'")
                response.headers['Content-Disposition'] = header_value
            except Exception as e:
                # Fallback to simple attachment with generic name
                app.logger.error(f"RFC2231 encoding failed: {e}, using fallback")
                response.headers['Content-Disposition'] = f'attachment; filename="download-{recording_id}.docx"'
        return response
        
    except Exception as e:
        app.logger.error(f"Error generating summary Word document: {e}")
        return jsonify({'error': 'Failed to generate Word document'}), 500

@app.route('/recording/<int:recording_id>/download/chat', methods=['POST'])
@login_required
def download_chat_word(recording_id):
    """Download chat conversation as a Word document."""
    try:
        from docx import Document
        from docx.shared import Inches
        import re
        from io import BytesIO
        
        recording = db.session.get(Recording, recording_id)
        if not recording:
            return jsonify({'error': 'Recording not found'}), 404
            
        # Check if the recording belongs to the current user
        if recording.user_id and recording.user_id != current_user.id:
            return jsonify({'error': 'You do not have permission to access this recording'}), 403
        
        # Get chat messages from request
        data = request.json
        if not data or 'messages' not in data:
            return jsonify({'error': 'No messages provided'}), 400
        
        messages = data['messages']
        if not messages:
            return jsonify({'error': 'No messages to download'}), 400
        
        # Create Word document
        doc = Document()
        
        # Add title
        title = doc.add_heading(f'Chat Conversation: {recording.title or "Untitled Recording"}', 0)
        
        # Add metadata
        doc.add_paragraph(f'Recording Date: {recording.created_at.strftime("%Y-%m-%d %H:%M")}')
        doc.add_paragraph(f'Chat Export Date: {datetime.utcnow().strftime("%Y-%m-%d %H:%M")}')
        doc.add_paragraph('')  # Empty line
        
        # Add chat messages
        for message in messages:
            role = message.get('role', 'unknown')
            content = message.get('content', '')
            thinking = message.get('thinking', '')
            
            # Add role header
            if role == 'user':
                p = doc.add_paragraph()
                p.add_run('You: ').bold = True
            elif role == 'assistant':
                p = doc.add_paragraph()
                p.add_run('Assistant: ').bold = True
            else:
                p = doc.add_paragraph()
                p.add_run(f'{role.title()}: ').bold = True
            
            # Add thinking content if present
            if thinking and role == 'assistant':
                p = doc.add_paragraph()
                p.add_run('[Model Reasoning]\n').italic = True
                p.add_run(thinking).italic = True
                doc.add_paragraph('')  # Empty line
            
            # Add message content
            # Convert markdown-like formatting
            lines = content.split('\n')
            current_paragraph = None
            
            for line in lines:
                if line.startswith('# '):
                    # Heading
                    doc.add_heading(line[2:], level=1)
                    current_paragraph = None
                elif line.startswith('## '):
                    # Subheading
                    doc.add_heading(line[3:], level=2)
                    current_paragraph = None
                elif line.startswith('### '):
                    # Sub-subheading
                    doc.add_heading(line[4:], level=3)
                    current_paragraph = None
                elif line.startswith('- ') or line.startswith('* '):
                    # Bullet point
                    p = doc.add_paragraph(style='List Bullet')
                    # Handle bold text
                    text = line[2:]
                    bold_pattern = r'\*\*(.*?)\*\*'
                    parts = re.split(bold_pattern, text)
                    for i, part in enumerate(parts):
                        if i % 2 == 0:  # Regular text
                            if part:
                                p.add_run(part)
                        else:  # Bold text
                            if part:
                                p.add_run(part).bold = True
                    current_paragraph = None
                elif re.match(r'^\d+\.', line):
                    # Numbered list
                    p = doc.add_paragraph(style='List Number')
                    text = re.sub(r'^\d+\.\s*', '', line)
                    # Handle bold text
                    bold_pattern = r'\*\*(.*?)\*\*'
                    parts = re.split(bold_pattern, text)
                    for i, part in enumerate(parts):
                        if i % 2 == 0:  # Regular text
                            if part:
                                p.add_run(part)
                        else:  # Bold text
                            if part:
                                p.add_run(part).bold = True
                    current_paragraph = None
                else:
                    # Regular paragraph
                    if current_paragraph is None:
                        current_paragraph = doc.add_paragraph()
                    else:
                        current_paragraph.add_run('\n')
                    
                    # Handle bold text
                    bold_pattern = r'\*\*(.*?)\*\*'
                    parts = re.split(bold_pattern, line)
                    for i, part in enumerate(parts):
                        if i % 2 == 0:  # Regular text
                            if part:
                                current_paragraph.add_run(part)
                        else:  # Bold text
                            if part:
                                current_paragraph.add_run(part).bold = True
            
            doc.add_paragraph('')  # Empty line between messages
        
        # Save to BytesIO
        doc_stream = BytesIO()
        doc.save(doc_stream)
        doc_stream.seek(0)
        
        # Create safe filename
        safe_title = re.sub(r'[<>:"/\\|?*]', '', recording.title or 'Untitled')
        safe_title = re.sub(r'[-\s]+', '-', safe_title).strip('-')
        filename = f'chat-{safe_title}.docx' if safe_title else f'chat-recording-{recording_id}.docx'

        # Create ASCII fallback for send_file - if title has non-ASCII chars, use generic name with ID
        ascii_filename = filename.encode('ascii', 'ignore').decode('ascii')
        if not ascii_filename.strip() or ascii_filename.strip() in ['chat-.docx', 'chat-recording-.docx']:
            ascii_filename = f'chat-recording-{recording_id}.docx'

        response = send_file(
            doc_stream,
            as_attachment=False,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

        # Properly encode filename for international characters
        # Check if filename contains non-ASCII characters
        try:
            # Try to encode as ASCII - if this works, use simple format
            filename.encode('ascii')
            # ASCII-only filename, use simple format
            response.headers['Content-Disposition'] = f'attachment; filename="{filename}"'
        except UnicodeEncodeError:
            # Contains non-ASCII characters, use proper RFC 2231 encoding
            try:
                # Use Python's built-in RFC 2231 encoder
                encoded_value = encode_rfc2231(filename, charset='utf-8')
                header_value = f'attachment; filename*={encoded_value}'
                app.logger.error(f"DEBUG CHINESE FILENAME (RFC2231): Original='{filename}', Header='{header_value}'")
                response.headers['Content-Disposition'] = header_value
            except Exception as e:
                # Fallback to simple attachment with generic name
                app.logger.error(f"RFC2231 encoding failed: {e}, using fallback")
                response.headers['Content-Disposition'] = f'attachment; filename="download-{recording_id}.docx"'
        return response
        
    except Exception as e:
        app.logger.error(f"Error generating chat Word document: {e}")
        return jsonify({'error': 'Failed to generate Word document'}), 500

@app.route('/api/recording/<int:recording_id>/events', methods=['GET'])
@login_required
def get_recording_events(recording_id):
    """Get all events extracted from a recording."""
    try:
        recording = db.session.get(Recording, recording_id)
        if not recording:
            return jsonify({'error': 'Recording not found'}), 404

        if recording.user_id and recording.user_id != current_user.id:
            return jsonify({'error': 'Unauthorized'}), 403

        events = Event.query.filter_by(recording_id=recording_id).all()
        return jsonify({'events': [event.to_dict() for event in events]})

    except Exception as e:
        app.logger.error(f"Error fetching events for recording {recording_id}: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/event/<int:event_id>/ics', methods=['GET'])
@login_required
def download_event_ics(event_id):
    """Generate and download an ICS file for a single event."""
    try:
        event = db.session.get(Event, event_id)
        if not event:
            return jsonify({'error': 'Event not found'}), 404

        # Check permissions through recording ownership
        if event.recording.user_id != current_user.id:
            return jsonify({'error': 'Unauthorized'}), 403

        # Generate ICS content
        ics_content = generate_ics_content(event)

        # Create response with ICS file
        response = make_response(ics_content)
        response.headers['Content-Type'] = 'text/calendar; charset=utf-8'
        response.headers['Content-Disposition'] = f'attachment; filename="{secure_filename(event.title)}.ics"'

        return response

    except Exception as e:
        app.logger.error(f"Error generating ICS for event {event_id}: {e}")
        return jsonify({'error': str(e)}), 500

def generate_ics_content(event):
    """Generate ICS calendar file content for an event."""
    import uuid
    from datetime import datetime, timedelta

    # Generate unique ID for the event
    uid = f"{event.id}-{uuid.uuid4()}@speakr.app"

    # Format dates in iCalendar format (YYYYMMDDTHHMMSS)
    def format_ical_date(dt):
        if dt:
            return dt.strftime('%Y%m%dT%H%M%S')
        return None

    # Start building ICS content
    lines = [
        'BEGIN:VCALENDAR',
        'VERSION:2.0',
        'PRODID:-//Speakr//Event Export//EN',
        'CALSCALE:GREGORIAN',
        'METHOD:PUBLISH',
        'BEGIN:VEVENT',
        f'UID:{uid}',
        f'DTSTAMP:{format_ical_date(datetime.utcnow())}',
    ]

    # Add event details
    if event.start_datetime:
        lines.append(f'DTSTART:{format_ical_date(event.start_datetime)}')

    if event.end_datetime:
        lines.append(f'DTEND:{format_ical_date(event.end_datetime)}')
    elif event.start_datetime:
        # If no end time, default to 1 hour after start
        end_time = event.start_datetime + timedelta(hours=1)
        lines.append(f'DTEND:{format_ical_date(end_time)}')

    # Add title and description
    lines.append(f'SUMMARY:{escape_ical_text(event.title)}')

    if event.description:
        lines.append(f'DESCRIPTION:{escape_ical_text(event.description)}')

    # Add location if available
    if event.location:
        lines.append(f'LOCATION:{escape_ical_text(event.location)}')

    # Add attendees if available
    if event.attendees:
        try:
            attendees_list = json.loads(event.attendees)
            for attendee in attendees_list:
                if attendee:
                    lines.append(f'ATTENDEE:CN={escape_ical_text(attendee)}:mailto:{attendee.replace(" ", ".").lower()}@example.com')
        except:
            pass

    # Add reminder/alarm if specified
    if event.reminder_minutes and event.reminder_minutes > 0:
        lines.extend([
            'BEGIN:VALARM',
            'TRIGGER:-PT{}M'.format(event.reminder_minutes),
            'ACTION:DISPLAY',
            f'DESCRIPTION:Reminder: {escape_ical_text(event.title)}',
            'END:VALARM'
        ])

    # Close event and calendar
    lines.extend([
        'STATUS:CONFIRMED',
        'TRANSP:OPAQUE',
        'END:VEVENT',
        'END:VCALENDAR'
    ])

    return '\r\n'.join(lines)

def escape_ical_text(text):
    """Escape special characters for iCalendar format."""
    if not text:
        return ''
    # Escape special characters
    text = str(text)
    text = text.replace('\\', '\\\\')
    text = text.replace(',', '\\,')
    text = text.replace(';', '\\;')
    text = text.replace('\n', '\\n')
    return text

@app.route('/recording/<int:recording_id>/download/notes')
@login_required
def download_notes_word(recording_id):
    """Download recording notes as a Word document."""
    try:
        from docx import Document
        from docx.shared import Inches
        import re
        from io import BytesIO
        
        recording = db.session.get(Recording, recording_id)
        if not recording:
            return jsonify({'error': 'Recording not found'}), 404
            
        # Check if the recording belongs to the current user
        if recording.user_id and recording.user_id != current_user.id:
            return jsonify({'error': 'You do not have permission to access this recording'}), 403
            
        if not recording.notes:
            return jsonify({'error': 'No notes available for this recording'}), 400
        
        # Create Word document
        doc = Document()
        
        # Add title
        title = doc.add_heading(f'Notes: {recording.title or "Untitled Recording"}', 0)
        
        # Add metadata
        doc.add_paragraph(f'Uploaded: {recording.created_at.strftime("%Y-%m-%d %H:%M")}')
        if recording.meeting_date:
            doc.add_paragraph(f'Recording Date: {recording.meeting_date.strftime("%Y-%m-%d")}')
        if recording.participants:
            doc.add_paragraph(f'Participants: {recording.participants}')
        if recording.tags:
            tags_str = ', '.join([tag.name for tag in recording.tags])
            doc.add_paragraph(f'Tags: {tags_str}')
        doc.add_paragraph('')  # Empty line
        
        # Convert markdown-like formatting to Word formatting
        def add_formatted_content(paragraph, text):
            # Handle bold text (**text**)
            bold_pattern = r'\*\*(.*?)\*\*'
            parts = re.split(bold_pattern, text)
            
            for i, part in enumerate(parts):
                if i % 2 == 0:  # Regular text
                    if part:
                        paragraph.add_run(part)
                else:  # Bold text
                    if part:
                        paragraph.add_run(part).bold = True
        
        # Split notes into paragraphs and process
        notes_lines = recording.notes.split('\n')
        current_paragraph = None
        
        for line in notes_lines:
            line = line.strip()
            if not line:
                current_paragraph = None
                continue
                
            # Check if line starts with markdown heading
            if line.startswith('# '):
                doc.add_heading(line[2:], 1)
                current_paragraph = None
            elif line.startswith('## '):
                doc.add_heading(line[3:], 2)
                current_paragraph = None
            elif line.startswith('### '):
                doc.add_heading(line[4:], 3)
                current_paragraph = None
            elif line.startswith('- ') or line.startswith('* '):
                # Bullet point
                p = doc.add_paragraph(style='List Bullet')
                add_formatted_content(p, line[2:])
                current_paragraph = None
            elif re.match(r'^\d+\.', line):
                # Numbered list
                p = doc.add_paragraph(style='List Number')
                add_formatted_content(p, re.sub(r'^\d+\.\s*', '', line))
                current_paragraph = None
            else:
                # Regular paragraph
                if current_paragraph is None:
                    current_paragraph = doc.add_paragraph()
                else:
                    current_paragraph.add_run('\n')
                add_formatted_content(current_paragraph, line)
        
        # Save to BytesIO
        doc_stream = BytesIO()
        doc.save(doc_stream)
        doc_stream.seek(0)
        
        # Create safe filename
        safe_title = re.sub(r'[<>:"/\\|?*]', '', recording.title or 'Untitled')
        safe_title = re.sub(r'[-\s]+', '-', safe_title).strip('-')
        filename = f'notes-{safe_title}.docx' if safe_title else f'notes-recording-{recording_id}.docx'

        # Create ASCII fallback for send_file - if title has non-ASCII chars, use generic name with ID
        ascii_filename = filename.encode('ascii', 'ignore').decode('ascii')
        if not ascii_filename.strip() or ascii_filename.strip() in ['notes-.docx', 'notes-recording-.docx']:
            ascii_filename = f'notes-recording-{recording_id}.docx'

        response = send_file(
            doc_stream,
            as_attachment=False,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        # Properly encode filename for international characters
        # Check if filename contains non-ASCII characters
        try:
            # Try to encode as ASCII - if this works, use simple format
            filename.encode('ascii')
            # ASCII-only filename, use simple format
            response.headers['Content-Disposition'] = f'attachment; filename="{filename}"'
        except UnicodeEncodeError:
            # Contains non-ASCII characters, use proper RFC 2231 encoding
            try:
                # Use Python's built-in RFC 2231 encoder
                encoded_value = encode_rfc2231(filename, charset='utf-8')
                header_value = f'attachment; filename*={encoded_value}'
                app.logger.error(f"DEBUG CHINESE FILENAME (RFC2231): Original='{filename}', Header='{header_value}'")
                response.headers['Content-Disposition'] = header_value
            except Exception as e:
                # Fallback to simple attachment with generic name
                app.logger.error(f"RFC2231 encoding failed: {e}, using fallback")
                response.headers['Content-Disposition'] = f'attachment; filename="download-{recording_id}.docx"'
        return response
        
    except Exception as e:
        app.logger.error(f"Error generating notes Word document: {e}")
        return jsonify({'error': 'Failed to generate Word document'}), 500

@app.route('/recording/<int:recording_id>/generate_summary', methods=['POST'])
@login_required
def generate_summary_endpoint(recording_id):
    """Generate summary for a recording that doesn't have one."""
    try:
        recording = db.session.get(Recording, recording_id)
        if not recording:
            return jsonify({'error': 'Recording not found'}), 404
            
        # Check if the recording belongs to the current user
        if recording.user_id and recording.user_id != current_user.id:
            return jsonify({'error': 'You do not have permission to generate summary for this recording'}), 403
            
        # Check if transcription exists
        if not recording.transcription or len(recording.transcription.strip()) < 10:
            return jsonify({'error': 'No valid transcription available for summary generation'}), 400
            
        # Check if already processing
        if recording.status in ['PROCESSING', 'SUMMARIZING']:
            return jsonify({'error': 'Recording is already being processed'}), 400
            
        # Check if OpenRouter client is available
        if client is None:
            return jsonify({'error': 'Summary service is not available (OpenRouter client not configured)'}), 503
            
        app.logger.info(f"Starting summary generation for recording {recording_id}")
        
        # Generate summary in background thread
        thread = threading.Thread(
            target=generate_summary_only_task,
            args=(app.app_context(), recording_id)
        )
        thread.start()
        
        return jsonify({
            'success': True, 
            'message': 'Summary generation started'
        })
        
    except Exception as e:
        db.session.rollback()
        app.logger.error(f"Error starting summary generation for recording {recording_id}: {e}")
        return jsonify({'error': str(e)}), 500

def update_speaker_usage(speaker_names):
    """Helper function to update speaker usage statistics."""
    if not speaker_names or not current_user.is_authenticated:
        return
    
    try:
        for name in speaker_names:
            name = name.strip()
            if not name:
                continue
                
            speaker = Speaker.query.filter_by(user_id=current_user.id, name=name).first()
            if speaker:
                speaker.use_count += 1
                speaker.last_used = datetime.utcnow()
            else:
                # Create new speaker
                speaker = Speaker(
                    name=name,
                    user_id=current_user.id,
                    use_count=1,
                    created_at=datetime.utcnow(),
                    last_used=datetime.utcnow()
                )
                db.session.add(speaker)
        
        db.session.commit()
    except Exception as e:
        app.logger.error(f"Error updating speaker usage: {e}")
        db.session.rollback()

@app.route('/recording/<int:recording_id>/update_speakers', methods=['POST'])
@login_required
def update_speakers(recording_id):
    """Updates speaker labels in a transcription with provided names."""
    try:
        recording = db.session.get(Recording, recording_id)
        if not recording:
            return jsonify({'error': 'Recording not found'}), 404

        if recording.user_id and recording.user_id != current_user.id:
            return jsonify({'error': 'You do not have permission to edit this recording'}), 403

        data = request.json
        speaker_map = data.get('speaker_map')
        regenerate_summary = data.get('regenerate_summary', False)

        if not speaker_map:
            return jsonify({'error': 'No speaker map provided'}), 400

        transcription_text = recording.transcription
        is_json = False
        try:
            transcription_data = json.loads(transcription_text)
            # Updated check for our new simplified JSON format (a list of segment objects)
            is_json = isinstance(transcription_data, list)
        except (json.JSONDecodeError, TypeError):
            is_json = False

        speaker_names_used = []

        if is_json:
            # Handle new simplified JSON transcript (list of segments)
            for segment in transcription_data:
                original_speaker_label = segment.get('speaker')
                if original_speaker_label in speaker_map:
                    new_name_info = speaker_map[original_speaker_label]
                    new_name = new_name_info.get('name', '').strip()
                    if new_name_info.get('isMe'):
                        new_name = current_user.name or 'Me'
                    
                    if new_name:
                        segment['speaker'] = new_name
                        if new_name not in speaker_names_used:
                            speaker_names_used.append(new_name)
            
            recording.transcription = json.dumps(transcription_data)
            
            # Update participants only from speakers that were actually given names (not default labels)
            final_speakers = set()
            for seg in transcription_data:
                speaker = seg.get('speaker')
                if speaker and str(speaker).strip():
                    # Only include speakers that have been given actual names (not default labels like "SPEAKER_01", "SPEAKER_09", etc.)
                    # Check if this speaker was updated with a real name (not a default label)
                    if not re.match(r'^SPEAKER_\d+$', str(speaker), re.IGNORECASE):
                        final_speakers.add(speaker)
            recording.participants = ', '.join(sorted(list(final_speakers)))

        else:
            # Handle plain text transcript
            new_participants = []
            for speaker_label, new_name_info in speaker_map.items():
                new_name = new_name_info.get('name', '').strip()
                if new_name_info.get('isMe'):
                    new_name = current_user.name or 'Me'
                
                if new_name:
                    transcription_text = re.sub(r'\[\s*' + re.escape(speaker_label) + r'\s*\]', f'[{new_name}]', transcription_text, flags=re.IGNORECASE)
                    if new_name not in new_participants:
                        new_participants.append(new_name)
            
            recording.transcription = transcription_text
            if new_participants:
                recording.participants = ', '.join(new_participants)
            speaker_names_used = new_participants

        # Update speaker usage statistics
        if speaker_names_used:
            update_speaker_usage(speaker_names_used)
        
        db.session.commit()

        if regenerate_summary:
            app.logger.info(f"Regenerating summary for recording {recording_id} after speaker update.")
            thread = threading.Thread(
                target=generate_summary_only_task,
                args=(app.app_context(), recording.id)
            )
            thread.start()
        
        return jsonify({
            'success': True, 
            'message': 'Speakers updated successfully.',
            'recording': recording.to_dict()
        })

    except Exception as e:
        db.session.rollback()
        app.logger.error(f"Error updating speakers for recording {recording_id}: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500

def identify_speakers_from_text(transcription):
    """
    Uses an LLM to identify speakers from a transcription.
    """
    if not TEXT_MODEL_API_KEY:
        raise ValueError("TEXT_MODEL_API_KEY not configured.")

    # The transcription passed here could be JSON, so we format it.
    formatted_transcription = format_transcription_for_llm(transcription)

    # Extract existing speaker labels (e.g., SPEAKER_00, SPEAKER_01) in order of appearance
    all_labels = re.findall(r'\[(SPEAKER_\d+)\]', formatted_transcription)
    seen = set()
    speaker_labels = [x for x in all_labels if not (x in seen or seen.add(x))]
    
    if not speaker_labels:
        return {}

    # Get configurable transcript length limit
    transcript_limit = SystemSetting.get_setting('transcript_length_limit', 30000)
    if transcript_limit == -1:
        # No limit
        transcript_text = formatted_transcription
    else:
        transcript_text = formatted_transcription[:transcript_limit]

    prompt = f"""Analyze the following transcription and identify the names of the speakers. The speakers are labeled as {', '.join(speaker_labels)}. Based on the context of the conversation, determine the most likely name for each speaker label.

Transcription:
---
{transcript_text}
---

Respond with a single JSON object where keys are the speaker labels (e.g., "SPEAKER_00") and values are the identified full names. If a name cannot be determined, use the value "Unknown".

Example:
{{
  "SPEAKER_00": "John Doe",
  "SPEAKER_01": "Jane Smith",
  "SPEAKER_02": "Unknown"
}}

JSON Response:
"""

    try:
        completion = call_llm_completion(
            messages=[
                {"role": "system", "content": "You are an expert in analyzing conversation transcripts to identify speakers. Your response must be a single, valid JSON object."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.2,
            response_format={"type": "json_object"}
        )
        response_content = completion.choices[0].message.content
        speaker_map = safe_json_loads(response_content, {})

        # Post-process the map to replace "Unknown" with an empty string
        for speaker_label, identified_name in speaker_map.items():
            if identified_name.strip().lower() == "unknown":
                speaker_map[speaker_label] = ""
                
        return speaker_map
    except Exception as e:
        app.logger.error(f"Error calling LLM for speaker identification: {e}")
        raise

def identify_unidentified_speakers_from_text(transcription, unidentified_speakers):
    """
    Uses an LLM to identify only the unidentified speakers from a transcription.
    """
    if not TEXT_MODEL_API_KEY:
        raise ValueError("TEXT_MODEL_API_KEY not configured.")

    # The transcription passed here could be JSON, so we format it.
    formatted_transcription = format_transcription_for_llm(transcription)

    if not unidentified_speakers:
        return {}

    # Get configurable transcript length limit
    transcript_limit = SystemSetting.get_setting('transcript_length_limit', 30000)
    if transcript_limit == -1:
        # No limit
        transcript_text = formatted_transcription
    else:
        transcript_text = formatted_transcription[:transcript_limit]

    prompt = f"""Analyze the following conversation transcript and identify the names of the UNIDENTIFIED speakers based on the context and content of their dialogue. 

The speakers that need to be identified are: {', '.join(unidentified_speakers)}

Look for clues in the conversation such as:
- Names mentioned by other speakers when addressing someone
- Self-introductions or references to their own name
- Context clues about roles, relationships, or positions
- Any direct mentions of names in the dialogue

Here is the complete conversation transcript:

{transcript_text}

Based on the conversation above, identify the most likely real names for the unidentified speakers. Pay close attention to how speakers address each other and any names that are mentioned in the dialogue.

Respond with a single JSON object where keys are the speaker labels (e.g., "SPEAKER_01") and values are the identified full names. If a name cannot be determined from the conversation context, use an empty string "".

Example format:
{{
  "SPEAKER_01": "Jane Smith",
  "SPEAKER_03": "Bob Johnson",
  "SPEAKER_05": ""
}}

JSON Response:
"""

    try:
        completion = call_llm_completion(
            messages=[
                {"role": "system", "content": "You are an expert in analyzing conversation transcripts to identify speakers based on contextual clues in the dialogue. Analyze the conversation carefully to find names mentioned when speakers address each other or introduce themselves. Your response must be a single, valid JSON object containing only the requested speaker identifications."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.2,
            response_format={"type": "json_object"}
        )
        response_content = completion.choices[0].message.content
        speaker_map = safe_json_loads(response_content, {})

        # Post-process the map to replace "Unknown" with an empty string
        for speaker_label, identified_name in speaker_map.items():
            if identified_name and identified_name.strip().lower() in ["unknown", "n/a", "not available", "unclear"]:
                speaker_map[speaker_label] = ""
                
        return speaker_map
    except Exception as e:
        app.logger.error(f"Error calling LLM for speaker identification: {e}")
        raise

@app.route('/recording/<int:recording_id>/auto_identify_speakers', methods=['POST'])
@login_required
def auto_identify_speakers(recording_id):
    """
    Automatically identifies speakers in a transcription using an LLM.
    Only identifies speakers that haven't been identified yet.
    """
    try:
        recording = db.session.get(Recording, recording_id)
        if not recording:
            return jsonify({'error': 'Recording not found'}), 404

        if recording.user_id and recording.user_id != current_user.id:
            return jsonify({'error': 'You do not have permission to modify this recording'}), 403

        if not recording.transcription:
            return jsonify({'error': 'No transcription available for speaker identification'}), 400

        # Get the current speaker map from the request
        data = request.json or {}
        current_speaker_map = data.get('current_speaker_map', {})
        
        # Extract all speaker labels from transcription
        formatted_transcription = format_transcription_for_llm(recording.transcription)
        all_labels = re.findall(r'\[(SPEAKER_\d+)\]', formatted_transcription)
        seen = set()
        speaker_labels = [x for x in all_labels if not (x in seen or seen.add(x))]
        
        # Filter out speakers that already have names assigned
        unidentified_speakers = []
        for speaker_label in speaker_labels:
            speaker_info = current_speaker_map.get(speaker_label, {})
            speaker_name = speaker_info.get('name', '').strip() if isinstance(speaker_info, dict) else str(speaker_info).strip()
            
            # Only include speakers that don't have names or have empty names
            if not speaker_name:
                unidentified_speakers.append(speaker_label)
        
        if not unidentified_speakers:
            return jsonify({'success': True, 'speaker_map': {}, 'message': 'All speakers are already identified'})

        # Call the helper function with only unidentified speakers
        speaker_map = identify_unidentified_speakers_from_text(recording.transcription, unidentified_speakers)

        return jsonify({'success': True, 'speaker_map': speaker_map})

    except ValueError as ve:
        # Handle cases where API key is not set
        return jsonify({'error': str(ve)}), 503
    except Exception as e:
        app.logger.error(f"Error during auto speaker identification for recording {recording_id}: {e}", exc_info=True)
        return jsonify({'error': f'An unexpected error occurred: {e}'}), 500

# --- Chat with Transcription ---
@app.route('/chat', methods=['POST'])
@login_required
def chat_with_transcription():
    try:
        data = request.json
        if not data:
            return jsonify({'error': 'No data provided'}), 400
        
        recording_id = data.get('recording_id')
        user_message = data.get('message')
        message_history = data.get('message_history', [])
        
        if not recording_id:
            return jsonify({'error': 'No recording ID provided'}), 400
        if not user_message:
            return jsonify({'error': 'No message provided'}), 400
            
        # Get the recording
        recording = db.session.get(Recording, recording_id)
        if not recording:
            return jsonify({'error': 'Recording not found'}), 404
            
        # Check if the recording belongs to the current user
        if recording.user_id and recording.user_id != current_user.id:
            return jsonify({'error': 'You do not have permission to chat with this recording'}), 403
            
        # Check if OpenRouter client is available
        if client is None:
            return jsonify({'error': 'Chat service is not available (OpenRouter client not configured)'}), 503
            
        # Prepare the system prompt with the transcription
        user_chat_output_language = current_user.output_language if current_user.is_authenticated else None
        
        language_instruction = ""
        if user_chat_output_language:
            language_instruction = f"Please provide all your responses in {user_chat_output_language}."
        
        user_name = current_user.name if current_user.is_authenticated and current_user.name else "User"
        user_title = current_user.job_title if current_user.is_authenticated and current_user.job_title else "a professional"
        user_company = current_user.company if current_user.is_authenticated and current_user.company else "their organization"

        formatted_transcription = format_transcription_for_llm(recording.transcription)
        
        # Get configurable transcript length limit for chat
        transcript_limit = SystemSetting.get_setting('transcript_length_limit', 30000)
        if transcript_limit == -1:
            # No limit
            chat_transcript = formatted_transcription
        else:
            chat_transcript = formatted_transcription[:transcript_limit]
        
        system_prompt = f"""You are a professional meeting and audio transcription analyst assisting {user_name}, who is a(n) {user_title} at {user_company}. {language_instruction} Analyze the following meeting information and respond to the specific request.

Following are the meeting participants and their roles:
{recording.participants or "No specific participants information provided."}

Following is the meeting transcript:
<<start transcript>>
{chat_transcript or "No transcript available."}
<<end transcript>>

Additional context and notes about the meeting:
{recording.notes or "none"}
"""
        
        # Prepare messages array with system prompt and conversation history
        messages = [{"role": "system", "content": system_prompt}]
        if message_history:
            messages.extend(message_history)
        messages.append({"role": "user", "content": user_message})

        def generate():
            try:
                # Enable streaming
                stream = call_llm_completion(
                    messages=messages,
                    temperature=0.7,
                    max_tokens=int(os.environ.get("CHAT_MAX_TOKENS", "2000")),
                    stream=True
                )
                
                # Use helper function to process streaming with thinking tag support
                for response in process_streaming_with_thinking(stream):
                    yield response

            except Exception as e:
                app.logger.error(f"Error during chat stream generation: {str(e)}")
                # Yield an error message in SSE format
                yield f"data: {json.dumps({'error': str(e)})}\n\n"

        return Response(generate(), mimetype='text/event-stream')

    except Exception as e:
        app.logger.error(f"Error in chat endpoint: {str(e)}")
        return jsonify({'error': str(e)}), 500


# --- Reprocessing Endpoints ---
@app.route('/recording/<int:recording_id>/reprocess_transcription', methods=['POST'])
@login_required
def reprocess_transcription(recording_id):
    """Reprocess transcription for a given recording."""
    try:
        recording = db.session.get(Recording, recording_id)
        if not recording:
            return jsonify({'error': 'Recording not found'}), 404

        if recording.user_id and recording.user_id != current_user.id:
            return jsonify({'error': 'You do not have permission to reprocess this recording'}), 403

        if not recording.audio_path or not os.path.exists(recording.audio_path):
            return jsonify({'error': 'Audio file not found for reprocessing'}), 404

        if recording.status in ['PROCESSING', 'SUMMARIZING']:
            return jsonify({'error': 'Recording is already being processed'}), 400

        # --- Convert file if necessary before reprocessing ---
        filepath = recording.audio_path
        filename_for_asr = recording.original_filename or os.path.basename(filepath)
        filename_lower = filename_for_asr.lower()

        supported_formats = ('.wav', '.mp3', '.flac')
        if not filename_lower.endswith(supported_formats):
            app.logger.info(f"Reprocessing: Converting {filename_lower} format to high-quality MP3.")
            base_filepath, file_ext = os.path.splitext(filepath)
            temp_mp3_filepath = f"{base_filepath}_temp.mp3"
            final_mp3_filepath = f"{base_filepath}.mp3"

            try:
                # Convert to high-quality MP3 (128kbps, 44.1kHz)
                subprocess.run(
                    ['ffmpeg', '-i', filepath, '-y', '-acodec', 'libmp3lame', '-b:a', '128k', '-ar', '44100', temp_mp3_filepath],
                    check=True, capture_output=True, text=True
                )
                app.logger.info(f"Successfully converted {filepath} to {temp_mp3_filepath} (128kbps MP3)")

                # If the original file is not the same as the final mp3 file, remove it
                if filepath.lower() != final_mp3_filepath.lower():
                    os.remove(filepath)
                
                # Rename the temporary file to the final filename
                os.rename(temp_mp3_filepath, final_mp3_filepath)
                
                filepath = final_mp3_filepath
                filename_for_asr = os.path.basename(filepath)
                
                # Update database with new path and mime type
                recording.audio_path = filepath
                recording.mime_type, _ = mimetypes.guess_type(filepath)
                db.session.commit()

            except FileNotFoundError:
                app.logger.error("ffmpeg command not found. Please ensure ffmpeg is installed and in the system's PATH.")
                return jsonify({'error': 'Audio conversion tool (ffmpeg) not found on server.'}), 500
            except subprocess.CalledProcessError as e:
                app.logger.error(f"ffmpeg conversion failed for {filepath}: {e.stderr}")
                return jsonify({'error': f'Failed to convert audio file: {e.stderr}'}), 500

        # --- Proceed with reprocessing ---
        recording.transcription = None
        recording.summary = None
        recording.status = 'PROCESSING'
        db.session.commit()

        app.logger.info(f"Starting transcription reprocessing for recording {recording_id}")

        # Decide which transcription method to use
        if USE_ASR_ENDPOINT:
            app.logger.info(f"Using ASR endpoint for reprocessing recording {recording_id}")
            
            data = request.json or {}
            language = data.get('language') or (recording.owner.transcription_language if recording.owner else None)
            min_speakers = data.get('min_speakers') or None
            max_speakers = data.get('max_speakers') or None
            
            # Convert to int if provided
            if min_speakers:
                try:
                    min_speakers = int(min_speakers)
                except (ValueError, TypeError):
                    min_speakers = None
            if max_speakers:
                try:
                    max_speakers = int(max_speakers)
                except (ValueError, TypeError):
                    max_speakers = None
            
            # Apply tag defaults if no user input provided (get merged defaults from all tags on this recording)
            if (min_speakers is None or max_speakers is None) and recording.tags:
                # Get tag defaults (use first non-None value from ordered tags)
                for tag_association in sorted(recording.recording_associations, key=lambda x: x.order):
                    tag = tag_association.tag
                    if min_speakers is None and tag.default_min_speakers:
                        min_speakers = tag.default_min_speakers
                    if max_speakers is None and tag.default_max_speakers:
                        max_speakers = tag.default_max_speakers
                    # Stop once we have both values
                    if min_speakers is not None and max_speakers is not None:
                        break
            
            # Apply environment variable defaults if still no values (reprocess hierarchy: user input > tag defaults > env vars > auto-detect)
            if min_speakers is None and ASR_MIN_SPEAKERS:
                try:
                    min_speakers = int(ASR_MIN_SPEAKERS)
                except (ValueError, TypeError):
                    min_speakers = None
            if max_speakers is None and ASR_MAX_SPEAKERS:
                try:
                    max_speakers = int(ASR_MAX_SPEAKERS)
                except (ValueError, TypeError):
                    max_speakers = None
            if 'ASR_DIARIZE' in os.environ:
                diarize_setting = ASR_DIARIZE
            elif USE_ASR_ENDPOINT:
                # When using ASR endpoint, use the configured ASR_DIARIZE value
                diarize_setting = ASR_DIARIZE
            else:
                diarize_setting = recording.owner.diarize if recording.owner else False

            start_time = datetime.utcnow()
            thread = threading.Thread(
                target=transcribe_audio_asr,
                args=(app.app_context(), recording.id, filepath, filename_for_asr, start_time, recording.mime_type, language, diarize_setting, min_speakers, max_speakers)
            )
        else:
            app.logger.info(f"Using standard transcription API for reprocessing recording {recording_id}")
            start_time = datetime.utcnow()
            thread = threading.Thread(
                target=transcribe_audio_task,
                args=(app.app_context(), recording.id, filepath, filename_for_asr, start_time)
            )
        
        thread.start()

        return jsonify({
            'success': True,
            'message': 'Transcription reprocessing started',
            'recording': recording.to_dict()
        })

    except Exception as e:
        db.session.rollback()
        app.logger.error(f"Error reprocessing transcription for recording {recording_id}: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/recording/<int:recording_id>/reprocess_summary', methods=['POST'])
@login_required
def reprocess_summary(recording_id):
    """Reprocess summary for a given recording (requires existing transcription)."""
    try:
        recording = db.session.get(Recording, recording_id)
        if not recording:
            return jsonify({'error': 'Recording not found'}), 404
            
        # Check if the recording belongs to the current user
        if recording.user_id and recording.user_id != current_user.id:
            return jsonify({'error': 'You do not have permission to reprocess this recording'}), 403
            
        # Check if transcription exists
        if not recording.transcription or len(recording.transcription.strip()) < 10:
            return jsonify({'error': 'No valid transcription available for summary generation'}), 400
            
        # Check if already processing
        if recording.status in ['PROCESSING', 'SUMMARIZING']:
            return jsonify({'error': 'Recording is already being processed'}), 400
            
        # Check if OpenRouter client is available
        if client is None:
            return jsonify({'error': 'Summary service is not available (OpenRouter client not configured)'}), 503
            
        # Set status to SUMMARIZING and clear existing summary
        recording.summary = None
        recording.status = 'SUMMARIZING'
        db.session.commit()
        
        app.logger.info(f"Starting summary reprocessing for recording {recording_id}")
        
        # Start summary generation in background thread
        def reprocess_summary_task(app_context, recording_id):
            app.logger.info(f"Starting summary reprocessing for recording {recording_id} using generate_summary_only_task")
            generate_summary_only_task(app_context, recording_id)
        
        thread = threading.Thread(
            target=reprocess_summary_task,
            args=(app.app_context(), recording.id)
        )
        thread.start()
        
        return jsonify({
            'success': True, 
            'message': 'Summary reprocessing started',
            'recording': recording.to_dict()
        })
        
    except Exception as e:
        db.session.rollback()
        app.logger.error(f"Error reprocessing summary for recording {recording_id}: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/recording/<int:recording_id>/reset_status', methods=['POST'])
@login_required
def reset_status(recording_id):
    """Resets the status of a stuck or failed recording."""
    try:
        recording = db.session.get(Recording, recording_id)
        if not recording:
            return jsonify({'error': 'Recording not found'}), 404

        if recording.user_id and recording.user_id != current_user.id:
            return jsonify({'error': 'You do not have permission to modify this recording'}), 403

        # Allow resetting if it's stuck or failed
        if recording.status in ['PROCESSING', 'SUMMARIZING', 'FAILED']:
            recording.status = 'FAILED'
            recording.error_message = "Manually reset from stuck or failed state."
            db.session.commit()
            app.logger.info(f"Manually reset status for recording {recording_id} to FAILED.")
            return jsonify({'success': True, 'message': 'Recording status has been reset.', 'recording': recording.to_dict()})
        else:
            return jsonify({'error': f'Recording is not in a state that can be reset. Current status: {recording.status}'}), 400

    except Exception as e:
        db.session.rollback()
        app.logger.error(f"Error resetting status for recording {recording_id}: {e}")
        return jsonify({'error': str(e)}), 500

# --- Authentication Routes ---
@app.route('/register', methods=['GET', 'POST'])
@limiter.limit("10 per minute")
def register():
    # Check if registration is allowed
    allow_registration = os.environ.get('ALLOW_REGISTRATION', 'true').lower() == 'true'
    
    if not allow_registration:
        flash('Registration is currently disabled. Please contact the administrator.', 'danger')
        return redirect(url_for('login'))
        
    if current_user.is_authenticated:
        return redirect(url_for('index'))
    
    form = RegistrationForm()
    if form.validate_on_submit():
        hashed_password = bcrypt.generate_password_hash(form.password.data).decode('utf-8')
        user = User(username=form.username.data, email=form.email.data, password=hashed_password)
        db.session.add(user)
        db.session.commit()
        flash('Your account has been created! You can now log in.', 'success')
        return redirect(url_for('login'))
    
    return render_template('register.html', title='Register', form=form)

def is_safe_url(target):
    ref_url = urlparse(request.host_url)
    test_url = urlparse(urljoin(request.host_url, target))
    return test_url.scheme in ('http', 'https') and ref_url.netloc == test_url.netloc

@app.route('/login', methods=['GET', 'POST'])
@limiter.limit("10 per minute")
def login():
    if current_user.is_authenticated:
        return redirect(url_for('index'))
    
    form = LoginForm()
    if form.validate_on_submit():
        user = User.query.filter_by(email=form.email.data).first()
        if user and bcrypt.check_password_hash(user.password, form.password.data):
            login_user(user, remember=form.remember.data)
            next_page = request.args.get('next')
            if not is_safe_url(next_page):
                return redirect(url_for('index'))
            return redirect(next_page) if next_page else redirect(url_for('index'))
        else:
            flash('Login unsuccessful. Please check email and password.', 'danger')
    
    return render_template('login.html', title='Login', form=form)

@app.route('/logout')
@csrf.exempt
def logout():
    logout_user()
    return redirect(url_for('login'))

@app.route('/account', methods=['GET', 'POST'])
@login_required
def account():
    if request.method == 'POST':
        # Only update fields that are present in the form submission
        # This prevents clearing data when switching between tabs
        
        # Check if this is the account information form (has user_name field)
        if 'user_name' in request.form:
            # Handle personal information updates
            user_name = request.form.get('user_name')
            user_job_title = request.form.get('user_job_title')
            user_company = request.form.get('user_company')
            ui_lang = request.form.get('ui_language')
            transcription_lang = request.form.get('transcription_language')
            output_lang = request.form.get('output_language')
            
            current_user.name = user_name if user_name else None
            current_user.job_title = user_job_title if user_job_title else None
            current_user.company = user_company if user_company else None
            current_user.ui_language = ui_lang if ui_lang else 'en'
            current_user.transcription_language = transcription_lang if transcription_lang else None
            current_user.output_language = output_lang if output_lang else None
        
        # Check if this is the custom prompts form (has summary_prompt field)
        elif 'summary_prompt' in request.form:
            # Handle custom prompt updates
            summary_prompt_text = request.form.get('summary_prompt')
            current_user.summary_prompt = summary_prompt_text if summary_prompt_text else None
            # Handle event extraction setting
            current_user.extract_events = 'extract_events' in request.form
        
        # Only update diarize if it's not locked by env var
        if 'ASR_DIARIZE' not in os.environ:
            current_user.diarize = 'diarize' in request.form
        
        db.session.commit()
        
        # Return JSON response for AJAX requests
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest' or request.accept_mimetypes.best == 'application/json':
            return jsonify({'success': True, 'message': 'Account details updated successfully!'})
        
        # Regular form submission with redirect
        flash('Account details updated successfully!', 'success')
        
        # Preserve the active tab when redirecting
        if 'summary_prompt' in request.form:
            return redirect(url_for('account') + '#prompts')
        else:
            return redirect(url_for('account'))
        
    # Get admin default prompt from system settings
    admin_default_prompt = SystemSetting.get_setting('admin_default_summary_prompt', None)
    if admin_default_prompt:
        default_summary_prompt_text = admin_default_prompt
    else:
        # Fallback to hardcoded default if admin hasn't set one
        default_summary_prompt_text = """Generate a comprehensive summary that includes the following sections:
- **Key Issues Discussed**: A bulleted list of the main topics
- **Key Decisions Made**: A bulleted list of any decisions reached
- **Action Items**: A bulleted list of tasks assigned, including who is responsible if mentioned"""
    
    asr_diarize_locked = 'ASR_DIARIZE' in os.environ
    
    return render_template('account.html', 
                           title='Account', 
                           default_summary_prompt_text=default_summary_prompt_text,
                           use_asr_endpoint=USE_ASR_ENDPOINT,
                           asr_diarize_locked=asr_diarize_locked,
                           asr_diarize_env_value=ASR_DIARIZE)

@app.route('/change_password', methods=['POST'])
@login_required
@limiter.limit("10 per minute")
def change_password():
    current_password = request.form.get('current_password')
    new_password = request.form.get('new_password')
    confirm_password = request.form.get('confirm_password')
    
    # Validate form data
    if not current_password or not new_password or not confirm_password:
        flash('All fields are required.', 'danger')
        return redirect(url_for('account'))
    
    if new_password != confirm_password:
        flash('New password and confirmation do not match.', 'danger')
        return redirect(url_for('account'))
        
    # Custom validation for new password
    try:
        password_check(None, type('obj', (object,), {'data': new_password}))
    except ValidationError as e:
        flash(str(e), 'danger')
        return redirect(url_for('account'))
    
    # Check if current password is correct
    if not bcrypt.check_password_hash(current_user.password, current_password):
        flash('Current password is incorrect.', 'danger')
        return redirect(url_for('account'))
    
    # Update password
    current_user.password = bcrypt.generate_password_hash(new_password).decode('utf-8')
    db.session.commit()
    
    flash('Your password has been updated successfully.', 'success')
    return redirect(url_for('account'))

# --- Admin Routes ---
@app.route('/admin', methods=['GET'])
@login_required
def admin():
    # Check if user is admin
    if not current_user.is_admin:
        flash('You do not have permission to access the admin page.', 'danger')
        return redirect(url_for('index'))
    return render_template('admin.html', title='Admin Dashboard', inquire_mode_enabled=ENABLE_INQUIRE_MODE)

@app.route('/admin/users', methods=['GET'])
@login_required
def admin_get_users():
    # Check if user is admin
    if not current_user.is_admin:
        return jsonify({'error': 'Unauthorized'}), 403
    
    users = User.query.all()
    user_data = []
    
    for user in users:
        # Get recordings count and storage used
        recordings_count = len(user.recordings)
        storage_used = sum(r.file_size for r in user.recordings if r.file_size) or 0
        
        user_data.append({
            'id': user.id,
            'username': user.username,
            'email': user.email,
            'is_admin': user.is_admin,
            'recordings_count': recordings_count,
            'storage_used': storage_used
        })
    
    return jsonify(user_data)

@app.route('/admin/users', methods=['POST'])
@login_required
def admin_add_user():
    # Check if user is admin
    if not current_user.is_admin:
        return jsonify({'error': 'Unauthorized'}), 403
    
    data = request.json
    if not data:
        return jsonify({'error': 'No data provided'}), 400
    
    # Validate required fields
    required_fields = ['username', 'email', 'password']
    for field in required_fields:
        if field not in data:
            return jsonify({'error': f'Missing required field: {field}'}), 400
    
    # Check if username or email already exists
    if User.query.filter_by(username=data['username']).first():
        return jsonify({'error': 'Username already exists'}), 400
    
    if User.query.filter_by(email=data['email']).first():
        return jsonify({'error': 'Email already exists'}), 400
    
    # Create new user
    hashed_password = bcrypt.generate_password_hash(data['password']).decode('utf-8')
    new_user = User(
        username=data['username'],
        email=data['email'],
        password=hashed_password,
        is_admin=data.get('is_admin', False)
    )
    
    db.session.add(new_user)
    db.session.commit()
    
    return jsonify({
        'id': new_user.id,
        'username': new_user.username,
        'email': new_user.email,
        'is_admin': new_user.is_admin,
        'recordings_count': 0,
        'storage_used': 0
    }), 201

@app.route('/admin/users/<int:user_id>', methods=['PUT'])
@login_required
def admin_update_user(user_id):
    # Check if user is admin
    if not current_user.is_admin:
        return jsonify({'error': 'Unauthorized'}), 403
    
    user = db.session.get(User, user_id)
    if not user:
        return jsonify({'error': 'User not found'}), 404
    
    data = request.json
    if not data:
        return jsonify({'error': 'No data provided'}), 400
    
    # Update user fields
    if 'username' in data and data['username'] != user.username:
        # Check if username already exists
        if User.query.filter_by(username=data['username']).first():
            return jsonify({'error': 'Username already exists'}), 400
        user.username = data['username']
    
    if 'email' in data and data['email'] != user.email:
        # Check if email already exists
        if User.query.filter_by(email=data['email']).first():
            return jsonify({'error': 'Email already exists'}), 400
        user.email = data['email']
    
    if 'password' in data and data['password']:
        user.password = bcrypt.generate_password_hash(data['password']).decode('utf-8')
    
    if 'is_admin' in data:
        user.is_admin = data['is_admin']
    
    db.session.commit()
    
    # Get recordings count and storage used
    recordings_count = len(user.recordings)
    storage_used = sum(r.file_size for r in user.recordings if r.file_size) or 0
    
    return jsonify({
        'id': user.id,
        'username': user.username,
        'email': user.email,
        'is_admin': user.is_admin,
        'recordings_count': recordings_count,
        'storage_used': storage_used
    })

@app.route('/admin/users/<int:user_id>', methods=['DELETE'])
@login_required
def admin_delete_user(user_id):
    # Check if user is admin
    if not current_user.is_admin:
        return jsonify({'error': 'Unauthorized'}), 403
    
    # Prevent deleting self
    if user_id == current_user.id:
        return jsonify({'error': 'Cannot delete your own account'}), 400
    
    user = db.session.get(User, user_id)
    if not user:
        return jsonify({'error': 'User not found'}), 404
    
    # Delete user's recordings and audio files
    total_chunks = 0
    if ENABLE_INQUIRE_MODE:
        total_chunks = TranscriptChunk.query.filter_by(user_id=user_id).count()
        if total_chunks > 0:
            app.logger.info(f"Deleting {total_chunks} transcript chunks with embeddings for user {user_id}")
    
    for recording in user.recordings:
        try:
            if recording.audio_path and os.path.exists(recording.audio_path):
                os.remove(recording.audio_path)
        except Exception as e:
            app.logger.error(f"Error deleting audio file {recording.audio_path}: {e}")
    
    # Delete user (cascade will handle all related data including chunks/embeddings)
    db.session.delete(user)
    db.session.commit()
    
    if ENABLE_INQUIRE_MODE and total_chunks > 0:
        app.logger.info(f"Successfully deleted {total_chunks} embeddings and chunks for user {user_id}")
    
    return jsonify({'success': True})

@app.route('/admin/users/<int:user_id>/toggle-admin', methods=['POST'])
@login_required
def admin_toggle_admin(user_id):
    # Check if user is admin
    if not current_user.is_admin:
        return jsonify({'error': 'Unauthorized'}), 403
    
    # Prevent changing own admin status
    if user_id == current_user.id:
        return jsonify({'error': 'Cannot change your own admin status'}), 400
    
    user = db.session.get(User, user_id)
    if not user:
        return jsonify({'error': 'User not found'}), 404
    
    # Toggle admin status
    user.is_admin = not user.is_admin
    db.session.commit()
    
    return jsonify({'success': True, 'is_admin': user.is_admin})

@app.route('/admin/stats', methods=['GET'])
@login_required
def admin_get_stats():
    # Check if user is admin
    if not current_user.is_admin:
        return jsonify({'error': 'Unauthorized'}), 403
    
    # Get total users
    total_users = User.query.count()
    
    # Get total recordings
    total_recordings = Recording.query.count()
    
    # Get recordings by status
    completed_recordings = Recording.query.filter_by(status='COMPLETED').count()
    processing_recordings = Recording.query.filter(Recording.status.in_(['PROCESSING', 'SUMMARIZING'])).count()
    pending_recordings = Recording.query.filter_by(status='PENDING').count()
    failed_recordings = Recording.query.filter_by(status='FAILED').count()
    
    # Get total storage used
    total_storage = db.session.query(db.func.sum(Recording.file_size)).scalar() or 0
    
    # Get top users by storage
    top_users_query = db.session.query(
        User.id,
        User.username,
        db.func.count(Recording.id).label('recordings_count'),
        db.func.sum(Recording.file_size).label('storage_used')
    ).join(Recording, User.id == Recording.user_id, isouter=True) \
     .group_by(User.id) \
     .order_by(db.func.sum(Recording.file_size).desc()) \
     .limit(5)
    
    top_users = []
    for user_id, username, recordings_count, storage_used in top_users_query:
        top_users.append({
            'id': user_id,
            'username': username,
            'recordings_count': recordings_count or 0,
            'storage_used': storage_used or 0
        })
    
    # Get total queries (chat requests)
    # This is a placeholder - you would need to track this in your database
    total_queries = 0
    
    return jsonify({
        'total_users': total_users,
        'total_recordings': total_recordings,
        'completed_recordings': completed_recordings,
        'processing_recordings': processing_recordings,
        'pending_recordings': pending_recordings,
        'failed_recordings': failed_recordings,
        'total_storage': total_storage,
        'top_users': top_users,
        'total_queries': total_queries
    })

# --- Transcript Template Routes ---
@app.route('/api/transcript-templates', methods=['GET'])
@login_required
def get_transcript_templates():
    """Get all transcript templates for the current user."""
    templates = TranscriptTemplate.query.filter_by(user_id=current_user.id).all()
    return jsonify([template.to_dict() for template in templates])

@app.route('/api/transcript-templates', methods=['POST'])
@login_required
def create_transcript_template():
    """Create a new transcript template."""
    data = request.json
    if not data or not data.get('name') or not data.get('template'):
        return jsonify({'error': 'Name and template are required'}), 400

    # If this is set as default, unset other defaults
    if data.get('is_default'):
        TranscriptTemplate.query.filter_by(
            user_id=current_user.id,
            is_default=True
        ).update({'is_default': False})

    template = TranscriptTemplate(
        user_id=current_user.id,
        name=data['name'],
        template=data['template'],
        description=data.get('description'),
        is_default=data.get('is_default', False)
    )

    db.session.add(template)
    db.session.commit()

    return jsonify(template.to_dict()), 201

@app.route('/api/transcript-templates/<int:template_id>', methods=['PUT'])
@login_required
def update_transcript_template(template_id):
    """Update an existing transcript template."""
    template = TranscriptTemplate.query.filter_by(
        id=template_id,
        user_id=current_user.id
    ).first()

    if not template:
        return jsonify({'error': 'Template not found'}), 404

    data = request.json
    if not data:
        return jsonify({'error': 'No data provided'}), 400

    # If this is set as default, unset other defaults
    if data.get('is_default'):
        TranscriptTemplate.query.filter_by(
            user_id=current_user.id,
            is_default=True
        ).update({'is_default': False})

    template.name = data.get('name', template.name)
    template.template = data.get('template', template.template)
    template.description = data.get('description', template.description)
    template.is_default = data.get('is_default', template.is_default)
    template.updated_at = datetime.utcnow()

    db.session.commit()

    return jsonify(template.to_dict())

@app.route('/api/transcript-templates/<int:template_id>', methods=['DELETE'])
@login_required
def delete_transcript_template(template_id):
    """Delete a transcript template."""
    template = TranscriptTemplate.query.filter_by(
        id=template_id,
        user_id=current_user.id
    ).first()

    if not template:
        return jsonify({'error': 'Template not found'}), 404

    db.session.delete(template)
    db.session.commit()

    return jsonify({'success': True})

@app.route('/api/transcript-templates/create-defaults', methods=['POST'])
@login_required
def create_default_templates():
    """Create default templates for the user if they don't have any."""
    existing_templates = TranscriptTemplate.query.filter_by(user_id=current_user.id).count()

    if existing_templates > 0:
        return jsonify({'message': 'User already has templates'}), 200

    # Default template 1: Simple with timestamps
    template1 = TranscriptTemplate(
        user_id=current_user.id,
        name="Simple with Timestamps",
        template="[{{start_time}} - {{end_time}}] {{speaker}}: {{text}}",
        description="Basic format with timestamps and speaker names",
        is_default=True
    )

    # Default template 2: Screenplay format
    template2 = TranscriptTemplate(
        user_id=current_user.id,
        name="Screenplay",
        template="{{speaker|upper}}\n({{start_time}})\n{{text}}\n",
        description="Screenplay-style format with speaker in caps",
        is_default=False
    )

    # Default template 3: Subtitle format
    template3 = TranscriptTemplate(
        user_id=current_user.id,
        name="SRT Subtitle",
        template="{{index}}\n{{start_time|srt}} --> {{end_time|srt}}\n{{text}}\n",
        description="SRT subtitle format for video editing",
        is_default=False
    )

    db.session.add(template1)
    db.session.add(template2)
    db.session.add(template3)
    db.session.commit()

    return jsonify({
        'success': True,
        'templates': [template1.to_dict(), template2.to_dict(), template3.to_dict()]
    }), 201

@app.route('/admin/settings', methods=['GET'])
@login_required
def admin_get_settings():
    # Check if user is admin
    if not current_user.is_admin:
        return jsonify({'error': 'Unauthorized'}), 403
    
    settings = SystemSetting.query.all()
    return jsonify([setting.to_dict() for setting in settings])

@app.route('/admin/settings', methods=['POST'])
@login_required
def admin_update_setting():
    # Check if user is admin
    if not current_user.is_admin:
        return jsonify({'error': 'Unauthorized'}), 403
    
    data = request.json
    if not data:
        return jsonify({'error': 'No data provided'}), 400
    
    key = data.get('key')
    value = data.get('value')
    description = data.get('description')
    setting_type = data.get('setting_type', 'string')
    
    if not key:
        return jsonify({'error': 'Setting key is required'}), 400
    
    # Validate setting type
    valid_types = ['string', 'integer', 'boolean', 'float']
    if setting_type not in valid_types:
        return jsonify({'error': f'Invalid setting type. Must be one of: {", ".join(valid_types)}'}), 400
    
    # Validate value based on type
    if setting_type == 'integer':
        try:
            int(value) if value is not None and value != '' else None
        except (ValueError, TypeError):
            return jsonify({'error': 'Value must be a valid integer'}), 400
    elif setting_type == 'float':
        try:
            float(value) if value is not None and value != '' else None
        except (ValueError, TypeError):
            return jsonify({'error': 'Value must be a valid number'}), 400
    elif setting_type == 'boolean':
        if value not in ['true', 'false', '1', '0', 'yes', 'no', True, False, 1, 0]:
            return jsonify({'error': 'Value must be a valid boolean (true/false, 1/0, yes/no)'}), 400
    
    try:
        setting = SystemSetting.set_setting(key, value, description, setting_type)
        return jsonify(setting.to_dict())
    except Exception as e:
        db.session.rollback()
        app.logger.error(f"Error updating setting {key}: {e}")
        return jsonify({'error': str(e)}), 500

# --- Configuration API ---
@app.route('/api/config', methods=['GET'])
def get_config():
    """Get application configuration settings for the frontend."""
    try:
        # Get configurable file size limit
        max_file_size_mb = SystemSetting.get_setting('max_file_size_mb', 250)
        
        # Get chunking configuration (supports both legacy and new formats)
        chunking_info = {}
        if ENABLE_CHUNKING and chunking_service:
            mode, limit_value = chunking_service.parse_chunk_limit()
            chunking_info = {
                'chunking_enabled': True,
                'chunking_mode': mode,  # 'size' or 'duration'
                'chunking_limit': limit_value,  # Value in MB or seconds
                'chunking_limit_display': f"{limit_value}{'MB' if mode == 'size' else 's'}"
            }
        else:
            chunking_info = {
                'chunking_enabled': False,
                'chunking_mode': 'size',
                'chunking_limit': 20,
                'chunking_limit_display': '20MB'
            }

        return jsonify({
            'max_file_size_mb': max_file_size_mb,
            'recording_disclaimer': SystemSetting.get_setting('recording_disclaimer', ''),
            'use_asr_endpoint': USE_ASR_ENDPOINT,
            **chunking_info
        })
    except Exception as e:
        app.logger.error(f"Error fetching configuration: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/csrf-token', methods=['GET'])
@csrf.exempt  # Exempt this endpoint from CSRF protection since it's providing tokens
def get_csrf_token():
    """Get a fresh CSRF token for the frontend."""
    try:
        from flask_wtf.csrf import generate_csrf
        token = generate_csrf()
        app.logger.info("Fresh CSRF token generated successfully")
        return jsonify({'csrf_token': token})
    except Exception as e:
        app.logger.error(f"Error generating CSRF token: {e}")
        return jsonify({'error': str(e)}), 500

# --- Flask Routes ---
@app.route('/')
@login_required
def index():
    # Pass the ASR config, inquire mode config, and user language preference to the template
    user_language = current_user.ui_language if current_user.is_authenticated and current_user.ui_language else 'en'
    return render_template('index.html', 
                         use_asr_endpoint=USE_ASR_ENDPOINT, 
                         inquire_mode_enabled=ENABLE_INQUIRE_MODE,
                         user_language=user_language)

@app.route('/inquire')
@login_required  
def inquire():
    # Check if inquire mode is enabled
    if not ENABLE_INQUIRE_MODE:
        flash('Inquire mode is not enabled on this server.', 'warning')
        return redirect(url_for('index'))
    # Render the inquire page with user context for theming
    return render_template('inquire.html', use_asr_endpoint=USE_ASR_ENDPOINT, current_user=current_user)

@app.route('/recordings', methods=['GET'])
def get_recordings():
    try:
        # Check if user is logged in
        if not current_user.is_authenticated:
            return jsonify([])  # Return empty array if not logged in
            
        # Filter recordings by the current user
        stmt = select(Recording).where(Recording.user_id == current_user.id).order_by(Recording.created_at.desc())
        recordings = db.session.execute(stmt).scalars().all()
        return jsonify([recording.to_dict() for recording in recordings])
    except Exception as e:
        app.logger.error(f"Error fetching recordings: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/recordings', methods=['GET'])
@login_required
@limiter.limit("1250 per hour")
def get_recordings_paginated():
    """Get recordings with pagination and server-side filtering."""
    try:
        # Parse query parameters
        page = request.args.get('page', 1, type=int)
        per_page = min(request.args.get('per_page', 25, type=int), 100)  # Cap at 100 per page
        search_query = request.args.get('q', '').strip()
        
        # Build base query
        stmt = select(Recording).where(Recording.user_id == current_user.id)
        
        # Apply search filters if provided
        if search_query:
            # Parse search query for special syntax
            import re
            
            # Extract date filters
            date_filters = re.findall(r'date:(\S+)', search_query.lower())
            date_from_filters = re.findall(r'date_from:(\S+)', search_query.lower())
            date_to_filters = re.findall(r'date_to:(\S+)', search_query.lower())
            tag_filters = re.findall(r'tag:(\S+)', search_query.lower())
            
            # Remove special syntax to get text search
            text_query = re.sub(r'date:\S+', '', search_query, flags=re.IGNORECASE)
            text_query = re.sub(r'date_from:\S+', '', text_query, flags=re.IGNORECASE)
            text_query = re.sub(r'date_to:\S+', '', text_query, flags=re.IGNORECASE)
            text_query = re.sub(r'tag:\S+', '', text_query, flags=re.IGNORECASE).strip()
            
            # Apply date filters
            for date_filter in date_filters:
                if date_filter == 'today':
                    today = datetime.now().date()
                    stmt = stmt.where(
                        db.or_(
                            db.func.date(Recording.meeting_date) == today,
                            db.and_(
                                Recording.meeting_date.is_(None),
                                db.func.date(Recording.created_at) == today
                            )
                        )
                    )
                elif date_filter == 'yesterday':
                    yesterday = datetime.now().date() - timedelta(days=1)
                    stmt = stmt.where(
                        db.or_(
                            db.func.date(Recording.meeting_date) == yesterday,
                            db.and_(
                                Recording.meeting_date.is_(None),
                                db.func.date(Recording.created_at) == yesterday
                            )
                        )
                    )
                elif date_filter == 'thisweek':
                    today = datetime.now().date()
                    start_of_week = today - timedelta(days=today.weekday())
                    stmt = stmt.where(
                        db.or_(
                            Recording.meeting_date >= start_of_week,
                            db.and_(
                                Recording.meeting_date.is_(None),
                                db.func.date(Recording.created_at) >= start_of_week
                            )
                        )
                    )
                elif date_filter == 'lastweek':
                    today = datetime.now().date()
                    end_of_last_week = today - timedelta(days=today.weekday())
                    start_of_last_week = end_of_last_week - timedelta(days=7)
                    stmt = stmt.where(
                        db.or_(
                            db.and_(
                                Recording.meeting_date >= start_of_last_week,
                                Recording.meeting_date < end_of_last_week
                            ),
                            db.and_(
                                Recording.meeting_date.is_(None),
                                db.func.date(Recording.created_at) >= start_of_last_week,
                                db.func.date(Recording.created_at) < end_of_last_week
                            )
                        )
                    )
                elif date_filter == 'thismonth':
                    today = datetime.now().date()
                    start_of_month = today.replace(day=1)
                    stmt = stmt.where(
                        db.or_(
                            Recording.meeting_date >= start_of_month,
                            db.and_(
                                Recording.meeting_date.is_(None),
                                db.func.date(Recording.created_at) >= start_of_month
                            )
                        )
                    )
                elif date_filter == 'lastmonth':
                    today = datetime.now().date()
                    first_day_this_month = today.replace(day=1)
                    last_day_last_month = first_day_this_month - timedelta(days=1)
                    first_day_last_month = last_day_last_month.replace(day=1)
                    stmt = stmt.where(
                        db.or_(
                            db.and_(
                                Recording.meeting_date >= first_day_last_month,
                                Recording.meeting_date <= last_day_last_month
                            ),
                            db.and_(
                                Recording.meeting_date.is_(None),
                                db.func.date(Recording.created_at) >= first_day_last_month,
                                db.func.date(Recording.created_at) <= last_day_last_month
                            )
                        )
                    )
                elif re.match(r'^\d{4}-\d{2}-\d{2}$', date_filter):
                    # Specific date format YYYY-MM-DD
                    target_date = datetime.strptime(date_filter, '%Y-%m-%d').date()
                    stmt = stmt.where(
                        db.or_(
                            db.func.date(Recording.meeting_date) == target_date,
                            db.and_(
                                Recording.meeting_date.is_(None),
                                db.func.date(Recording.created_at) == target_date
                            )
                        )
                    )
                elif re.match(r'^\d{4}-\d{2}$', date_filter):
                    # Month format YYYY-MM
                    year, month = map(int, date_filter.split('-'))
                    stmt = stmt.where(
                        db.or_(
                            db.and_(
                                db.extract('year', Recording.meeting_date) == year,
                                db.extract('month', Recording.meeting_date) == month
                            ),
                            db.and_(
                                Recording.meeting_date.is_(None),
                                db.extract('year', Recording.created_at) == year,
                                db.extract('month', Recording.created_at) == month
                            )
                        )
                    )
                elif re.match(r'^\d{4}$', date_filter):
                    # Year format YYYY
                    year = int(date_filter)
                    stmt = stmt.where(
                        db.or_(
                            db.extract('year', Recording.meeting_date) == year,
                            db.and_(
                                Recording.meeting_date.is_(None),
                                db.extract('year', Recording.created_at) == year
                            )
                        )
                    )
            
            # Apply date range filters
            if date_from_filters and date_from_filters[0]:
                try:
                    date_from = datetime.strptime(date_from_filters[0], '%Y-%m-%d').date()
                    stmt = stmt.where(
                        db.or_(
                            Recording.meeting_date >= date_from,
                            db.and_(
                                Recording.meeting_date.is_(None),
                                db.func.date(Recording.created_at) >= date_from
                            )
                        )
                    )
                except ValueError:
                    pass  # Invalid date format, ignore
            
            if date_to_filters and date_to_filters[0]:
                try:
                    date_to = datetime.strptime(date_to_filters[0], '%Y-%m-%d').date()
                    stmt = stmt.where(
                        db.or_(
                            Recording.meeting_date <= date_to,
                            db.and_(
                                Recording.meeting_date.is_(None),
                                db.func.date(Recording.created_at) <= date_to
                            )
                        )
                    )
                except ValueError:
                    pass  # Invalid date format, ignore
            
            # Apply tag filters
            if tag_filters:
                # Join with tags table and filter by tag names
                tag_conditions = []
                for tag_filter in tag_filters:
                    # Replace underscores back to spaces for matching
                    tag_name = tag_filter.replace('_', ' ')
                    tag_conditions.append(Tag.name.ilike(f'%{tag_name}%'))
                
                stmt = stmt.join(RecordingTag).join(Tag).where(db.or_(*tag_conditions))
            
            # Apply text search
            if text_query:
                text_conditions = [
                    Recording.title.ilike(f'%{text_query}%'),
                    Recording.participants.ilike(f'%{text_query}%'),
                    Recording.transcription.ilike(f'%{text_query}%'),
                    Recording.notes.ilike(f'%{text_query}%')
                ]
                stmt = stmt.where(db.or_(*text_conditions))
        
        # Apply ordering (most recent first based on meeting_date or created_at)
        stmt = stmt.order_by(
            db.case(
                (Recording.meeting_date.is_not(None), Recording.meeting_date),
                else_=db.func.date(Recording.created_at)
            ).desc(),
            Recording.created_at.desc()
        )
        
        # Get total count for pagination info
        count_stmt = select(db.func.count()).select_from(stmt.subquery())
        total_count = db.session.execute(count_stmt).scalar()
        
        # Apply pagination
        offset = (page - 1) * per_page
        stmt = stmt.offset(offset).limit(per_page)
        
        # Execute query
        recordings = db.session.execute(stmt).scalars().all()
        
        # Calculate pagination metadata
        total_pages = (total_count + per_page - 1) // per_page
        has_next = page < total_pages
        has_prev = page > 1
        
        return jsonify({
            'recordings': [recording.to_dict() for recording in recordings],
            'pagination': {
                'page': page,
                'per_page': per_page,
                'total': total_count,
                'total_pages': total_pages,
                'has_next': has_next,
                'has_prev': has_prev
            }
        })
        
    except Exception as e:
        app.logger.error(f"Error fetching paginated recordings: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/inbox_recordings', methods=['GET'])
@login_required
@limiter.limit("1250 per hour")  # Allow frequent polling for inbox recordings
def get_inbox_recordings():
    """Get recordings that are in the inbox and currently processing."""
    try:
        stmt = select(Recording).where(
            Recording.user_id == current_user.id,
            Recording.is_inbox == True,
            Recording.status.in_(['PENDING', 'PROCESSING', 'SUMMARIZING'])
        ).order_by(Recording.created_at.desc())
        
        recordings = db.session.execute(stmt).scalars().all()
        return jsonify([recording.to_dict() for recording in recordings])
    except Exception as e:
        app.logger.error(f"Error fetching inbox recordings: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/save', methods=['POST'])
@login_required
def save_metadata():
    try:
        data = request.json
        if not data: return jsonify({'error': 'No data provided'}), 400
        recording_id = data.get('id')
        if not recording_id: return jsonify({'error': 'No recording ID provided'}), 400

        recording = db.session.get(Recording, recording_id)
        if not recording: return jsonify({'error': 'Recording not found'}), 404
        
        # Check if the recording belongs to the current user
        if recording.user_id and recording.user_id != current_user.id:
            return jsonify({'error': 'You do not have permission to edit this recording'}), 403

        # Update fields if provided (with sanitization for notes and summary)
        if 'title' in data: recording.title = data['title']
        if 'participants' in data: recording.participants = data['participants']
        if 'notes' in data: recording.notes = sanitize_html(data['notes']) if data['notes'] else data['notes']
        if 'summary' in data: recording.summary = sanitize_html(data['summary']) if data['summary'] else data['summary']
        if 'is_inbox' in data: recording.is_inbox = data['is_inbox']
        if 'is_highlighted' in data: recording.is_highlighted = data['is_highlighted']
        if 'meeting_date' in data:
            try:
                # Attempt to parse date string (e.g., "YYYY-MM-DD")
                date_str = data['meeting_date']
                if date_str:
                    recording.meeting_date = datetime.strptime(date_str, '%Y-%m-%d').date()
                else:
                    recording.meeting_date = None # Allow clearing the date
            except (ValueError, TypeError) as e:
                app.logger.warning(f"Could not parse meeting_date '{data.get('meeting_date')}': {e}")
                # Optionally return an error or just ignore the invalid date
                # return jsonify({'error': f"Invalid date format for meeting_date. Use YYYY-MM-DD."}), 400

        # Do not update transcription or status here
        db.session.commit()
        return jsonify({'success': True, 'recording': recording.to_dict()})

    except Exception as e:
        db.session.rollback()
        app.logger.error(f"Error saving metadata for recording {recording_id}: {e}", exc_info=True)
        return jsonify({'error': 'An unexpected error occurred while saving.'}), 500

@app.route('/recording/<int:recording_id>/update_transcription', methods=['POST'])
@login_required
def update_transcription(recording_id):
    """Updates the transcription content for a recording."""
    try:
        recording = db.session.get(Recording, recording_id)
        if not recording:
            return jsonify({'error': 'Recording not found'}), 404

        if recording.user_id and recording.user_id != current_user.id:
            return jsonify({'error': 'You do not have permission to edit this recording'}), 403

        data = request.json
        new_transcription = data.get('transcription')

        if new_transcription is None:
            return jsonify({'error': 'No transcription data provided'}), 400

        # The incoming data could be a JSON string (from ASR edit) or plain text
        recording.transcription = new_transcription
        
        # Optional: If the transcription changes, we might want to indicate that the summary is outdated.
        # For now, we'll just save the transcript. A "regenerate summary" button could be a good follow-up.

        db.session.commit()
        app.logger.info(f"Transcription for recording {recording_id} was updated.")
        
        return jsonify({'success': True, 'message': 'Transcription updated successfully.', 'recording': recording.to_dict()})

    except Exception as e:
        db.session.rollback()
        app.logger.error(f"Error updating transcription for recording {recording_id}: {e}", exc_info=True)
        return jsonify({'error': 'An unexpected error occurred while updating the transcription.'}), 500

# Toggle inbox status endpoint
@app.route('/recording/<int:recording_id>/toggle_inbox', methods=['POST'])
@login_required
def toggle_inbox(recording_id):
    try:
        recording = db.session.get(Recording, recording_id)
        if not recording:
            return jsonify({'error': 'Recording not found'}), 404
            
        # Check if the recording belongs to the current user
        if recording.user_id and recording.user_id != current_user.id:
            return jsonify({'error': 'You do not have permission to modify this recording'}), 403

        # Toggle the inbox status
        recording.is_inbox = not recording.is_inbox
        db.session.commit()
        
        return jsonify({'success': True, 'is_inbox': recording.is_inbox})
    except Exception as e:
        db.session.rollback()
        app.logger.error(f"Error toggling inbox status for recording {recording_id}: {e}", exc_info=True)
        return jsonify({'error': 'An unexpected error occurred.'}), 500

# Toggle highlighted status endpoint
@app.route('/recording/<int:recording_id>/toggle_highlight', methods=['POST'])
@login_required
def toggle_highlight(recording_id):
    try:
        recording = db.session.get(Recording, recording_id)
        if not recording:
            return jsonify({'error': 'Recording not found'}), 404
            
        # Check if the recording belongs to the current user
        if recording.user_id and recording.user_id != current_user.id:
            return jsonify({'error': 'You do not have permission to modify this recording'}), 403

        # Toggle the highlighted status
        recording.is_highlighted = not recording.is_highlighted
        db.session.commit()
        
        return jsonify({'success': True, 'is_highlighted': recording.is_highlighted})
    except Exception as e:
        db.session.rollback()
        app.logger.error(f"Error toggling highlighted status for recording {recording_id}: {e}", exc_info=True)
        return jsonify({'error': 'An unexpected error occurred.'}), 500


@app.route('/upload', methods=['POST'])
@login_required
def upload_file():
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400

        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400

        original_filename = file.filename
        safe_filename = secure_filename(original_filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], f"{datetime.now().strftime('%Y%m%d%H%M%S')}_{safe_filename}")

        # Get original file size
        file.seek(0, os.SEEK_END)
        original_file_size = file.tell()
        file.seek(0)

        # Check size limit before saving - only enforce if chunking is disabled or using ASR endpoint
        max_content_length = app.config.get('MAX_CONTENT_LENGTH')
        
        # Skip size check if chunking is enabled and using OpenAI Whisper API
        should_enforce_size_limit = True
        if ENABLE_CHUNKING and chunking_service and not USE_ASR_ENDPOINT:
            should_enforce_size_limit = False
            # Get chunking mode for better logging
            mode, limit_value = chunking_service.parse_chunk_limit()
            if mode == 'size':
                app.logger.info(f"Size-based chunking enabled ({limit_value}MB limit) - skipping {original_file_size/1024/1024:.1f}MB size limit check")
            else:
                app.logger.info(f"Duration-based chunking enabled ({limit_value}s limit) - skipping {original_file_size/1024/1024:.1f}MB size limit check")
        
        if should_enforce_size_limit and max_content_length and original_file_size > max_content_length:
            raise RequestEntityTooLarge()

        file.save(filepath)
        app.logger.info(f"File saved to {filepath}")

        # --- Convert files only when chunking is needed ---
        filename_lower = original_filename.lower()
        
        # Check if chunking will be needed for this file
        needs_chunking_for_processing = (chunking_service and 
                                       ENABLE_CHUNKING and 
                                       not USE_ASR_ENDPOINT and
                                       chunking_service.needs_chunking(filepath, USE_ASR_ENDPOINT))
        
        # Define supported formats based on whether chunking is needed
        if needs_chunking_for_processing:
            # For chunking: only support formats that work well with chunking
            supported_formats = ('.wav', '.mp3', '.flac')
            convertible_formats = ('.amr', '.3gp', '.3gpp', '.m4a', '.aac', '.ogg', '.wma', '.webm')
        else:
            # For direct transcription: support WebM and other formats directly
            supported_formats = ('.wav', '.mp3', '.flac', '.webm', '.m4a', '.aac', '.ogg')
            convertible_formats = ('.amr', '.3gp', '.3gpp', '.wma')
        
        # Special handling for problematic AAC files when using ASR endpoint
        is_problematic_aac = (USE_ASR_ENDPOINT and 
                             (filename_lower.endswith('.aac') or 
                              'aac' in filename_lower.lower()))
        
        # Convert if file is not in supported formats OR is problematic AAC for ASR
        should_convert = ((not filename_lower.endswith(supported_formats) and needs_chunking_for_processing) or 
                         is_problematic_aac)
        
        if should_convert:
            if is_problematic_aac:
                app.logger.info(f"Converting AAC-encoded file {filename_lower} to high-quality MP3 for ASR endpoint compatibility.")
            elif filename_lower.endswith(convertible_formats):
                app.logger.info(f"Converting {filename_lower} format to high-quality MP3 for chunking processing.")
            else:
                app.logger.info(f"Attempting to convert unknown format ({filename_lower}) to high-quality MP3 for chunking.")
            
            base_filepath, _ = os.path.splitext(filepath)
            temp_mp3_filepath = f"{base_filepath}_temp.mp3"
            mp3_filepath = f"{base_filepath}.mp3"

            try:
                # Convert to high-quality MP3 (128kbps, 44.1kHz) for better transcription accuracy
                subprocess.run(
                    ['ffmpeg', '-i', filepath, '-y', '-acodec', 'libmp3lame', '-b:a', '128k', '-ar', '44100', '-ac', '1', temp_mp3_filepath],
                    check=True, capture_output=True, text=True
                )
                app.logger.info(f"Successfully converted {filepath} to {temp_mp3_filepath} (128kbps MP3)")
                
                # If the original file is not the same as the final mp3 file, remove it
                if filepath.lower() != mp3_filepath.lower():
                    os.remove(filepath)
                
                # Rename the temporary file to the final filename
                os.rename(temp_mp3_filepath, mp3_filepath)
                
                filepath = mp3_filepath
            except FileNotFoundError:
                app.logger.error("ffmpeg command not found. Please ensure ffmpeg is installed and in the system's PATH.")
                return jsonify({'error': 'Audio conversion tool (ffmpeg) not found on server.'}), 500
            except subprocess.CalledProcessError as e:
                app.logger.error(f"ffmpeg conversion failed for {filepath}: {e.stderr}")
                return jsonify({'error': f'Failed to convert audio file: {e.stderr}'}), 500
        elif not filename_lower.endswith(supported_formats):
            # File is not supported and chunking is not needed - log but don't convert
            app.logger.info(f"File format {filename_lower} will be processed directly without conversion (chunking not needed)")

        # Get final file size (of original or converted file)
        final_file_size = os.path.getsize(filepath)

        # Determine MIME type of the final file
        mime_type, _ = mimetypes.guess_type(filepath)
        app.logger.info(f"Final MIME type: {mime_type} for file {filepath}")

        # Get notes from the form
        notes = request.form.get('notes')
        
        # Get selected tags if provided (multiple tags support)
        selected_tags = []
        tag_index = 0
        while True:
            tag_id_key = f'tag_ids[{tag_index}]'
            tag_id = request.form.get(tag_id_key)
            if not tag_id:
                break
            
            tag = Tag.query.filter_by(id=tag_id, user_id=current_user.id).first()
            if tag:
                selected_tags.append(tag)
            tag_index += 1
        
        # For backward compatibility with single tag uploads
        if not selected_tags:
            single_tag_id = request.form.get('tag_id')
            if single_tag_id:
                tag = Tag.query.filter_by(id=single_tag_id, user_id=current_user.id).first()
                if tag:
                    selected_tags.append(tag)
        
        # Get ASR advanced options if provided
        language = request.form.get('language', '')
        min_speakers = request.form.get('min_speakers') or None
        max_speakers = request.form.get('max_speakers') or None
        
        # Convert to int if provided
        if min_speakers:
            try:
                min_speakers = int(min_speakers)
            except (ValueError, TypeError):
                min_speakers = None
        if max_speakers:
            try:
                max_speakers = int(max_speakers)
            except (ValueError, TypeError):
                max_speakers = None
        
        # Apply precedence hierarchy: user input > tag defaults > environment variables > auto-detect
        
        # Apply tag defaults if tags are selected and values are not explicitly provided by user
        # Use first tag's defaults (highest priority)
        if selected_tags:
            first_tag = selected_tags[0]
            if not language and first_tag.default_language:
                language = first_tag.default_language
            if min_speakers is None and first_tag.default_min_speakers:
                min_speakers = first_tag.default_min_speakers
            if max_speakers is None and first_tag.default_max_speakers:
                max_speakers = first_tag.default_max_speakers
        
        # Apply environment variable defaults if still no values are set
        if min_speakers is None and ASR_MIN_SPEAKERS:
            try:
                min_speakers = int(ASR_MIN_SPEAKERS)
            except (ValueError, TypeError):
                min_speakers = None
        if max_speakers is None and ASR_MAX_SPEAKERS:
            try:
                max_speakers = int(ASR_MAX_SPEAKERS)
            except (ValueError, TypeError):
                max_speakers = None

        # Create initial database entry
        now = datetime.utcnow()
        recording = Recording(
            audio_path=filepath,
            original_filename=original_filename,
            title=f"Recording - {original_filename}",
            file_size=final_file_size,
            status='PENDING',
            meeting_date=now.date(),
            user_id=current_user.id,
            mime_type=mime_type,
            notes=notes,
            processing_source='upload'  # Track that this was manually uploaded
        )
        db.session.add(recording)
        db.session.commit()
        
        # Add tags to recording if selected (preserve order)
        for order, tag in enumerate(selected_tags, 1):
            new_association = RecordingTag(
                recording_id=recording.id,
                tag_id=tag.id,
                order=order,
                added_at=datetime.utcnow()
            )
            db.session.add(new_association)
        
        if selected_tags:
            db.session.commit()
            tag_names = [tag.name for tag in selected_tags]
            app.logger.info(f"Added {len(selected_tags)} tags to recording {recording.id}: {', '.join(tag_names)}")
        
        app.logger.info(f"Initial recording record created with ID: {recording.id}")

        # --- Start transcription & summarization in background thread ---
        start_time = datetime.utcnow()
        
        # Pass ASR parameters and first tag to the transcription task (for compatibility with existing functions)
        first_tag = selected_tags[0] if selected_tags else None
        if USE_ASR_ENDPOINT:
            app.logger.info(f"Starting ASR transcription thread for recording {recording.id} with params: language={language}, min_speakers={min_speakers}, max_speakers={max_speakers}, tag_id={first_tag.id if first_tag else None}")
            thread = threading.Thread(
                target=transcribe_audio_task,
                args=(app.app_context(), recording.id, filepath, os.path.basename(filepath), start_time),
                kwargs={'language': language, 'min_speakers': min_speakers, 'max_speakers': max_speakers, 'tag_id': first_tag.id if first_tag else None}
            )
        else:
            app.logger.info(f"Starting Whisper transcription thread for recording {recording.id} with tag_id={first_tag.id if first_tag else None}")
            thread = threading.Thread(
                target=transcribe_audio_task,
                args=(app.app_context(), recording.id, filepath, os.path.basename(filepath), start_time),
                kwargs={'tag_id': first_tag.id if first_tag else None}
            )
        thread.start()
        app.logger.info(f"Background processing thread started for recording ID: {recording.id}")

        return jsonify(recording.to_dict()), 202

    except RequestEntityTooLarge:
        max_size_mb = app.config['MAX_CONTENT_LENGTH'] / (1024 * 1024)
        app.logger.warning(f"Upload failed: File too large (>{max_size_mb}MB)")
        return jsonify({
            'error': f'File too large. Maximum size is {max_size_mb:.0f} MB.',
            'max_size_mb': max_size_mb
        }), 413
    except Exception as e:
        db.session.rollback()
        app.logger.error(f"Error during file upload: {e}", exc_info=True)
        return jsonify({'error': 'An unexpected error occurred during upload.'}), 500


# Status Endpoint
@app.route('/status/<int:recording_id>', methods=['GET'])
@login_required
@limiter.limit("1250 per hour")  # Allow frequent polling for status checks
def get_status(recording_id):
    """Endpoint to check the transcription/summarization status."""
    try:
        recording = db.session.get(Recording, recording_id)
        if not recording:
            return jsonify({'error': 'Recording not found'}), 404
        
        # Check if the recording belongs to the current user
        if recording.user_id and recording.user_id != current_user.id:
            return jsonify({'error': 'You do not have permission to view this recording'}), 403
            
        return jsonify(recording.to_dict())
    except Exception as e:
        app.logger.error(f"Error fetching status for recording {recording_id}: {e}", exc_info=True)
        return jsonify({'error': 'An unexpected error occurred.'}), 500

# Get Audio Endpoint
@app.route('/audio/<int:recording_id>')
@login_required
def get_audio(recording_id):
    try:
        recording = db.session.get(Recording, recording_id)
        if not recording or not recording.audio_path:
            return jsonify({'error': 'Recording or audio file not found'}), 404
            
        # Check if the recording belongs to the current user
        if recording.user_id and recording.user_id != current_user.id:
            return jsonify({'error': 'You do not have permission to access this audio file'}), 403
        if not os.path.exists(recording.audio_path):
            app.logger.error(f"Audio file missing from server: {recording.audio_path}")
            return jsonify({'error': 'Audio file missing from server'}), 404
        return send_file(recording.audio_path)
    except Exception as e:
        app.logger.error(f"Error serving audio for recording {recording_id}: {e}", exc_info=True)
        return jsonify({'error': 'An unexpected error occurred.'}), 500

@app.route('/share/audio/<string:public_id>')
def get_shared_audio(public_id):
    try:
        share = Share.query.filter_by(public_id=public_id).first_or_404()
        recording = share.recording
        if not recording or not recording.audio_path:
            return jsonify({'error': 'Recording or audio file not found'}), 404
        if not os.path.exists(recording.audio_path):
            app.logger.error(f"Audio file missing from server: {recording.audio_path}")
            return jsonify({'error': 'Audio file missing from server'}), 404
        return send_file(recording.audio_path)
    except Exception as e:
        app.logger.error(f"Error serving shared audio for public_id {public_id}: {e}", exc_info=True)
        return jsonify({'error': 'An unexpected error occurred.'}), 500

# Delete Recording Endpoint
@app.route('/recording/<int:recording_id>', methods=['DELETE'])
@login_required
def delete_recording(recording_id):
    try:
        recording = db.session.get(Recording, recording_id)
        if not recording:
            return jsonify({'error': 'Recording not found'}), 404
            
        # Check if the recording belongs to the current user
        if recording.user_id and recording.user_id != current_user.id:
            return jsonify({'error': 'You do not have permission to delete this recording'}), 403

        # Delete the audio file first
        try:
            if recording.audio_path and os.path.exists(recording.audio_path):
                os.remove(recording.audio_path)
                app.logger.info(f"Deleted audio file: {recording.audio_path}")
        except Exception as e:
            app.logger.error(f"Error deleting audio file {recording.audio_path}: {e}")

        # Log embeddings cleanup for Inquire Mode if enabled
        if ENABLE_INQUIRE_MODE:
            chunk_count = TranscriptChunk.query.filter_by(recording_id=recording_id).count()
            if chunk_count > 0:
                app.logger.info(f"Deleting {chunk_count} transcript chunks with embeddings for recording {recording_id}")

        # Delete the database record (cascade will handle chunks/embeddings)
        db.session.delete(recording)
        db.session.commit()
        app.logger.info(f"Deleted recording record ID: {recording_id}")
        
        if ENABLE_INQUIRE_MODE and chunk_count > 0:
            app.logger.info(f"Successfully deleted embeddings and chunks for recording {recording_id}")

        return jsonify({'success': True})
    except Exception as e:
        db.session.rollback()
        app.logger.error(f"Error deleting recording {recording_id}: {e}", exc_info=True)
        return jsonify({'error': 'An unexpected error occurred while deleting.'}), 500


# --- Inquire Mode API Endpoints ---

@app.route('/api/inquire/sessions', methods=['GET'])
@login_required
def get_inquire_sessions():
    """Get all inquire sessions for the current user."""
    if not ENABLE_INQUIRE_MODE:
        return jsonify({'error': 'Inquire mode is not enabled'}), 403
    try:
        sessions = InquireSession.query.filter_by(user_id=current_user.id).order_by(InquireSession.last_used.desc()).all()
        return jsonify([session.to_dict() for session in sessions])
    except Exception as e:
        app.logger.error(f"Error getting inquire sessions: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/inquire/sessions', methods=['POST'])
@login_required
def create_inquire_session():
    """Create a new inquire session with filters."""
    if not ENABLE_INQUIRE_MODE:
        return jsonify({'error': 'Inquire mode is not enabled'}), 403
    try:
        data = request.json
        if not data:
            return jsonify({'error': 'No data provided'}), 400
        
        session = InquireSession(
            user_id=current_user.id,
            session_name=data.get('session_name'),
            filter_tags=json.dumps(data.get('filter_tags', [])),
            filter_speakers=json.dumps(data.get('filter_speakers', [])),
            filter_date_from=datetime.fromisoformat(data['filter_date_from']).date() if data.get('filter_date_from') else None,
            filter_date_to=datetime.fromisoformat(data['filter_date_to']).date() if data.get('filter_date_to') else None,
            filter_recording_ids=json.dumps(data.get('filter_recording_ids', []))
        )
        
        db.session.add(session)
        db.session.commit()
        
        return jsonify(session.to_dict()), 201
        
    except Exception as e:
        app.logger.error(f"Error creating inquire session: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/inquire/search', methods=['POST'])
@login_required
def inquire_search():
    """Perform semantic search within filtered transcriptions."""
    if not ENABLE_INQUIRE_MODE:
        return jsonify({'error': 'Inquire mode is not enabled'}), 403
    try:
        data = request.json
        if not data:
            return jsonify({'error': 'No data provided'}), 400
        
        query = data.get('query')
        if not query:
            return jsonify({'error': 'No query provided'}), 400
        
        # Build filters from request
        filters = {}
        if data.get('filter_tags'):
            filters['tag_ids'] = data['filter_tags']
        if data.get('filter_speakers'):
            filters['speaker_names'] = data['filter_speakers']
        if data.get('filter_recording_ids'):
            filters['recording_ids'] = data['filter_recording_ids']
        if data.get('filter_date_from'):
            filters['date_from'] = datetime.fromisoformat(data['filter_date_from']).date()
        if data.get('filter_date_to'):
            filters['date_to'] = datetime.fromisoformat(data['filter_date_to']).date()
        
        # Perform semantic search
        top_k = data.get('top_k', 5)
        chunk_results = semantic_search_chunks(current_user.id, query, filters, top_k)
        
        # Format results
        results = []
        for chunk, similarity in chunk_results:
            result = chunk.to_dict()
            result['similarity'] = similarity
            result['recording_title'] = chunk.recording.title
            result['recording_meeting_date'] = f"{chunk.recording.meeting_date.isoformat()}T00:00:00" if chunk.recording.meeting_date else None
            results.append(result)
        
        return jsonify({'results': results})
        
    except Exception as e:
        app.logger.error(f"Error in inquire search: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/inquire/chat', methods=['POST'])
@login_required
def inquire_chat():
    """Chat with filtered transcriptions using RAG."""
    if not ENABLE_INQUIRE_MODE:
        return jsonify({'error': 'Inquire mode is not enabled'}), 403
    try:
        data = request.json
        if not data:
            return jsonify({'error': 'No data provided'}), 400
        
        user_message = data.get('message')
        message_history = data.get('message_history', [])
        
        if not user_message:
            return jsonify({'error': 'No message provided'}), 400
        
        # Check if OpenRouter client is available
        if client is None:
            return jsonify({'error': 'Chat service is not available (OpenRouter client not configured)'}), 503
        
        # Build filters from request
        filters = {}
        if data.get('filter_tags'):
            filters['tag_ids'] = data['filter_tags']
        if data.get('filter_speakers'):
            filters['speaker_names'] = data['filter_speakers']
        if data.get('filter_recording_ids'):
            filters['recording_ids'] = data['filter_recording_ids']
        if data.get('filter_date_from'):
            filters['date_from'] = datetime.fromisoformat(data['filter_date_from']).date()
        if data.get('filter_date_to'):
            filters['date_to'] = datetime.fromisoformat(data['filter_date_to']).date()
        
        # Debug logging
        app.logger.info(f"Inquire chat - User: {current_user.username}, Query: '{user_message}', Filters: {filters}")
        
        # Capture user context before generator to avoid current_user being None
        user_id = current_user.id
        user_name = current_user.name if current_user.name else "the user"
        user_title = current_user.job_title if current_user.job_title else "professional"
        user_company = current_user.company if current_user.company else "their organization"
        user_output_language = current_user.output_language if current_user.output_language else None
        
        # Enhanced query processing with enrichment and debugging
        def create_status_response(status, message):
            """Helper to create SSE status updates"""
            return f"data: {json.dumps({'status': status, 'message': message})}\n\n"
        
        def generate_enhanced_chat():
            # Explicitly reference outer scope variables
            nonlocal user_id, user_name, user_title, user_company, user_output_language, data, filters
            
            try:
                # Send initial status
                yield create_status_response('processing', 'Analyzing your query...')
                
                # Step 1: Router - Determine if RAG lookup is needed
                router_prompt = f"""Analyze this user query to determine if it requires searching through transcription content or if it's a simple formatting/clarification request.

User query: "{user_message}"

Respond with ONLY "RAG" if the query requires searching transcriptions (asking about content, conversations, specific information from recordings).
Respond with ONLY "DIRECT" if it's a formatting request, clarification about previous responses, or doesn't require searching transcriptions.

Examples:
- "What did Beth say about the budget?" → RAG
- "Can you format this in separate headings?" → DIRECT  
- "Who mentioned the timeline?" → RAG
- "Make this more structured" → DIRECT"""

                try:
                    router_response = call_llm_completion(
                        messages=[
                            {"role": "system", "content": "You are a query router. Respond with only 'RAG' or 'DIRECT'."},
                            {"role": "user", "content": router_prompt}
                        ],
                        temperature=0.1,
                        max_tokens=10
                    )
                    
                    route_decision = router_response.choices[0].message.content.strip().upper()
                    app.logger.info(f"Router decision: {route_decision}")
                    
                    if route_decision == "DIRECT":
                        # Direct response without RAG lookup
                        yield create_status_response('responding', 'Generating direct response...')
                        
                        direct_prompt = f"""You are assisting {user_name}. Respond to their request directly using proper markdown formatting.

User request: "{user_message}"

Previous conversation context (if relevant):
{json.dumps(message_history[-2:] if message_history else [])}

Use proper markdown formatting including headings (##), bold (**text**), bullet points (-), etc."""

                        stream = call_llm_completion(
                            messages=[
                                {"role": "system", "content": direct_prompt},
                                {"role": "user", "content": user_message}
                            ],
                            temperature=0.7,
                            max_tokens=int(os.environ.get("CHAT_MAX_TOKENS", "2000")),
                            stream=True
                        )
                        
                        # Use helper function to process streaming with thinking tag support
                        for response in process_streaming_with_thinking(stream):
                            yield response
                        return
                        
                except Exception as e:
                    app.logger.warning(f"Router failed, defaulting to RAG: {e}")
                
                # Step 2: Query enrichment - generate better search terms based on user intent
                yield create_status_response('enriching', 'Enriching search query...')
                
                # Use captured user context for personalized search terms
                
                enrichment_prompt = f"""You are a query enhancement assistant. Given a user's question about transcribed meetings/recordings, generate 3-5 alternative search terms or phrases that would help find relevant content in a semantic search system.

User context:
- Name: {user_name}
- Title: {user_title}  
- Company: {user_company}

User question: "{user_message}"
Available context: Transcribed meetings and recordings with speakers: {', '.join(data.get('filter_speakers', []))}.

Generate search terms that would find relevant content. Focus on:
1. Key concepts and topics using the user's actual name instead of generic terms like "me"
2. Specific terminology that might be used in their professional context
3. Alternative phrasings of the question with proper names
4. Related terms that might appear in transcripts from their meetings

Examples:
- Instead of "what Beth told me" use "what Beth told {user_name}"
- Instead of "my last conversation" use "{user_name}'s conversation"
- Use their job title and company context when relevant

Respond with only a JSON array of strings: ["term1", "term2", "term3", ...]"""
                
                try:
                    enrichment_response = call_llm_completion(
                        messages=[
                            {"role": "system", "content": "You are a query enhancement assistant. Respond only with valid JSON arrays of search terms."},
                            {"role": "user", "content": enrichment_prompt}
                        ],
                        temperature=0.3,
                        max_tokens=200
                    )
                    
                    enriched_terms = json.loads(enrichment_response.choices[0].message.content.strip())
                    app.logger.info(f"Enriched search terms: {enriched_terms}")
                    
                    # Combine original query with enriched terms for search
                    search_queries = [user_message] + enriched_terms[:3]  # Use original + top 3 enriched terms
                    
                except Exception as e:
                    app.logger.warning(f"Query enrichment failed, using original query: {e}")
                    search_queries = [user_message]
                
                # Step 2: Semantic search with multiple queries
                yield create_status_response('searching', 'Searching transcriptions...')
                
                all_chunks = []
                seen_chunk_ids = set()
                
                for query in search_queries:
                    with app.app_context():
                        chunk_results = semantic_search_chunks(user_id, query, filters, 8)
                        app.logger.info(f"Search query '{query}' returned {len(chunk_results)} chunks")
                    
                    for chunk, similarity in chunk_results:
                        if chunk and chunk.id not in seen_chunk_ids:
                            all_chunks.append((chunk, similarity))
                            seen_chunk_ids.add(chunk.id)
                
                # Sort by similarity and take top results
                all_chunks.sort(key=lambda x: x[1], reverse=True)
                chunk_results = all_chunks[:data.get('context_chunks', 8)]
                
                app.logger.info(f"Final chunk results: {len(chunk_results)} chunks with similarities: {[f'{s:.3f}' for _, s in chunk_results]}")
                
                # Step 2.5: Auto-detect mentioned speakers and apply filters if needed
                with app.app_context():
                    # Get available speakers
                    recordings_with_participants = Recording.query.filter_by(user_id=user_id).filter(
                        Recording.participants.isnot(None),
                        Recording.participants != ''
                    ).all()
                    
                    available_speakers = set()
                    for recording in recordings_with_participants:
                        if recording.participants:
                            participants = [p.strip() for p in recording.participants.split(',') if p.strip()]
                            available_speakers.update(participants)
                    
                    # Check if any speakers are mentioned in the user query but missing from results
                    mentioned_speakers = []
                    for speaker in available_speakers:
                        if speaker.lower() in user_message.lower():
                            # Check if this speaker appears in current chunk results
                            speaker_in_results = False
                            for chunk, _ in chunk_results:
                                if chunk and (
                                    (chunk.speaker_name and speaker.lower() in chunk.speaker_name.lower()) or
                                    (chunk.recording and chunk.recording.participants and speaker.lower() in chunk.recording.participants.lower())
                                ):
                                    speaker_in_results = True
                                    break
                            
                            if not speaker_in_results:
                                mentioned_speakers.append(speaker)
                    
                    # If we found mentioned speakers not in results, automatically apply speaker filter
                    if mentioned_speakers and not data.get('filter_speakers'):  # Only if no speaker filter already applied
                        app.logger.info(f"Auto-detected mentioned speakers not in results: {mentioned_speakers}")
                        yield create_status_response('filtering', f'Detected mention of {", ".join(mentioned_speakers)}, applying speaker filter...')
                        
                        # Apply automatic speaker filter
                        auto_filters = filters.copy()
                        auto_filters['speaker_names'] = mentioned_speakers
                        
                        # Re-run semantic search with speaker filter
                        auto_filtered_chunks = []
                        auto_filtered_seen_ids = set()
                        
                        for query in search_queries:
                            with app.app_context():
                                auto_filtered_results = semantic_search_chunks(user_id, query, auto_filters, data.get('context_chunks', 8))
                                app.logger.info(f"Auto-filtered search for '{query}' with speakers {mentioned_speakers} returned {len(auto_filtered_results)} chunks")
                            
                            for chunk, similarity in auto_filtered_results:
                                if chunk and chunk.id not in auto_filtered_seen_ids:
                                    auto_filtered_chunks.append((chunk, similarity))
                                    auto_filtered_seen_ids.add(chunk.id)
                        
                        # If auto-filter found better results, use them
                        if len(auto_filtered_chunks) > 0:
                            auto_filtered_chunks.sort(key=lambda x: x[1], reverse=True)
                            chunk_results = auto_filtered_chunks[:data.get('context_chunks', 8)]
                            app.logger.info(f"Auto speaker filter found {len(chunk_results)} relevant chunks, using filtered results")
                            filters = auto_filters  # Update filters for context building
                
                # Step 3: Evaluate results and re-query if needed
                if len(chunk_results) < 2:  # If we got very few results, try a broader search
                    yield create_status_response('requerying', 'Expanding search scope...')
                    
                    # Try without speaker filter if it was applied
                    broader_filters = filters.copy()
                    if 'speaker_names' in broader_filters:
                        del broader_filters['speaker_names']
                        app.logger.info("Retrying search without speaker filter...")
                        
                        for query in search_queries:
                            with app.app_context():
                                chunk_results_broader = semantic_search_chunks(user_id, query, broader_filters, 6)
                            for chunk, similarity in chunk_results_broader:
                                if chunk and chunk.id not in seen_chunk_ids:
                                    all_chunks.append((chunk, similarity))
                                    seen_chunk_ids.add(chunk.id)
                        
                        # Re-sort and limit
                        all_chunks.sort(key=lambda x: x[1], reverse=True)
                        chunk_results = all_chunks[:data.get('context_chunks', 8)]
                        app.logger.info(f"Broader search returned {len(chunk_results)} total chunks")
                
                # Build context from retrieved chunks
                yield create_status_response('contextualizing', 'Building context...')
                
                # Group chunks by recording and organize properly
                recording_chunks = {}
                recording_ids_in_context = set()
                
                for chunk, similarity in chunk_results:
                    if not chunk or not chunk.recording:
                        continue
                    recording_id = chunk.recording.id
                    recording_ids_in_context.add(recording_id)
                    
                    if recording_id not in recording_chunks:
                        recording_chunks[recording_id] = {
                            'recording': chunk.recording,
                            'chunks': []
                        }
                    
                    recording_chunks[recording_id]['chunks'].append({
                        'chunk': chunk,
                        'similarity': similarity
                    })
                
                # Build organized context pieces
                context_pieces = []
                
                for recording_id, data in recording_chunks.items():
                    recording = data['recording']
                    chunks = data['chunks']
                    
                    # Sort chunks by their index to maintain chronological order
                    chunks.sort(key=lambda x: x['chunk'].chunk_index)
                    
                    # Build recording header with complete metadata
                    header = f"=== {recording.title} [Recording ID: {recording_id}] ==="
                    if recording.meeting_date:
                        header += f" ({recording.meeting_date})"
                    
                    # Add participants information
                    if recording.participants:
                        participants_list = [p.strip() for p in recording.participants.split(',') if p.strip()]
                        header += f"\\nParticipants: {', '.join(participants_list)}"
                    
                    context_piece = header + "\\n\\n"
                    
                    # Process chunks and detect non-continuity
                    prev_chunk_index = None
                    for chunk_data in chunks:
                        chunk = chunk_data['chunk']
                        similarity = chunk_data['similarity']
                        
                        # Check for non-continuity
                        if prev_chunk_index is not None and chunk.chunk_index != prev_chunk_index + 1:
                            context_piece += "\\n[... gap in transcript - non-consecutive chunks ...]\\n\\n"
                        
                        # Add speaker information if available
                        speaker_info = ""
                        if chunk.speaker_name:
                            speaker_info = f"{chunk.speaker_name}: "
                        elif chunk.start_time is not None:
                            speaker_info = f"[{chunk.start_time:.1f}s]: "
                        
                        # Add timing info if available
                        timing_info = ""
                        if chunk.start_time is not None and chunk.end_time is not None:
                            timing_info = f" [{chunk.start_time:.1f}s-{chunk.end_time:.1f}s]"
                        
                        context_piece += f"{speaker_info}{chunk.content}{timing_info} (similarity: {similarity:.3f})\\n\\n"
                        prev_chunk_index = chunk.chunk_index
                    
                    context_pieces.append(context_piece)
                
                app.logger.info(f"Built context from {len(chunk_results)} chunks across {len(recording_chunks)} recordings")
                
                # Generate response
                yield create_status_response('responding', 'Generating response...')
                
                # Prepare system prompt
                language_instruction = f"Please provide all your responses in {user_output_language}." if user_output_language else ""
                
                # Build filter description for context
                filter_description = []
                with app.app_context():
                    if data.get('filter_tags'):
                        tag_names = [tag.name for tag in Tag.query.filter(Tag.id.in_(data['filter_tags'])).all()]
                        filter_description.append(f"tags: {', '.join(tag_names)}")
                if data.get('filter_speakers'):
                    filter_description.append(f"speakers: {', '.join(data['filter_speakers'])}")
                if data.get('filter_date_from') or data.get('filter_date_to'):
                    date_range = []
                    if data.get('filter_date_from'):
                        date_range.append(f"from {data['filter_date_from']}")
                    if data.get('filter_date_to'):
                        date_range.append(f"to {data['filter_date_to']}")
                    filter_description.append(f"dates: {' '.join(date_range)}")
                
                filter_text = f" (filtered by {'; '.join(filter_description)})" if filter_description else ""
                
                context_text = "\n\n".join(context_pieces) if context_pieces else "No relevant context found."
                
                # Get transcript length limit setting and available speakers
                with app.app_context():
                    transcript_limit = SystemSetting.get_setting('transcript_length_limit', 30000)
                    
                    # Get all available speakers for this user
                    recordings_with_participants = Recording.query.filter_by(user_id=user_id).filter(
                        Recording.participants.isnot(None),
                        Recording.participants != ''
                    ).all()
                    
                    available_speakers = set()
                    for recording in recordings_with_participants:
                        if recording.participants:
                            participants = [p.strip() for p in recording.participants.split(',') if p.strip()]
                            available_speakers.update(participants)
                    
                    available_speakers = sorted(list(available_speakers))
                
                system_prompt = f"""You are a professional meeting and audio transcription analyst assisting {user_name}, who is a(n) {user_title} at {user_company}. {language_instruction}

You are analyzing transcriptions from multiple recordings{filter_text}. The following context has been retrieved based on semantic similarity to the user's question:

<<start context>>
{context_text}
<<end context>>

The system has automatically analyzed your query and retrieved the most relevant context from your transcriptions. The search returned {len(chunk_results)} chunks from {len(recording_ids_in_context)} recording(s).

**Available speakers in your recordings**: {', '.join(available_speakers) if available_speakers else 'None available'}

**Recording IDs in context**: {list(recording_ids_in_context)}

IMPORTANT FORMATTING INSTRUCTIONS:
You MUST use proper markdown formatting in your responses. Structure your response as follows:

1. **Always use markdown syntax** - Use `#`, `##`, `###` for headings, `**bold**`, `*italic*`, `-` for lists, etc.
2. Start with a brief summary or preamble if helpful
3. Organize information by source transcript using clear markdown headings
4. Use the format: `## [Recording Title] - [Date if available]` 
5. Under each heading, provide the relevant information from that specific recording using bullet points and formatting
6. Include speaker names when referring to specific statements using **bold** formatting
7. Use bullet points (`-`) and sub-bullets for organizing information clearly

**Required Example Structure:**
Brief summary with **key points** highlighted...

## Meeting Discussion on Project Implementation - 2024-06-18
- **Speaker A** mentioned that "there's significant support needed for implementation"
- **Speaker B** confirmed the upcoming meeting with the technical team
- Key topics discussed:
  - Budget planning considerations
  - Timeline coordination needs

## Budget Planning Meeting - 2024-05-30  
- **Speaker A** reviewed the budget document
- **Speaker C** will approve the final version for submission
- Important details:
  - Budget represents approximately 1/3 of the project total
  - Coordination needed for upcoming milestones

Order your response with notes from the most recent meetings first. Always use proper markdown formatting and structure by source recording for maximum clarity and readability."""
        
                # Prepare messages array
                messages = [{"role": "system", "content": system_prompt}]
                if message_history:
                    messages.extend(message_history)
                messages.append({"role": "user", "content": user_message})

                # Enable streaming
                stream = call_llm_completion(
                    messages=messages,
                    temperature=0.7,
                    max_tokens=int(os.environ.get("CHAT_MAX_TOKENS", "2000")),
                    stream=True
                )
                
                # Buffer content to detect full transcript requests
                response_buffer = ""
                
                # Buffer content to detect full transcript requests
                response_buffer = ""
                content_buffer = ""
                in_thinking = False
                thinking_buffer = ""
                
                for chunk in stream:
                    content = chunk.choices[0].delta.content
                    if content:
                        response_buffer += content
                        content_buffer += content
                        
                        # Check if this is a full transcript request
                        if response_buffer.strip().startswith("REQUEST_FULL_TRANSCRIPT:"):
                            lines = response_buffer.split('\n')
                            request_line = lines[0].strip()
                            
                            if ':' in request_line:
                                try:
                                    recording_id = int(request_line.split(':')[1])
                                    app.logger.info(f"Agent requested full transcript for recording {recording_id}")
                                    
                                    # Fetch full transcript
                                    yield create_status_response('fetching', f'Retrieving full transcript for recording {recording_id}...')
                                    
                                    with app.app_context():
                                        recording = db.session.get(Recording, recording_id)
                                        if recording and recording.user_id == user_id and recording.transcription:
                                            # Apply transcript length limit
                                            if transcript_limit == -1:
                                                full_transcript = recording.transcription
                                            else:
                                                full_transcript = recording.transcription[:transcript_limit]
                                            
                                            # Add full transcript to context
                                            full_context = f"{context_text}\n\n<<FULL TRANSCRIPT - {recording.title}>>\n{full_transcript}\n<<END FULL TRANSCRIPT>>"
                                            
                                            # Update system prompt with full transcript
                                            updated_system_prompt = system_prompt.replace(
                                                f"<<start context>>\n{context_text}\n<<end context>>",
                                                f"<<start context>>\n{full_context}\n<<end context>>"
                                            )
                                            
                                            # Create new messages with updated context
                                            updated_messages = [{"role": "system", "content": updated_system_prompt}]
                                            if message_history:
                                                updated_messages.extend(message_history)
                                            updated_messages.append({"role": "user", "content": user_message})
                                            
                                            # Generate new response with full context
                                            yield create_status_response('responding', 'Analyzing full transcript...')
                                            
                                            new_stream = call_llm_completion(
                                                messages=updated_messages,
                                                temperature=0.7,
                                                max_tokens=int(os.environ.get("CHAT_MAX_TOKENS", "2000")),
                                                stream=True
                                            )
                                            
                                            # Use helper function to process streaming with thinking tag support
                                            for response in process_streaming_with_thinking(new_stream):
                                                yield response
                                            return
                                        else:
                                            # Recording not found or no permission
                                            error_msg = f"\n\nError: Unable to access full transcript for recording {recording_id}. Recording may not exist or you may not have permission."
                                            yield f"data: {json.dumps({'delta': error_msg})}\n\n"
                                            yield f"data: {json.dumps({'end_of_stream': True})}\n\n"
                                            return
                                            
                                except (ValueError, IndexError):
                                    app.logger.warning(f"Invalid transcript request format: {request_line}")
                                    # Continue with normal streaming
                                    pass
                        
                        # Process the buffer to detect and handle thinking tags
                        while True:
                            if not in_thinking:
                                # Look for opening thinking tag
                                think_start = re.search(r'<think(?:ing)?>', content_buffer, re.IGNORECASE)
                                if think_start:
                                    # Send any content before the thinking tag
                                    before_thinking = content_buffer[:think_start.start()]
                                    if before_thinking:
                                        yield f"data: {json.dumps({'delta': before_thinking})}\n\n"
                                    
                                    # Start capturing thinking content
                                    in_thinking = True
                                    content_buffer = content_buffer[think_start.end():]
                                    thinking_buffer = ""
                                else:
                                    # No thinking tag found, send accumulated content
                                    if content_buffer:
                                        yield f"data: {json.dumps({'delta': content_buffer})}\n\n"
                                    content_buffer = ""
                                    break
                            else:
                                # We're inside a thinking tag, look for closing tag
                                think_end = re.search(r'</think(?:ing)?>', content_buffer, re.IGNORECASE)
                                if think_end:
                                    # Capture thinking content up to the closing tag
                                    thinking_buffer += content_buffer[:think_end.start()]
                                    
                                    # Send the thinking content as a special type
                                    if thinking_buffer.strip():
                                        yield f"data: {json.dumps({'thinking': thinking_buffer.strip()})}\n\n"
                                    
                                    # Continue processing after the closing tag
                                    in_thinking = False
                                    content_buffer = content_buffer[think_end.end():]
                                    thinking_buffer = ""
                                else:
                                    # Still inside thinking tag, accumulate content
                                    thinking_buffer += content_buffer
                                    content_buffer = ""
                                    break
                
                # Handle any remaining content
                if in_thinking and thinking_buffer:
                    # Unclosed thinking tag - send as thinking content
                    yield f"data: {json.dumps({'thinking': thinking_buffer.strip()})}\n\n"
                elif content_buffer:
                    # Regular content
                    yield f"data: {json.dumps({'delta': content_buffer})}\n\n"
                
                yield f"data: {json.dumps({'end_of_stream': True})}\n\n"
                
            except Exception as e:
                app.logger.error(f"Error in enhanced chat generation: {e}")
                yield f"data: {json.dumps({'error': str(e)})}\n\n"
        
        return Response(generate_enhanced_chat(), mimetype='text/event-stream')
        
    except Exception as e:
        app.logger.error(f"Error in inquire chat endpoint: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/recording/<int:recording_id>/process_chunks', methods=['POST'])
@login_required
def process_recording_chunks_endpoint(recording_id):
    """Process chunks for a specific recording."""
    try:
        recording = db.session.get(Recording, recording_id)
        if not recording:
            return jsonify({'error': 'Recording not found'}), 404
        
        if recording.user_id != current_user.id:
            return jsonify({'error': 'Permission denied'}), 403
        
        success = process_recording_chunks(recording_id)
        if success:
            return jsonify({'message': 'Chunks processed successfully'})
        else:
            return jsonify({'error': 'Failed to process chunks'}), 500
            
    except Exception as e:
        app.logger.error(f"Error in process chunks endpoint: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/inquire/available_filters', methods=['GET'])
@login_required
def get_available_filters():
    """Get available filter options for the user."""
    if not ENABLE_INQUIRE_MODE:
        return jsonify({'error': 'Inquire mode is not enabled'}), 403
    try:
        # Get user's tags
        tags = Tag.query.filter_by(user_id=current_user.id).all()
        
        # Get unique speakers from user's recordings participants field
        recordings_with_participants = Recording.query.filter_by(user_id=current_user.id).filter(
            Recording.participants.isnot(None),
            Recording.participants != ''
        ).all()
        
        speaker_names = set()
        for recording in recordings_with_participants:
            if recording.participants:
                # Split participants by comma and clean up
                participants = [p.strip() for p in recording.participants.split(',') if p.strip()]
                speaker_names.update(participants)
        
        speaker_names = sorted(list(speaker_names))
        
        # Get user's recordings for recording-specific filtering
        recordings = Recording.query.filter_by(user_id=current_user.id).filter(
            Recording.status == 'COMPLETED'
        ).order_by(Recording.created_at.desc()).all()
        
        return jsonify({
            'tags': [tag.to_dict() for tag in tags],
            'speakers': speaker_names,
            'recordings': [{'id': r.id, 'title': r.title, 'meeting_date': f"{r.meeting_date.isoformat()}T00:00:00" if r.meeting_date else None} for r in recordings]
        })
        
    except Exception as e:
        app.logger.error(f"Error getting available filters: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/admin/migrate_recordings', methods=['POST'])
@login_required
def migrate_existing_recordings_api():
    """API endpoint to migrate existing recordings for inquire mode (admin only)."""
    if not current_user.is_admin:
        return jsonify({'error': 'Unauthorized. Admin access required.'}), 403
    
    try:
        # Count recordings that need processing
        completed_recordings = Recording.query.filter_by(status='COMPLETED').all()
        recordings_needing_processing = []
        
        for recording in completed_recordings:
            if recording.transcription:  # Has transcription
                chunk_count = TranscriptChunk.query.filter_by(recording_id=recording.id).count()
                if chunk_count == 0:  # No chunks yet
                    recordings_needing_processing.append(recording)
        
        if len(recordings_needing_processing) == 0:
            return jsonify({
                'success': True,
                'message': 'All recordings are already processed for inquire mode',
                'processed': 0,
                'total': len(completed_recordings)
            })
        
        # Process in small batches to avoid timeout
        batch_size = min(5, len(recordings_needing_processing))  # Process max 5 at a time
        processed = 0
        errors = 0
        
        for i in range(min(batch_size, len(recordings_needing_processing))):
            recording = recordings_needing_processing[i]
            try:
                success = process_recording_chunks(recording.id)
                if success:
                    processed += 1
                else:
                    errors += 1
            except Exception as e:
                app.logger.error(f"Error processing recording {recording.id} for migration: {e}")
                errors += 1
        
        remaining = max(0, len(recordings_needing_processing) - batch_size)
        
        return jsonify({
            'success': True,
            'message': f'Processed {processed} recordings. {remaining} remaining.',
            'processed': processed,
            'errors': errors,
            'remaining': remaining,
            'total': len(recordings_needing_processing)
        })
        
    except Exception as e:
        app.logger.error(f"Error in migration API: {e}")
        return jsonify({'error': str(e)}), 500


# --- Auto-Processing File Monitor Integration ---
def initialize_file_monitor():
    """Initialize file monitor after app is fully loaded to avoid circular imports."""
    try:
        # Import here to avoid circular imports
        import src.file_monitor as file_monitor
        file_monitor.start_file_monitor()
        app.logger.info("File monitor initialization completed")
    except Exception as e:
        app.logger.warning(f"File monitor initialization failed: {e}")

def get_file_monitor_functions():
    """Get file monitor functions, handling import errors gracefully."""
    try:
        import src.file_monitor as file_monitor
        return file_monitor.start_file_monitor, file_monitor.stop_file_monitor, file_monitor.get_file_monitor_status
    except ImportError as e:
        app.logger.warning(f"File monitor not available: {e}")
        
        # Create stub functions if file_monitor is not available
        def start_file_monitor():
            pass
        def stop_file_monitor():
            pass
        def get_file_monitor_status():
            return {'running': False, 'error': 'File monitor module not available'}
        
        return start_file_monitor, stop_file_monitor, get_file_monitor_status

# --- Auto-Processing API Endpoints ---
@app.route('/admin/auto-process/status', methods=['GET'])
@login_required
def admin_get_auto_process_status():
    """Get the status of the automated file processing system."""
    if not current_user.is_admin:
        return jsonify({'error': 'Unauthorized'}), 403
    
    try:
        _, _, get_file_monitor_status = get_file_monitor_functions()
        status = get_file_monitor_status()
        
        # Add configuration info
        config = {
            'enabled': os.environ.get('ENABLE_AUTO_PROCESSING', 'false').lower() == 'true',
            'watch_directory': os.environ.get('AUTO_PROCESS_WATCH_DIR', '/data/auto-process'),
            'check_interval': int(os.environ.get('AUTO_PROCESS_CHECK_INTERVAL', '30')),
            'mode': os.environ.get('AUTO_PROCESS_MODE', 'admin_only'),
            'default_username': os.environ.get('AUTO_PROCESS_DEFAULT_USERNAME')
        }
        
        return jsonify({
            'status': status,
            'config': config
        })
        
    except Exception as e:
        app.logger.error(f"Error getting auto-process status: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/admin/auto-process/start', methods=['POST'])
@login_required
def admin_start_auto_process():
    """Start the automated file processing system."""
    if not current_user.is_admin:
        return jsonify({'error': 'Unauthorized'}), 403
    
    try:
        start_file_monitor, _, _ = get_file_monitor_functions()
        start_file_monitor()
        return jsonify({'success': True, 'message': 'Auto-processing started'})
    except Exception as e:
        app.logger.error(f"Error starting auto-process: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/admin/auto-process/stop', methods=['POST'])
@login_required
def admin_stop_auto_process():
    """Stop the automated file processing system."""
    if not current_user.is_admin:
        return jsonify({'error': 'Unauthorized'}), 403
    
    try:
        _, stop_file_monitor, _ = get_file_monitor_functions()
        stop_file_monitor()
        return jsonify({'success': True, 'message': 'Auto-processing stopped'})
    except Exception as e:
        app.logger.error(f"Error stopping auto-process: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/admin/auto-process/config', methods=['POST'])
@login_required
def admin_update_auto_process_config():
    """Update auto-processing configuration (requires restart)."""
    if not current_user.is_admin:
        return jsonify({'error': 'Unauthorized'}), 403
    
    try:
        data = request.json
        if not data:
            return jsonify({'error': 'No configuration data provided'}), 400
        
        # This endpoint would typically update environment variables or config files
        # For now, we'll just return the current config and note that restart is required
        return jsonify({
            'success': True, 
            'message': 'Configuration updated. Restart required to apply changes.',
            'note': 'Environment variables need to be updated manually and application restarted.'
        })
        
    except Exception as e:
        app.logger.error(f"Error updating auto-process config: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/admin/inquire/process-recordings', methods=['POST'])
@login_required
def admin_process_recordings_for_inquire():
    """Process all remaining recordings for inquire mode (chunk and embed them)."""
    if not current_user.is_admin:
        return jsonify({'error': 'Unauthorized'}), 403
    
    try:
        # Get optional parameters from request
        data = request.json or {}
        batch_size = data.get('batch_size', 10)
        max_recordings = data.get('max_recordings', None)
        
        # Find recordings that need processing
        completed_recordings = Recording.query.filter_by(status='COMPLETED').all()
        recordings_needing_processing = []
        
        for recording in completed_recordings:
            if recording.transcription:  # Has transcription
                chunk_count = TranscriptChunk.query.filter_by(recording_id=recording.id).count()
                if chunk_count == 0:  # No chunks yet
                    recordings_needing_processing.append(recording)
                    if max_recordings and len(recordings_needing_processing) >= max_recordings:
                        break
        
        total_to_process = len(recordings_needing_processing)
        
        if total_to_process == 0:
            return jsonify({
                'success': True,
                'message': 'All recordings are already processed for inquire mode.',
                'processed': 0,
                'total': 0
            })
        
        # Process recordings in batches
        processed = 0
        failed = []
        
        for recording in recordings_needing_processing:
            try:
                success = process_recording_chunks(recording.id)
                if success:
                    processed += 1
                    app.logger.info(f"Admin API: Processed chunks for recording: {recording.title} ({recording.id})")
                else:
                    failed.append({'id': recording.id, 'title': recording.title, 'reason': 'Processing returned false'})
            except Exception as e:
                app.logger.error(f"Admin API: Failed to process recording {recording.id}: {e}")
                failed.append({'id': recording.id, 'title': recording.title, 'reason': str(e)})
            
            # Commit after each batch
            if processed % batch_size == 0:
                db.session.commit()
        
        # Final commit
        db.session.commit()
        
        return jsonify({
            'success': True,
            'message': f'Processed {processed} out of {total_to_process} recordings.',
            'processed': processed,
            'total': total_to_process,
            'failed': failed
        })
        
    except Exception as e:
        app.logger.error(f"Error in admin process recordings endpoint: {e}")
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

@app.route('/admin/inquire/status', methods=['GET'])
@login_required  
def admin_inquire_status():
    """Get the status of recordings for inquire mode."""
    if not current_user.is_admin:
        return jsonify({'error': 'Unauthorized'}), 403
    
    try:
        # Count total completed recordings
        total_completed = Recording.query.filter_by(status='COMPLETED').count()
        
        # Count recordings with transcriptions
        recordings_with_transcriptions = Recording.query.filter(
            Recording.status == 'COMPLETED',
            Recording.transcription.isnot(None),
            Recording.transcription != ''
        ).count()
        
        # Count recordings that have been processed for inquire mode
        processed_recordings = db.session.query(Recording.id).join(
            TranscriptChunk, Recording.id == TranscriptChunk.recording_id
        ).distinct().count()
        
        # Count recordings that still need processing
        completed_recordings = Recording.query.filter_by(status='COMPLETED').all()
        need_processing = 0
        
        for recording in completed_recordings:
            if recording.transcription:  # Has transcription
                chunk_count = TranscriptChunk.query.filter_by(recording_id=recording.id).count()
                if chunk_count == 0:  # No chunks yet
                    need_processing += 1
        
        # Get total chunks and embeddings count
        total_chunks = TranscriptChunk.query.count()
        
        return jsonify({
            'total_completed_recordings': total_completed,
            'recordings_with_transcriptions': recordings_with_transcriptions,
            'processed_for_inquire': processed_recordings,
            'need_processing': need_processing,
            'total_chunks': total_chunks,
            'embeddings_available': EMBEDDINGS_AVAILABLE
        })
        
    except Exception as e:
        app.logger.error(f"Error getting inquire status: {e}")
        return jsonify({'error': str(e)}), 500

with app.app_context():
    # Set dynamic MAX_CONTENT_LENGTH based on database setting
    max_file_size_mb = SystemSetting.get_setting('max_file_size_mb', 250)
    app.config['MAX_CONTENT_LENGTH'] = max_file_size_mb * 1024 * 1024
    app.logger.info(f"Set MAX_CONTENT_LENGTH to {max_file_size_mb}MB from database setting")

    # Initialize file monitor after app setup
    initialize_file_monitor()

if __name__ == '__main__':
    # Consider using waitress or gunicorn for production
    # waitress-serve --host 0.0.0.0 --port 8899 app:app
    # For development:
    app.run(host='0.0.0.0', port=8899, debug=True) # Set debug=False if thread issues arise
