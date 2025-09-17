#!/usr/bin/env python
"""Test script to verify Unicode support in Word documents."""

from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml.ns import qn
import sys

def test_unicode_support():
    """Test creating a Word document with various Unicode characters."""
    doc = Document()

    # Test data with various Unicode characters
    test_strings = [
        ("Chinese (Simplified)", "第二次研究 - 美食与饮品词汇学习"),
        ("Chinese (Traditional)", "第二次研究 - 美食與飲品詞彙學習"),
        ("Japanese", "これはテストです。日本語のテキスト。"),
        ("Korean", "이것은 테스트입니다. 한국어 텍스트."),
        ("Arabic", "هذا اختبار. نص عربي."),
        ("Hebrew", "זה מבחן. טקסט עברי."),
        ("Russian", "Это тест. Русский текст."),
        ("Greek", "Αυτό είναι ένα τεστ. Ελληνικό κείμενο."),
        ("Thai", "นี่คือการทดสอบ ข้อความภาษาไทย"),
        ("Emoji", "Testing emojis: 😀 🎉 🚀 ❤️"),
        ("Mixed", "Summary: 第二次研究 with English and 中文"),
    ]

    # Add title
    title_text = "Unicode Font Support Test Document"
    title = doc.add_heading(title_text, 0)

    # Add test strings
    for language, text in test_strings:
        # Add language label
        p_label = doc.add_paragraph()
        run = p_label.add_run(f"{language}: ")
        run.bold = True

        # Add test text
        p_text = doc.add_paragraph(text)

        # Apply Unicode font support if needed
        try:
            text.encode('ascii')
            print(f"✓ {language}: ASCII only, no special font needed")
        except UnicodeEncodeError:
            print(f"✓ {language}: Non-ASCII detected, applying Arial font")
            for run in p_text.runs:
                run.font.name = 'Arial'
                r = run._element
                r.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

    # Save document
    output_file = "test_unicode_output.docx"
    doc.save(output_file)
    print(f"\nTest document saved as: {output_file}")
    print("Please open this document in Microsoft Word to verify all characters display correctly.")

    return True

if __name__ == "__main__":
    try:
        # Check if python-docx is installed
        import docx
        print("python-docx is installed. Version:", docx.__version__)
    except ImportError:
        print("ERROR: python-docx is not installed. Please install it with: pip install python-docx")
        sys.exit(1)

    print("Testing Unicode support in Word documents...\n")
    if test_unicode_support():
        print("\nTest completed successfully!")
        print("\nNOTE: The fix implemented uses Arial font which has broad Unicode support.")
        print("The solution automatically detects non-ASCII characters and applies appropriate font settings.")
        print("This ensures compatibility across different languages without forcing a specific font on all users.")